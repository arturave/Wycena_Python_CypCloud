#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
wycena_optimized.py — szybka analiza XLSX i generowanie WZ (DOCX) z prostym GUI.

- Adnotacje typów (PEP 484)
- Docstringi kompatybilne z Google/NumPy
- Logowanie
- Cache cennika (lru_cache)
- Obsługa błędów z oknami komunikatów (messagebox)

Uruchom:
    python wycena_optimized.py
"""
from __future__ import annotations

import hashlib
import logging
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from typing import List, Optional

import tkinter as tk
from tkinter import filedialog, messagebox, ttk

from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------- Logging ----------
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s"
)
LOG = logging.getLogger("wycena_optimized")


# ---------- Data structures ----------
@dataclass(frozen=True)
class Item:
    """Single line item parsed from XLSX inputs.

    Attributes:
        lp: Ordinal number (1-based).
        symbol: Internal part code or SKU.
        name: Human-readable part name.
        qty: Quantity (integer).
        unit: Unit of measure (default 'szt.').
    """
    lp: int
    symbol: str
    name: str
    qty: int
    unit: str = "szt."


@dataclass
class AnalysisResult:
    """Container for analysis output.

    Attributes:
        items: List of parsed Item entries.
        total_qty: Sum of quantities across items.
        source_files: XLSX file paths used as inputs.
    """
    items: List[Item]
    total_qty: int
    source_files: List[Path]


# ---------- Utility helpers ----------
def _hash_file(path: Path) -> str:
    """Return SHA1 of file for cache invalidation."""
    import hashlib
    h = hashlib.sha1()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


@lru_cache(maxsize=8)
def load_price_list(price_path: str, file_hash: Optional[str] = None) -> dict:
    """Load and cache a price list (CSV or XLSX).

    The second parameter `file_hash` is only for cache-busting when a file changes.
    If the file does not exist, returns an empty dict.

    Args:
        price_path: Path to CSV/XLSX with columns [symbol, price].
        file_hash: Optional content hash, ignored but contributes to cache key.

    Returns:
        Mapping symbol -> price (float).
    """
    p = Path(price_path)
    if not p.exists():
        LOG.warning("Price list not found: %s", p)
        return {}

    prices = {}
    if p.suffix.lower() == ".csv":
        text = p.read_text(encoding="utf-8", errors="ignore")
        for line in text.splitlines():
            parts = [x.strip() for x in line.split(",")]
            if len(parts) >= 2:
                sym, price = parts[0], parts[1].replace(",", ".")
                try:
                    prices[sym] = float(price)
                except ValueError:
                    continue
    else:
        # XLSX
        wb = load_workbook(p, read_only=True, data_only=True)
        ws = wb.active
        for row in ws.iter_rows(min_row=1, values_only=True):
            if not row or len(row) < 2:
                continue
            sym = str(row[0]).strip()
            try:
                price = float(str(row[1]).replace(",", "."))
            except Exception:
                continue
            prices[sym] = price
    LOG.info("Loaded %d prices from %s", len(prices), p.name)
    return prices


def _cell(ws: Worksheet, r: int, c: int) -> Optional[str]:
    v = ws.cell(r, c).value
    if v is None:
        return None
    return str(v).strip()


def parse_items_from_xlsx(path: Path) -> List[Item]:
    """Parse one XLSX file for [Lp, Symbol, Nazwa, Ilość] columns.

    The function scans the sheet header to find these columns by name,
    then parses all subsequent rows until an empty Lp or Symbol is encountered.

    Args:
        path: Path to the XLSX file.

    Returns:
        List of Item objects.
    """
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb.active

    # Detect header row by scanning first 20 rows for column names (case-insensitive)
    header_row_idx = None
    headers = {}
    targets = {"lp": "lp", "symbol": "symbol", "nazwa": "name", "ilość": "qty", "ilosc": "qty"}

    for r in range(1, 21):
        row_vals = [(_cell(ws, r, c) or "").lower() for c in range(1, 20)]
        for c, val in enumerate(row_vals, start=1):
            key = val.replace(" ", "").replace(".", "")
            if key in targets:
                headers[targets[key]] = c
        if {"lp", "symbol", "name", "qty"}.issubset(headers.keys()):
            header_row_idx = r
            break

    if header_row_idx is None:
        raise ValueError("Nie znaleziono wiersza nagłówków z kolumnami: Lp, Symbol, Nazwa, Ilość")

    items: List[Item] = []
    r = header_row_idx + 1
    lp_col, sym_col, name_col, qty_col = (
        headers["lp"],
        headers["symbol"],
        headers["name"],
        headers["qty"],
    )
    while True:
        lp_val = _cell(ws, r, lp_col)
        sym_val = _cell(ws, r, sym_col)
        name_val = _cell(ws, r, name_col)
        qty_val = _cell(ws, r, qty_col)
        r += 1

        if not (lp_val or sym_val or name_val or qty_val):
            # reached the end
            break
        if not sym_val:
            continue

        try:
            lp = int(str(lp_val).strip()) if lp_val else len(items) + 1
        except Exception:
            lp = len(items) + 1
        try:
            qty = int(float(str(qty_val).replace(",", ".") if qty_val else "0"))
        except Exception:
            qty = 0

        items.append(Item(lp=lp, symbol=str(sym_val), name=str(name_val or ""), qty=qty))

    return items


def analyze_folder(folder: Path) -> AnalysisResult:
    """Analyze a folder of XLSX files and aggregate items.

    Args:
        folder: Directory with XLSX files.

    Returns:
        AnalysisResult with all parsed items.
    """
    xlsx_files = sorted(folder.glob("*.xlsx"))
    if not xlsx_files:
        raise FileNotFoundError("Brak plików XLSX w wybranym folderze.")

    all_items: List[Item] = []
    for f in xlsx_files:
        try:
            items = parse_items_from_xlsx(f)
            all_items.extend(items)
        except Exception as e:
            LOG.warning("Pominięto %s (%s)", f.name, e)

    if not all_items:
        raise ValueError("Nie udało się odczytać pozycji z żadnego pliku XLSX.")

    total_qty = sum(x.qty for x in all_items)
    return AnalysisResult(items=all_items, total_qty=total_qty, source_files=xlsx_files)


def generate_wz_doc(
    out_path: Path,
    wz_number: str,
    issue_date: str,
    place: str,
    warehouse: str,
    issuer: dict,
    recipient: dict,
    items: List[Item],
) -> Path:
    """Generate a WZ DOCX with a minimal, performant layout.

    Returns:
        Path to saved DOCX.
    """
    doc = Document()

    # Margins
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Title
    title = doc.add_paragraph()
    run = title.add_run("WZ – Wydanie na Zewnątrz")
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["WZ nr:", "Data wystawienia:", "Miejsce wystawienia:", "Magazyn:"]
    values = [wz_number, issue_date, place, warehouse]
    for i in range(4):
        cell = meta.cell(0, i)
        r = cell.paragraphs[0].add_run(labels[i])
        r.bold = True
        meta.cell(1, i).text = values[i]

    doc.add_paragraph("")

    # Parties
    parties = doc.add_table(rows=1, cols=2)
    parties.style = "Table Grid"
    left, right = parties.cell(0, 0), parties.cell(0, 1)

    p = left.paragraphs[0].add_run("Wystawca (wydający):\n")
    p.bold = True
    left.add_paragraph(f"Nazwa: {issuer.get('name','')}")
    left.add_paragraph(f"Adres: {issuer.get('address','')}")
    left.add_paragraph(
        f"NIP: {issuer.get('nip','')}    REGON: {issuer.get('regon','')}    KRS: {issuer.get('krs','')}"
    )
    left.add_paragraph(f"Telefon: {issuer.get('phone','')}")

    p = right.paragraphs[0].add_run("Odbiorca:\n")
    p.bold = True
    right.add_paragraph(f"Nazwa: {recipient.get('name','')}")
    right.add_paragraph(f"Adres: {recipient.get('address','')}")
    if recipient.get("nip"):
        right.add_paragraph(f"NIP: {recipient.get('nip')}")
    if recipient.get("phone"):
        right.add_paragraph(f"Telefon: {recipient.get('phone')}")

    doc.add_paragraph("")

    # Items table
    headers = ["Lp.", "Symbol", "Nazwa towaru", "Ilość"]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True

    for it in items:
        row = tbl.add_row().cells
        row[0].text = str(it.lp)
        row[1].text = it.symbol
        row[2].text = it.name
        row[3].text = str(it.qty)

    doc.add_paragraph("")
    sumtbl = doc.add_table(rows=1, cols=2)
    sumtbl.style = "Table Grid"
    sumtbl.cell(0, 0).text = f"Razem pozycji: {len(items)}"
    sumtbl.cell(0, 1).text = f"Suma ilości (szt.): {sum(x.qty for x in items)}"

    doc.save(out_path)
    return out_path


# ---------- GUI ----------
class App(tk.Tk):
    """Minimal GUI to run analysis and create WZ quickly."""

    def __init__(self) -> None:
        super().__init__()
        self.title("Wycena / WZ – szybka analiza")
        self.geometry("720x360")

        self.var_folder = tk.StringVar()
        self.var_price = tk.StringVar()

        frm = ttk.Frame(self, padding=12)
        frm.pack(fill="both", expand=True)

        ttk.Label(frm, text="Folder z XLSX:").grid(row=0, column=0, sticky="w")
        ttk.Entry(frm, textvariable=self.var_folder, width=60).grid(row=0, column=1, sticky="we")
        ttk.Button(frm, text="Wybierz…", command=self.choose_folder).grid(row=0, column=2, padx=6)

        ttk.Label(frm, text="Cennik (opcjonalnie):").grid(row=1, column=0, sticky="w")
        ttk.Entry(frm, textvariable=self.var_price, width=60).grid(row=1, column=1, sticky="we")
        ttk.Button(frm, text="Wybierz…", command=self.choose_price).grid(row=1, column=2, padx=6)

        ttk.Separator(frm).grid(row=2, column=0, columnspan=3, sticky="ew", pady=8)

        ttk.Button(frm, text="Analizuj i wygeneruj WZ", command=self.run_all).grid(
            row=3, column=0, columnspan=3, pady=8
        )

        frm.columnconfigure(1, weight=1)

    def choose_folder(self) -> None:
        path = filedialog.askdirectory()
        if path:
            self.var_folder.set(path)

    def choose_price(self) -> None:
        path = filedialog.askopenfilename(filetypes=[("CSV/XLSX", "*.csv *.xlsx")])
        if path:
            self.var_price.set(path)

    def run_all(self) -> None:
        folder = Path(self.var_folder.get())
        if not folder.exists():
            messagebox.showerror("Błąd", "Nie wybrano poprawnego folderu z XLSX.")
            return
        try:
            result = analyze_folder(folder)
        except Exception as e:
            messagebox.showerror("Błąd analizy", str(e))
            return

        # optional price list
        price_path = self.var_price.get().strip()
        if price_path:
            try:
                load_price_list.cache_clear()
                load_price_list(price_path, _hash_file(Path(price_path)))
            except Exception as e:
                messagebox.showwarning("Cennik", f"Nie udało się wczytać cennika: {e}")

        # Dummy issuer/recipient – replace in your workflow
        issuer = {
            "name": "LP Konstal sp. z o.o.",
            "address": "Pisarzowice 203b, 59-800 Lubań, PL",
            "nip": "101-000-45-08",
            "regon": "021470270",
            "krs": "0000379723",
            "phone": "+48 75 775 93 16",
        }
        recipient = {"name": "Odbiorca", "address": "Adres", "nip": "", "phone": ""}

        out = folder / "WZ_wygenerowane.docx"
        try:
            generate_wz_doc(
                out_path=out,
                wz_number="AUTO/GEN",
                issue_date="2025-08-22",
                place="Lubań",
                warehouse="Magazyn",
                issuer=issuer,
                recipient=recipient,
                items=result.items,
            )
        except Exception as e:
            messagebox.showerror("Błąd generowania DOCX", str(e))
            return

        messagebox.showinfo(
            "Gotowe",
            f"Wygenerowano WZ ({len(result.items)} pozycji, {result.total_qty} szt.).\\n\\nPlik: {out}",
        )


if __name__ == "__main__":
    App().mainloop()
