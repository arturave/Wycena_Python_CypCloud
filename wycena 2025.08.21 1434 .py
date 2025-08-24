#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import datetime
from datetime import timedelta
import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from openpyxl import load_workbook
# DOCX (raport)
from docx import Document
from docx.shared import Inches, RGBColor, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io  # Do konwersji obrazu na bajty, jeśli potrzeba do raportu
from PIL import Image, ImageTk
import subprocess

# Locale do formatowania PLN
import locale
try:
    locale.setlocale(locale.LC_ALL, 'pl_PL.UTF-8')
except Exception:
    pass

def format_pln(value):
    try:
        s = locale.format_string('%.2f', float(value), grouping=True)
        return s.replace('.', ',')
    except Exception:
        try:
            return f"{float(value):.2f}".replace('.', ',')
        except Exception:
            return "0,00"

def sanitize_filename(name):
    for ch in r'< > : " / \ | ? *':
        name = name.replace(ch, '_')
    return name

def _norm_s(s):
    return (str(s).strip().upper() if s is not None else "")

def _parse_float(val):
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip().replace(" ", "").replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None

def _map_gas_to_key(gas_raw: str) -> str:
    g = _norm_s(gas_raw)
    if g in {"NITROGEN", "AZOT", "氮气", "N"}:
        return "N"
    if g in {"OXYGEN", "TLEN", "氧气", "O"}:
        return "O"
    return ""

# ---- cenniki ----
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MATERIALS_FILE = os.path.join(SCRIPT_DIR, "materials prices.xlsx")
CUTTING_FILE   = os.path.join(SCRIPT_DIR, "cutting prices.xlsx")
material_prices = {}  # (MAT, THK)-> PLN/kg
cutting_prices  = {}  # (THK, MAT, GAS)-> PLN/m
_mat_set, _thk_set, _gas_set = set(), set(), set()

# ---- GUI ----
root = tk.Tk()
root.title("Generator Raportu Kosztów — MERGED (fixed)")
root.configure(bg="#2c2c2c")  # Dark background for modern look

# Use a modern ttk theme
style = ttk.Style(root)
style.theme_use('clam')  # Or 'alt', 'default', etc. for better visuals
style.configure("TLabel", foreground="white", background="#2c2c2c", font=("Arial", 10))
style.configure("TEntry", fieldbackground="#3c3c3c", foreground="white")
style.configure("TButton", background="#4c4c4c", foreground="white", borderwidth=0)
style.configure("TCombobox", fieldbackground="#3c3c3c", foreground="white")
style.configure("Treeview", background="#3c3c3c", foreground="white", fieldbackground="#3c3c3c", rowheight=80)  # Increased row height for better thumbnails
style.configure("Treeview.Heading", background="#4c4c4c", foreground="white")
style.map("TButton", background=[('active', '#5c5c5c')])
style.map("Treeview", background=[('selected', '#5c5c5c')])

folder_var   = tk.StringVar()
customer_var = tk.StringVar()
offer_var    = tk.StringVar()
date_var     = tk.StringVar(value=datetime.datetime.now().strftime("%Y-%m-%d"))
validity_var = tk.StringVar(value=(datetime.datetime.now() + timedelta(days=14)).strftime("%Y-%m-%d"))
logo_var     = tk.StringVar()

default_logo_path = os.path.join(SCRIPT_DIR, "Logo.jpg")
if os.path.exists(default_logo_path):
    logo_var.set(default_logo_path)

# LEFT
left_frame = tk.Frame(root, bg="#2c2c2c")
left_frame.pack(side="left", padx=10, pady=10, fill="y")

def update_file_list(folder_path):
    file_list.delete(0, tk.END)
    try:
        for f in os.listdir(folder_path):
            if f.lower().endswith(".xlsx"):
                file_list.insert(tk.END, f)
    except Exception:
        pass

def select_folder():
    p = filedialog.askdirectory()
    if p:
        folder_var.set(p); update_file_list(p)

ttk.Label(left_frame, text="Wybierz folder:").grid(row=0, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=folder_var, width=50).grid(row=0, column=1)
ttk.Button(left_frame, text="Przeglądaj", command=select_folder).grid(row=0, column=2)

ttk.Label(left_frame, text="Nazwa klienta:").grid(row=1, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=customer_var).grid(row=1, column=1)
ttk.Label(left_frame, text="Numer oferty:").grid(row=2, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=offer_var).grid(row=2, column=1)
ttk.Label(left_frame, text="Data oferty:").grid(row=3, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=date_var).grid(row=3, column=1)
ttk.Label(left_frame, text="Okres ważności:").grid(row=4, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=validity_var).grid(row=4, column=1)

def upload_logo():
    p = filedialog.askopenfilename(filetypes=[("Pliki obrazów", "*.png;*.jpg;*.jpeg")])
    if p: logo_var.set(p)

ttk.Label(left_frame, text="Wczytaj logo:").grid(row=5, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=logo_var, width=50).grid(row=5, column=1)
ttk.Button(left_frame, text="Przeglądaj", command=upload_logo).grid(row=5, column=2)

ttk.Label(left_frame, text="Szczegóły kontaktowe:").grid(row=6, column=0, sticky="ne")
contact_text = tk.Text(left_frame, height=5, width=50, bg="#3c3c3c", fg="white", insertbackground="white")
contact_text.grid(row=6, column=1)
contact_text.insert(tk.INSERT,
    "LP KONSTAL Sp. z o.o.\n59-800 Lubań\nPisarzowice 203B\nNIP: 1010004508\n\n"
    "Laser Team\nMateusz Brzostek M. +48 537 883 393\n"
    "Artur Jednoróg M. +48 515 803 333\nE. laser@konstal.com"
)

ttk.Label(left_frame, text="Tekst poprzedzający:").grid(row=7, column=0, sticky="ne")
preceding_text_var = tk.Text(left_frame, height=5, width=50, bg="#3c3c3c", fg="white", insertbackground="white")
preceding_text_var.grid(row=7, column=1)
preceding_text_var.insert(tk.INSERT,
    "Szanowni Państwo,\n\n"
    "dziękujemy za przesłanie zapytania ofertowego dotyczącego usługi cięcia laserem blach. "
    "Z przyjemnością przedstawiamy przygotowaną dla Państwa ofertę..."
)

ttk.Label(left_frame, text="").grid(row=8, column=0, pady=10)
ttk.Label(left_frame, text="Tekst kończący:").grid(row=9, column=0, sticky="ne")
finishing_text_var = tk.Text(left_frame, height=10, width=50, bg="#3c3c3c", fg="white", insertbackground="white")
finishing_text_var.grid(row=9, column=1)
finishing_text_var.insert(tk.INSERT, "Wyłączenia odpowiedzialności ...")

ttk.Label(left_frame, text="Odczytane pliki:").grid(row=10, column=0, sticky="ne")
file_list = tk.Listbox(left_frame, height=5, width=50, bg="#3c3c3c", fg="white")
file_list.grid(row=10, column=1)

def open_selected_file(event=None):
    sel = file_list.curselection()
    if sel:
        f = os.path.join(folder_var.get(), file_list.get(sel[0]))
        try:
            os.startfile(f)
        except Exception:
            pass

file_list.bind('<Double-Button-1>', open_selected_file)

buttons_frame = tk.Frame(left_frame, bg="#2c2c2c")
buttons_frame.grid(row=11, column=1, sticky="s")

# RIGHT
right_frame = tk.Frame(root, bg="#2c2c2c")
right_frame.pack(side="right", padx=10, pady=10, fill="both", expand=True)

right_paned = tk.PanedWindow(right_frame, orient=tk.VERTICAL, bg="#2c2c2c", sashrelief="raised", borderwidth=1)
right_paned.pack(fill="both", expand=True)

panel_a = tk.PanedWindow(right_paned, orient=tk.VERTICAL, bg="#2c2c2c", sashrelief="raised", borderwidth=1)

# --- PANEL 1 ---
subpanel1 = tk.LabelFrame(panel_a, text="PANEL 1 — PODGLĄD", bg="#2c2c2c", fg="white")
columns = ("1", "2", "3", "4", "5", "6", "7", "8", "9")
tree = ttk.Treeview(subpanel1, columns=columns, show="tree headings")
tree.column("#0", width=150, minwidth=100, stretch=tk.NO)  # Increased width for better thumbnail visibility
tree.heading("1", text="Nr");     tree.column("1", minwidth=50,  width=50,  stretch=tk.NO)
tree.heading("2", text="SubNr");  tree.column("2", minwidth=50,  width=50,  stretch=tk.NO)
tree.heading("3", text="Nazwa");  tree.column("3", minwidth=150, width=400, stretch=tk.NO)
tree.heading("4", text="Materiał"); tree.column("4", minwidth=50, width=80, stretch=tk.NO)
tree.heading("5", text="Grubość");  tree.column("5", minwidth=50, width=80, stretch=tk.NO, anchor="e")
tree.heading("6", text="Ilość");    tree.column("6", minwidth=50, width=80, stretch=tk.NO, anchor="e")
tree.heading("7", text="Laser");    tree.column("7", minwidth=50, width=100, stretch=tk.NO, anchor="e")
tree.heading("8", text="Gięcie/szt."); tree.column("8", minwidth=50, width=100, stretch=tk.NO, anchor="e")
tree.heading("9", text="Dodatkowe/szt."); tree.column("9", minwidth=50, width=120, stretch=tk.NO, anchor="e")

# Add scrollbar for treeview
scrollbar = ttk.Scrollbar(subpanel1, orient="vertical", command=tree.yview)
tree.configure(yscrollcommand=scrollbar.set)
tree.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")

def edit_cell(event):
    item = tree.identify_row(event.y)
    column = tree.identify_column(event.x)
    if not item or not column:
        return
    col_index = int(column[1:]) - 1
    if col_index in [5, 6, 7, 8]:
        x, y, w, h = tree.bbox(item, column)
        e = tk.Entry(subpanel1, bg="#3c3c3c", fg="white", insertbackground="white")
        e.place(x=x, y=y, width=w, height=h)
        e.insert(0, tree.item(item, 'values')[col_index])
        e.focus()
        def save_edit(_):
            vals = list(tree.item(item, 'values'))
            vals[col_index] = e.get()
            tree.item(item, values=vals)
            e.destroy()
        e.bind("<Return>", save_edit); e.bind("<FocusOut>", save_edit)

tree.bind("<Double-1>", edit_cell)
panel_a.add(subpanel1, minsize=220)

# --- PANEL 2 ---
subpanel2 = tk.LabelFrame(panel_a, text="PANEL 2 — STAŁE KOSZTY", bg="#2c2c2c", fg="white")
ttk.Label(subpanel2, text="Koszty operacyjne za arkusz:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
op_cost_entry = ttk.Entry(subpanel2); op_cost_entry.grid(row=0, column=1, padx=5, pady=5); op_cost_entry.insert(tk.INSERT, "40,00")
ttk.Label(subpanel2, text="Technologia za zlecenie:").grid(row=1, column=0, sticky="w", padx=5, pady=5)
tech_order_entry = ttk.Entry(subpanel2); tech_order_entry.grid(row=1, column=1, padx=5, pady=5); tech_order_entry.insert(tk.INSERT, "50,00")
ttk.Label(subpanel2, text="Dodatkowe koszty dla zlecenia (np. narzędzia):").grid(row=2, column=0, sticky="w", padx=5, pady=5)
add_order_cost_entry = ttk.Entry(subpanel2); add_order_cost_entry.grid(row=2, column=1, padx=5, pady=5); add_order_cost_entry.insert(tk.INSERT, "0,00")
subpanel2.update_idletasks()
panel2_height = subpanel2.winfo_reqheight() + 20  # Dodatkowe miejsce na marginesy
panel_a.add(subpanel2, height=panel2_height, minsize=panel2_height)

# --- PANEL 3 ---
subpanel3 = tk.LabelFrame(panel_a, text="PANEL 3 — CENNIKI I TESTY", bg="#2c2c2c", fg="white", padx=6, pady=6)

mat_frame = tk.LabelFrame(subpanel3, text="Cennik materiałów (PLN/kg)", bg="#2c2c2c", fg="white")
mat_frame.grid(row=0, column=0, sticky="nwe", padx=4, pady=4)
def _update_led(canvas, ok): canvas.delete("all"); canvas.create_oval(2,2,18,18, fill=("green" if ok else "red"))
btn_load_mat = ttk.Button(mat_frame, text="Załaduj cennik materiałów", command=lambda: load_material_prices(preview=True))
btn_load_mat.grid(row=0, column=0, sticky="w")
material_led = tk.Canvas(mat_frame, width=20, height=20, bg="#2c2c2c", highlightthickness=0); material_led.grid(row=0, column=1, padx=(6,0)); _update_led(material_led, False)
ttk.Label(mat_frame, text="Materiał:").grid(row=1, column=0, sticky="e", pady=(6,0))
material_var = tk.StringVar(); material_cb = ttk.Combobox(mat_frame, textvariable=material_var, width=22, state="readonly"); material_cb.grid(row=1, column=1, sticky="w", pady=(6,0))
ttk.Label(mat_frame, text="Grubość [mm]:").grid(row=2, column=0, sticky="e")
thickness_mat_var = tk.StringVar(); thickness_mat_cb = ttk.Combobox(mat_frame, textvariable=thickness_mat_var, width=12, state="readonly"); thickness_mat_cb.grid(row=2, column=1, sticky="w")
btn_find_mat = ttk.Button(mat_frame, text="Znajdź cenę materiału", command=lambda: ui_find_material_price()); btn_find_mat.grid(row=3, column=0, columnspan=2, pady=4, sticky="we")
material_result_label = ttk.Label(mat_frame, text="Wynik Materiał: —"); material_result_label.grid(row=4, column=0, columnspan=2, sticky="w")

cut_frame = tk.LabelFrame(subpanel3, text="Cennik cięcia (PLN/m)", bg="#2c2c2c", fg="white")
cut_frame.grid(row=0, column=1, sticky="nwe", padx=4, pady=4)
btn_load_cut = ttk.Button(cut_frame, text="Załaduj cennik cięcia", command=lambda: load_cutting_prices(preview=True))
btn_load_cut.grid(row=0, column=0, sticky="w")
cutting_led = tk.Canvas(cut_frame, width=20, height=20, bg="#2c2c2c", highlightthickness=0); cutting_led.grid(row=0, column=1, padx=(6,0)); _update_led(cutting_led, False)
ttk.Label(cut_frame, text="Materiał:").grid(row=1, column=0, sticky="e", pady=(6,0))
material_cut_var = tk.StringVar(); material_cut_cb = ttk.Combobox(cut_frame, textvariable=material_cut_var, width=22, state="readonly"); material_cut_cb.grid(row=1, column=1, sticky="w", pady=(6,0))
ttk.Label(cut_frame, text="Grubość [mm]:").grid(row=2, column=0, sticky="e")
thickness_cut_var = tk.StringVar(); thickness_cut_cb = ttk.Combobox(cut_frame, textvariable=thickness_cut_var, width=12, state="readonly"); thickness_cut_cb.grid(row=2, column=1, sticky="w")
ttk.Label(cut_frame, text="Gaz:").grid(row=3, column=0, sticky="e")
gas_var = tk.StringVar(); gas_cb = ttk.Combobox(cut_frame, textvariable=gas_var, width=12, state="readonly"); gas_cb.grid(row=3, column=1, sticky="w")
btn_find_cut = ttk.Button(cut_frame, text="Znajdź cenę cięcia", command=lambda: ui_find_cutting_price()); btn_find_cut.grid(row=4, column=0, columnspan=2, pady=4, sticky="we")
cutting_result_label = ttk.Label(cut_frame, text="Wynik Cięcie: —"); cutting_result_label.grid(row=5, column=0, columnspan=2, sticky="w")

btn_load_both = ttk.Button(subpanel3, text="Załaduj oba cenniki i odśwież listy",
                          command=lambda: (load_material_prices(True), load_cutting_prices(True)))
btn_load_both.grid(row=1, column=0, columnspan=2, sticky="we", padx=4, pady=(2,6))
subpanel3.grid_columnconfigure(0, weight=1); subpanel3.grid_columnconfigure(1, weight=1)

panel_a.add(subpanel3, minsize=200)
right_paned.add(panel_a)

# ---- Loader'y cenników ----
def _tree_preview_clear_and_headers(headers):
    for item in tree.get_children():
        tree.delete(item)
    tree.insert('', 'end', values=(0, '', ' | '.join(headers), '', '', '', '', '', ''))

def load_material_prices(preview=False):
    global material_prices, _mat_set, _thk_set
    material_prices.clear(); _mat_set.clear(); _thk_set.clear()
    try:
        if not os.path.exists(MATERIALS_FILE):
            raise FileNotFoundError(f"Nie znaleziono pliku: {MATERIALS_FILE}")
        wb = load_workbook(MATERIALS_FILE, data_only=True)
        sheet = wb.active
        headers = [str(c.value).strip().lower() if c.value is not None else "" for c in next(sheet.iter_rows(min_row=1, max_row=1))]
        need = ("material", "thickness", "price")
        idx = {n: headers.index(n) for n in need if n in headers}
        if not set(need).issubset(idx):
            raise ValueError("Brak wymaganych kolumn: material, thickness, price")
        if preview: _tree_preview_clear_and_headers(["materials prices.xlsx → material/thickness/price"])
        for row in sheet.iter_rows(min_row=2, values_only=True):
            mat = _norm_s(row[idx["material"]]); thk = _parse_float(row[idx["thickness"]]); prc = _parse_float(row[idx["price"]])
            if mat and thk is not None and prc is not None:
                material_prices[(mat, thk)] = prc; _mat_set.add(mat); _thk_set.add(thk)
                if preview: tree.insert('', 'end', values=("", "", f"{mat} @ {thk:.2f} mm → {format_pln(prc)} PLN/kg", "", "", "", "", "", ""))
        mats_sorted = sorted(_mat_set); thk_sorted = [f"{t:.2f}".rstrip("0").rstrip(".") for t in sorted(_thk_set)]
        material_cb["values"] = mats_sorted; material_cut_cb["values"] = mats_sorted
        thickness_mat_cb["values"] = thk_sorted
        if not thickness_cut_cb["values"]: thickness_cut_cb["values"] = thk_sorted
        _update_led(material_led, len(material_prices) > 0)
    except Exception as e:
        _update_led(material_led, False); messagebox.showerror("Błąd", f"Ładowanie cen materiałów:\n{e}")

def load_cutting_prices(preview=False):
    global cutting_prices, _mat_set, _thk_set, _gas_set
    cutting_prices.clear(); _gas_set.clear()
    try:
        if not os.path.exists(CUTTING_FILE):
            raise FileNotFoundError(f"Nie znaleziono pliku: {CUTTING_FILE}")
        wb = load_workbook(CUTTING_FILE, data_only=True)
        sheet = wb.active
        headers = [str(c.value).strip().lower() if c.value is not None else "" for c in next(sheet.iter_rows(min_row=1, max_row=1))]
        need = ("thickness", "material", "gas", "price")
        idx = {n: headers.index(n) for n in need if n in headers}
        if not set(need).issubset(idx):
            raise ValueError("Brak wymaganych kolumn: thickness, material, gas, price")
        if preview: _tree_preview_clear_and_headers(["cutting prices.xlsx → thickness/material/gas/price"])
        for row in sheet.iter_rows(min_row=2, values_only=True):
            thk = _parse_float(row[idx["thickness"]]); mat = _norm_s(row[idx["material"]]); gas = _norm_s(row[idx["gas"]]); prc = _parse_float(row[idx["price"]])
            if thk is not None and mat and gas and prc is not None:
                cutting_prices[(thk, mat, gas)] = prc; _mat_set.add(mat); _thk_set.add(thk); _gas_set.add(gas)
                if preview: tree.insert('', 'end', values=("", "", f"{thk:.2f} mm / {mat} / {gas} → {format_pln(prc)} PLN/m", "", "", "", "", "", ""))
        mats_sorted = sorted(_mat_set); thk_sorted = [f"{t:.2f}".rstrip("0").rstrip(".") for t in sorted(_thk_set)]; gas_sorted = sorted(_gas_set)
        material_cut_cb["values"] = mats_sorted
        if not material_cb["values"]: material_cb["values"] = mats_sorted
        thickness_cut_cb["values"] = thk_sorted
        if not thickness_mat_cb["values"]: thickness_mat_cb["values"] = thk_sorted
        gas_cb["values"] = gas_sorted
        _update_led(cutting_led, len(cutting_prices) > 0)
    except Exception as e:
        _update_led(cutting_led, False); messagebox.showerror("Błąd", f"Ładowanie cen cięcia:\n{e}")

# ---- UI testy (Panel 3) ----
def ui_find_material_price():
    mat = _norm_s(material_var.get()); thk = _parse_float(thickness_mat_var.get())
    if not mat or thk is None:
        messagebox.showerror("Błąd", "Uzupełnij Materiał i Grubość (mm)."); return
    price = material_prices.get((mat, thk))
    material_result_label.config(text="Wynik Materiał: nie znaleziono" if price is None else f"Wynik Materiał: {format_pln(price)} PLN/kg")

def ui_find_cutting_price():
    mat = _norm_s(material_cut_var.get()); thk = _parse_float(thickness_cut_var.get()); gas = _norm_s(gas_var.get())
    if not mat or thk is None or not gas:
        messagebox.showerror("Błąd", "Uzupełnij Materiał, Grubość (mm) i Gaz."); return
    price = cutting_prices.get((thk, mat, gas))
    cutting_result_label.config(text="Wynik Cięcie: nie znaleziono" if price is None else f"Wynik Cięcie: {format_pln(price)} PLN/m")

# ---- Analiza folderu ----
last_groups = []; last_total_cost = 0.0; last_folder_path = ""

def _ensure_cenniki_loaded():
    ok = True
    if not material_prices:
        try: load_material_prices(preview=False)
        except Exception: ok = False
    if not cutting_prices:
        try: load_cutting_prices(preview=False)
        except Exception: ok = False
    return ok


def get_total_cut_length(ws, text="Total") -> float:
    """
    Szuka w kolumnie A pierwszej komórki zawierającej 'text' (case-insensitive),
    po czym odczytuje wartość z kolumny H (8) w tym samym wierszu.
    Zwraca float; radzi sobie z polskim formatem '312,51'.
    """
    # Przechodzimy jedynie po kolumnie A — to szybkie i proste
    for cell in ws['A']:
        val = cell.value
        if val and str(text).lower() in str(val).lower():
            raw = ws.cell(row=cell.row, column=8).value  # kol. H
            # openpyxl zwykle zwraca już float dla liczbowych komórek;
            # jeśli to string w formacie '312,51', zamień przecinek:
            if isinstance(raw, (int, float)):
                return float(raw)
            try:
                return float(str(raw).replace(" ", "").replace("\xa0", "").replace(",", "."))
            except Exception:
                return 0.0
    raise ValueError(f"Nie znaleziono wiersza z tekstem '{text}' w kolumnie A")



def parse_duration_to_hours(value) -> float:
    """
    Zamienia '1h26min21s', '1h26m21s', '86min', '90s', '1:26:21', '1:26' itp. na float godzin.
    Zwraca 0.0, gdy nie da się sparsować.
    """
    if value is None:
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)

    s = str(value).strip().lower().replace(" ", "")

    # H:M:S lub H:M
    if ":" in s:
        parts = s.split(":")
        try:
            def to_f(x): return float(x.replace(",", "."))
            if len(parts) == 3:
                h, m, sec = parts
            elif len(parts) == 2:
                h, m = parts; sec = "0"
            else:
                h, m, sec = parts[0], "0", "0"
            return to_f(h) + to_f(m)/60.0 + to_f(sec)/3600.0
        except Exception:
            pass

    def take(pattern):
        m = re.search(pattern, s)
        return float(m.group(1).replace(",", ".")) if m else 0.0

    h = take(r"(\d+(?:[.,]\d+)?)h")
    m = take(r"(\d+(?:[.,]\d+)?)(?:m|min)")
    sec = take(r"(\d+(?:[.,]\d+)?)s")

    if h == 0.0 and m == 0.0 and sec == 0.0:
        # Samo '123' → potraktuj jako sekundy
        try:
            seconds = float(s.replace(",", "."))
            return seconds / 3600.0
        except Exception:
            return 0.0

    return h + m/60.0 + sec/3600.0



# ---- Dynamic pricing based on thickness, length and time ----
_THK_SPEED_MPM = [
    (1.0, 18.0),
    (1.5, 18.0),
    (2.0, 14.0),
    (3.0, 4.0),
    (5.0, 3.5),
    (6.0, 3.0),
    (8.0, 2.7),
    (10.0, 2.1),
    (12.0, 2.1),
    (15.0, 2.1),
]

def _interp(x, pts):
    """Linear interpolation over sorted (x, y) points. Clamps outside the range."""
    pts = sorted(pts, key=lambda p: p[0])
    if x <= pts[0][0]:
        return pts[0][1]
    if x >= pts[-1][0]:
        return pts[-1][1]
    for (x0, y0), (x1, y1) in zip(pts[:-1], pts[1:]):
        if x0 <= x <= x1:
            if x1 == x0:
                return y0
            t = (x - x0) / (x1 - x0)
            return y0 + t * (y1 - y0)
    return pts[-1][1]

def get_speed_mpm(thickness_mm: float) -> float:
    """Return cutting speed [m/min] for a given thickness using piecewise-linear interpolation."""
    if thickness_mm is None:
        return 0.0
    return float(_interp(float(thickness_mm), _THK_SPEED_MPM))

def get_time_thresholds_minutes(thickness_mm: float):
    """
    Zwraca (t_min, t_neutral) w minutach dla danej grubości.
    Kotwice: 1 mm → (1, 45), 15 mm → (5, 90). Liniowo i z klamrowaniem.
    """
    thk = max(1.0, min(float(thickness_mm or 1.0), 15.0))
    t_min = 1.0 + 4.0 * (thk - 1.0) / 14.0       # 1 → 5
    t_neutral = 45.0 + 45.0 * (thk - 1.0) / 14.0 # 45 → 90
    return (t_min, t_neutral)

def compute_effective_minutes(thickness_mm: float, total_cut_length_m: float, cut_time_hours: float) -> float:
    """
    Efektywny czas zlecenia [min] na podstawie długości ścieżki i podanego cut_time.
    Liczymy czas z długości: L / v(thk), bierzemy maksimum z tym z komórki.
    """
    speed = get_speed_mpm(thickness_mm)  # m/min
    t_from_len = 0.0
    if total_cut_length_m is not None and speed > 0:
        try:
            t_from_len = float(total_cut_length_m) / float(speed)
        except Exception:
            t_from_len = 0.0
    t_from_cell = float(cut_time_hours or 0.0) * 60.0
    if t_from_len <= 0:
        return t_from_cell
    if t_from_cell <= 0:
        return t_from_len
    return max(t_from_len, t_from_cell)

def compute_boost_factor(thickness_mm: float, effective_minutes: float, max_boost: float = 3.5) -> float:
    """
    Mnożnik w przedziale [1.0, max_boost].
    = max_boost, gdy czas ≤ t_min; = 1.0, gdy czas ≥ t_neutral; liniowo pomiędzy.
    """
    t_min, t_neutral = get_time_thresholds_minutes(thickness_mm)
    if effective_minutes <= t_min:
        return float(max_boost)
    if effective_minutes >= t_neutral:
        return 1.0
    alpha = (t_neutral - effective_minutes) / (t_neutral - t_min)  # (0,1)
    return 1.0 + (max_boost - 1.0) * alpha

def apply_dynamic_pricing(price_per_kg: float, rate_per_m: float, thickness_mm: float, total_cut_length_m: float, cut_time_hours: float):
    """
    Zasady:
    - materiał: MINIMUM +7% powyżej ceny zakupu,
    - materiał i cięcie: dodatkowy mnożnik 'boost' zależny od czasu (1.0 → 3.5 = +250%) wg progów,
    - materiał: max(1.07, boost) × cena zakupu,
      cięcie:   boost × stawka za metr.
    Zwraca (material_price_adj, cutting_rate_adj, debug_dict).
    """
    try:
        base_material = float(price_per_kg or 0.0)
        base_rate = float(rate_per_m or 0.0)
    except Exception:
        base_material, base_rate = 0.0, 0.0

    eff_minutes = compute_effective_minutes(thickness_mm, total_cut_length_m, cut_time_hours)
    boost = compute_boost_factor(thickness_mm, eff_minutes, max_boost=3.5)

    material_factor = max(1.07, boost)  # minimalnie +7%
    cutting_factor = boost

    mat_adj = base_material * material_factor
    rate_adj = base_rate * cutting_factor

    dbg = {
        'eff_minutes': eff_minutes,
        't_min_t_neutral': get_time_thresholds_minutes(thickness_mm),
        'boost': boost,
        'material_factor': material_factor,
        'cutting_factor': cutting_factor,
    }
    return mat_adj, rate_adj, dbg

# Referencje do PhotoImage, żeby obrazy nie znikały (GC)
tree_img_refs = []
thumbnail_imgs = []

def analyze_xlsx_folder():
    for item in tree.get_children():
        tree.delete(item)
    thumbnail_imgs.clear()
    global last_groups, last_total_cost, last_folder_path
    folder_path = folder_var.get()
    if not folder_path:
        messagebox.showerror("Błąd", "Proszę wybrać folder."); return
    try:
        files = [f for f in os.listdir(folder_path) if f.lower().endswith(".xlsx")]
    except Exception:
        files = []
    if not files:
        messagebox.showerror("Błąd", "Brak plików .xlsx w wybranym folderze."); return
    if not _ensure_cenniki_loaded():
        messagebox.showwarning("Uwaga", "Cenniki niezaładowane — użyję 0.00, sprawdź Panel 3.")

    # wyczyść tabelę
    for item in tree.get_children(): tree.delete(item)

    # stałe koszty
    def to_float(s):
        try: return float(str(s).replace(",", "."))
        except Exception: return 0.0
    op_cost_per_sheet = to_float(op_cost_entry.get() or "0")
    tech_per_order     = to_float(tech_order_entry.get() or "0")
    add_costs_order    = to_float(add_order_cost_entry.get() or "0")
    total_extra = tech_per_order + add_costs_order

    total_sheets = 0; total_parts_qty = 0
    groups = []; all_parts = []; subnr = 0 

    for fname in files:
        path = os.path.join(folder_path, fname)
        try:
            wb = load_workbook(path, data_only=True)
            if "All Task List" not in wb.sheetnames: raise KeyError("Brak arkusza 'All Task List'")
            # pobieranie obrazu z arkusza 'All Parts List'            
            all_task = wb["All Task List"]
            all_part_list = wb["All Parts List"]
            thumbnails = {}
            for img in all_part_list._images:
                row = img.anchor._from.row + 1
                col = img.anchor._from.col + 1
                if col == 2:  # Column B
                    img_data = img._data()
                    thumbnails[row] = img_data
            # Czas cięcia (All Task List!F4) → na '0,00h'
            cut_time = parse_duration_to_hours(all_task['F4'].value)
            total_cut_length = get_total_cut_length(all_task,"Total")
            material_name = all_task["B4"].value; thickness_raw = all_task["C4"].value; gas_raw = all_task["E4"].value
            mat_norm = _norm_s(material_name); thk_val = _parse_float(thickness_raw); gas_key = _map_gas_to_key(gas_raw)
            if not mat_norm: raise ValueError("All Task List!B4 (Material) — brak wartości")
            if thk_val is None: raise ValueError("All Task List!C4 (Thickness(mm)) — brak liczby")
            if not gas_key: raise ValueError("All Task List!E4 (Gas) — nieobsługiwany typ gazu")

            price_per_kg = material_prices.get((mat_norm, thk_val), 0.0)
            rate_per_cut_length = cutting_prices.get((thk_val, mat_norm, gas_key), 0.0)

            if "Cost List" not in wb.sheetnames: raise KeyError("Brak arkusza 'Cost List'")
            cost_sheet = wb["Cost List"]

            # Average utilization → szukamy tekstu, wartość w kol. K
            util_row = None
            for r in range(1, cost_sheet.max_row + 1):
                for c in range(1, cost_sheet.max_column + 1):
                    v = cost_sheet.cell(row=r, column=c).value
                    if v and "Average utilization:" in str(v):
                        util_row = r; break
                if util_row: break
            if util_row is None: raise ValueError("Nie znaleziono 'Average utilization:'")
            util_str = cost_sheet.cell(row=util_row, column=11).value
            util_val = _parse_float(str(util_str).replace("%","")) if util_str is not None else None
            utilization_rate = (util_val/100.0) if (util_val is not None) else 0.0
            if utilization_rate <= 0 or utilization_rate > 1:
                messagebox.showwarning("Uwaga", f"Average utilization poza zakresem ({utilization_rate}).")

            # 'Material Price' w kol. A -> stawki kontur/marking/defilm
            mat_price_row = None
            for r in range(1, cost_sheet.max_row + 1):
                v = cost_sheet.cell(row=r, column=1).value
                if v and "Material Price" in str(v): mat_price_row = r; break
            if mat_price_row is None: raise ValueError("Brak wiersza 'Material Price'")

            def parse_num(cellv):
                if cellv is None: return 0.0
                s = str(cellv).strip(); s = ''.join(ch for ch in s.split()[0] if ch.isdigit() or ch in ('.',',')); s = s.replace(",", ".")
                try: return float(s or "0.0")
                except Exception: return 0.0

            rate_per_contour        = parse_num(cost_sheet.cell(row=mat_price_row, column=7).value)   # G
            rate_per_marking_length = parse_num(cost_sheet.cell(row=mat_price_row, column=9).value)   # I
            rate_per_defilm_length  = parse_num(cost_sheet.cell(row=mat_price_row, column=10).value)  # J


            # Dynamic pricing: +7% min. na materiale oraz boost zależny od czasu/ długości
            #• Dla bardzo krótkich zleceń (np. ≤ 1 min @ 1 mm lub ≤ 5 min @ 15 mm) — materiał i stawka cięcia ×3.5.
            #• Dla długich zleceń (≥ 45 min @ 1 mm, ≥ 90 min @ 15 mm) — cięcie bez zmian, materiał min. ×1.07.
            #• Pośrednio — liniowa interpolacja między tymi wartościami.
            #• Prędkości cięcia (m/min) interpolowane po podanych punktach (1, 1.5, 2, 3, 5, 6, 8, 10, 12, 15 mm).
            
            price_per_kg, rate_per_cut_length, _dpdbg = apply_dynamic_pricing(
                price_per_kg, rate_per_cut_length, thk_val, total_cut_length, cut_time
            )



            # Suma arkuszy: "Cut number:" w kol. D (od wiersza 8) w All Task List
            r_idx = 8
            while all_task.cell(row=r_idx, column=4).value is not None:
                v = all_task.cell(row=r_idx, column=4).value
                if isinstance(v, (int, float)):
                    total_sheets += int(v)
                r_idx += 1

            # Dane części — pierwszy numeryczny wiersz kol. A
            start_row = None
            for r in range(1, cost_sheet.max_row + 1):
                a_val = cost_sheet.cell(row=r, column=1).value
                if a_val and isinstance(a_val, (int, float)):
                    start_row = r; break
            if start_row is None: raise ValueError("Nie znaleziono wiersza startowego (kol. A — ID)")

            parts_for_group = []; subnr += 1; lp = 0; row = start_row
            while row <= cost_sheet.max_row and isinstance(cost_sheet.cell(row=row, column=1).value, (int,float)):
                lp += 1
                part_name      = cost_sheet.cell(row=row, column=2).value
                part_qty       = cost_sheet.cell(row=row, column=5).value or 0
                weight         = parse_num(cost_sheet.cell(row=row, column=6).value)
                contours_qty   = parse_num(cost_sheet.cell(row=row, column=7).value)
                cut_length     = parse_num(cost_sheet.cell(row=row, column=8).value)
                marking_length = parse_num(cost_sheet.cell(row=row, column=9).value)
                defilm_length  = parse_num(cost_sheet.cell(row=row, column=10).value)

                adj_weight   = (weight / utilization_rate) if utilization_rate > 0 else weight
                material_cost= adj_weight * price_per_kg
                contours_cost= contours_qty * rate_per_contour
                cut_cost     = cut_length * rate_per_cut_length
                marking_cost = marking_length * rate_per_marking_length
                defilm_cost  = defilm_length * rate_per_defilm_length
                total_part   = material_cost + contours_cost + cut_cost + marking_cost + defilm_cost

                # --- Pobierz thumbnail z arkusza All Parts List ---
                thumbnail_photo = None
                all_parts_row = 2 + lp  # Assuming row 3 is first part
                if all_parts_row in thumbnails:
                    data = thumbnails[all_parts_row]
                    try:
                        pil_img = Image.open(io.BytesIO(data))
                        max_w, max_h = 140, 70
                        w, h = pil_img.size
                        ratio = min(max_w / w, max_h / h, 1.0)
                        new_w = int(w * ratio)
                        new_h = int(h * ratio)
                        pil_img = pil_img.resize((new_w, new_h), Image.LANCZOS)
                        thumbnail_photo = ImageTk.PhotoImage(pil_img)
                        thumbnail_imgs.append(thumbnail_photo)
                    except Exception:
                        pass

                all_parts.append({
                'id': lp,
                'subnr': subnr,
                'name': part_name,
                'material': material_name,
                'thickness': thk_val,
                'qty': int(part_qty) if isinstance(part_qty,(int,float)) else 0,
                'cost_per_unit': float(f"{total_part:.2f}"),
                'thumb': thumbnail_photo, # <-- nowość
                })

                parts_for_group.append((part_name, float(f"{total_part:.2f}"),
                                        int(part_qty) if isinstance(part_qty,(int,float)) else 0))
                total_parts_qty += int(part_qty) if isinstance(part_qty,(int,float)) else 0
                row += 1

            groups.append((material_name, thk_val, parts_for_group))

        except Exception as e:
            messagebox.showerror("Błąd", f"Błąd podczas przetwarzania pliku {fname}: {e}")
            return

    # rozdział overheadów na sztuki
    op_cost_total = total_sheets * op_cost_per_sheet
    extra_total   = total_extra
    if total_parts_qty > 0:
        extra_per_part  = extra_total / total_parts_qty
        op_cost_per_part= op_cost_total / total_parts_qty
    else:
        extra_per_part = op_cost_per_part = 0.0

    for p in all_parts:
        p['cost_per_unit'] = float(f"{(p['cost_per_unit'] + extra_per_part + op_cost_per_part):.2f}")

    # tabela
    for i, p in enumerate(all_parts, start=1):
        item_values = (
            i,
            p['subnr'],
            p['name'],
            p['material'],
            f"{p['thickness']}",
            p['qty'],
            format_pln(p['cost_per_unit']),
            "",  # Gięcie/szt.
            "",  # Dodatkowe/szt.
        )
        opts = {'values': item_values}
        thumb = p.get('thumb')
        if thumb:  # only include when a PhotoImage is present
            opts['image'] = thumb

        tree.insert('', 'end', **opts)

    total_sum = 0.0; merged_groups = []
    for (mat_name, thk, parts) in groups:
        adj = []
        for (nm, cost, qty) in parts:
            c = float(f"{(cost + extra_per_part + op_cost_per_part):.2f}")
            adj.append((nm, c, qty)); total_sum += c * qty
        merged_groups.append((mat_name, thk, adj))

    last_groups[:] = merged_groups
    last_total_cost = total_sum
    last_folder_path = folder_path
    messagebox.showinfo("Analiza", "Analiza plików XLSX zakończona. Dane w Panelu 1 uzupełnione.")

# raport
def generate_report():
    if not last_groups:
        messagebox.showerror("Błąd", "Brak danych do raportu. Najpierw 'Analizuj XLSX'."); return
    folder_path = (folder_var.get().strip() or last_folder_path)
    if not folder_path or not os.path.isdir(folder_path):
        messagebox.showerror("Błąd", "Nieprawidłowy folder docelowy."); return
    customer_name = customer_var.get().strip() or "Klient"
    offer_number  = offer_var.get().strip() or "0001"
    offer_date    = date_var.get().strip() or datetime.datetime.now().strftime("%Y-%m-%d")
    validity      = validity_var.get().strip() or (datetime.datetime.now() + timedelta(days=14)).strftime("%Y-%m-%d")
    logo_path     = logo_var.get().strip()
    contact_details = contact_text.get("1.0", tk.END).strip()
    preceding_text  = preceding_text_var.get("1.0", tk.END).strip()
    finishing_text  = finishing_text_var.get("1.0", tk.END).strip()

    doc = Document()
    if logo_path and os.path.exists(logo_path):
        try: doc.add_picture(logo_path, width=Inches(3.0))
        except Exception: pass
    if contact_details:
        p = doc.add_paragraph(contact_details)
        for r in p.runs: r.bold = False

    doc.add_heading(f"Oferta dla {customer_name}", level=1)
    p = doc.add_paragraph(f"Numer oferty: {offer_number}"); p.runs[0].bold = True
    doc.add_paragraph(f"Data oferty: {offer_date}")
    doc.add_paragraph(f"Okres ważności: {validity}")
    if preceding_text: doc.add_paragraph(preceding_text)

    table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'
    hdr = table.rows[0].cells; hdr[0].text='Lp.'; hdr[1].text='Nazwa części'; hdr[2].text='Ilość'; hdr[3].text='Koszt (PLN)'; hdr[4].text='Razem (PLN)'
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '006995'); tcPr.append(shd)
        run = cell.paragraphs[0].runs[0]; run.font.color.rgb = RGBColor(255,255,255); run.bold = True; run.font.size = Pt(11)

    total = 0.0; lp = 1
    for mat_name, thk, parts in last_groups:
        row = table.add_row().cells; row[0].text=""; row[1].text=f"Materiał: {mat_name}, Grubość: {thk} mm"; row[1].merge(row[4])
        run = row[1].paragraphs[0].runs[0]; run.font.size = Pt(9); run.italic = True
        for nm, cost_per_unit, qty in parts:
            r = table.add_row().cells
            r[0].text = str(lp); r[1].text = str(nm) if nm else "Brak nazwy"
            r[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; r[2].paragraphs[0].add_run(f"{int(qty)}  ").font.size = Pt(10)
            r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; r[3].paragraphs[0].add_run(f"{format_pln(cost_per_unit)}  ").font.size = Pt(10)
            row_total = cost_per_unit * qty
            r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; r[4].paragraphs[0].add_run(f"{format_pln(row_total)}  ").font.size = Pt(10)
            total += row_total; lp += 1

    srow = table.add_row().cells; srow[1].text="Razem"; srow[4].text=format_pln(total); srow[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in srow[1].paragraphs[0].runs: r.bold=True
    for r in srow[4].paragraphs[0].runs: r.bold=True

    widths = [Cm(1), Cm(8), Cm(2), Cm(3), Cm(3)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells: cell.width = w

    p = doc.add_paragraph(f"Całkowity koszt: {format_pln(total)} PLN"); p.paragraph_format.space_before = Pt(12)
    for r in p.runs: r.font.size = Pt(14)
    if finishing_text:
        pf = doc.add_paragraph(finishing_text)
        for r in pf.runs: r.font.size = Pt(9)

    current_date = datetime.datetime.now().strftime("%Y%m%d")
    fname = f"Oferta_{sanitize_filename(customer_name) or 'Klient'}_{current_date}_{offer_number}.docx"
    full = os.path.join(folder_path, fname)
    try:
        doc.save(full); messagebox.showinfo("Sukces", f"Raport wygenerowany:\n{fname}")
    except Exception as e:
        messagebox.showerror("Błąd", f"Nie udało się zapisać pliku DOCX:\n{e}")

# przyciski lewe
ttk.Button(buttons_frame, text="Analizuj XLSX", command=analyze_xlsx_folder).pack(side="left", padx=5)
ttk.Button(buttons_frame, text="Generuj raport", command=generate_report).pack(side="left")

# ---- ustawienie sashy ----
def set_sash_positions(attempt=1):
    try:
        root.update_idletasks()
        panes = panel_a.panes()
        sash_count = max(len(panes) - 1, 0)
        if sash_count == 0:
            root.after(60, lambda: set_sash_positions(attempt+1)); return

        h = panel_a.winfo_height()
        if h < 400 and attempt < 10:
            # okno jeszcze się rozciąga — spróbuj później
            root.after(80, lambda: set_sash_positions(attempt+1)); return

        # rozkład: Panel1 ~ 50% wysokości, Panel2 ~ 20%, Panel3 reszta (minsize chroni przed 0px)
        y1 = max(220, int(h * 0.50))
        try: panel_a.sash_place(0, 0, y1)
        except Exception: pass

        if sash_count >= 2:
            y2 = min(h - 200, y1 + panel2_height)
            try: panel_a.sash_place(1, 0, y2)
            except Exception: pass
    except tk.TclError:
        if attempt < 10:
            root.after(80, lambda: set_sash_positions(attempt+1))

root.after_idle(set_sash_positions)

# run
root.geometry("1280x800")
root.mainloop()
