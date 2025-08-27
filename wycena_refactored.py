#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
wycena.py - Main GUI script for analyzing XLSX files and generating cost reports.

Usage instructions:
1. Run the script in a Python 3 environment with installed libraries: tkinter, openpyxl, docx, Pillow, requests.
2. Select the folder with XLSX files.
3. Analyze XLSX to fill the table.
4. Edit values in the table if needed (quantity, laser, bending, additional).
5. Click "Generate report" to create the DOCX offer, XLSX reports, and log.

The script is optimized for readability and performance, with full documentation.
"""

import os
import datetime
from datetime import timedelta
import re
from tkinter import ttk
import tkinter as tk
from tkinter import filedialog, messagebox
from openpyxl import load_workbook
from openpyxl import Workbook
from openpyxl.drawing.image import Image as OpenpyxlImage
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart.label import DataLabelList
from docx import Document
from docx.shared import Inches, RGBColor, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from PIL import Image, ImageTk
import locale
import requests

# total order value

def SetTotalPricePerOrder(value):
    global total_price_per_order
    total_price_per_order = value   

try:
    locale.setlocale(locale.LC_ALL, 'pl_PL.UTF-8')
except Exception:
    pass

def format_pln(value):
    """Formats the value to PLN with a comma and thousands grouping."""
    try:
        s = locale.format_string('%.2f', float(value), grouping=True)
        return s.replace('.', ',')
    except Exception:
        try:
            return f"{float(value):.2f}".replace('.', ',')
        except Exception:
            return "0,00"

def sanitize_filename(name):
    """Sanitizes the file name by replacing disallowed characters."""
    for ch in r'< > : " / \ | ? *':
        name = name.replace(ch, '_')
    return name

def _norm_s(s):
    """Normalizes the string to uppercase, removes spaces."""
    return (str(s).strip().upper() if s is not None else "")

def _parse_float(val):
    """Parses the value to float, handles commas."""
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip().replace(" ", "").replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None

def _map_gas_to_key(gas_raw: str) -> str:
    """Maps the gas name to the key 'N' or 'O'."""
    g = _norm_s(gas_raw)
    if g in {"NITROGEN", "AZOT", "氮气", "N"}:
        return "N"
    if g in {"OXYGEN", "TLEN", "氧气", "O"}:
        return "O"
    return ""

# ---- price lists ----
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MATERIALS_FILE = os.path.join(SCRIPT_DIR, "materials prices.xlsx")
CUTTING_FILE   = os.path.join(SCRIPT_DIR, "cutting prices.xlsx")
material_prices = {}  # (MAT, THK)-> PLN/kg
cutting_prices  = {}  # (THK, MAT, GAS)-> PLN/m
_mat_set, _thk_set, _gas_set = set(), set(), set()

# Global variables for shared data
all_parts = []
last_groups = []
last_total_cost = 0.0
last_folder_path = ""
total_sheets = 0
total_parts_qty = 0

# ---- GUI ----
root = tk.Tk()
root.title("Cost Report Generator — MERGED (fixed)")
root.configure(bg="#2c2c2c")  # Dark background for modern look

# Use a modern ttk theme
style = ttk.Style(root)
style.theme_use('clam')  # Or 'alt', 'default', etc. for better visuals
style.configure("TLabel", foreground="white", background="#2c2c2c", font=("Arial", 10))
style.configure("TEntry", fieldbackground="#3c3c3c", foreground="white")
style.configure("TButton", background="#4c4c4c", foreground="white", borderwidth=0)
style.configure("TCombobox", fieldbackground="#3c3c3c", foreground="white")
style.configure("Treeview", background="#3c3c3c", foreground="white", fieldbackground="#3c3c3c", rowheight=80)  # Increased row height for better thumbnails
style.configure("Treeview.Heading", background="#4c4c4c", foreground="white")
style.map("TButton", background=[('active', '#5c5c5c')])
style.map("Treeview", background=[('selected', '#5c5c5c')])

folder_var   = tk.StringVar()
customer_var = tk.StringVar()
offer_var    = tk.StringVar()
date_var     = tk.StringVar(value=datetime.datetime.now().strftime("%Y-%m-%d"))
validity_var = tk.StringVar(value=(datetime.datetime.now() + timedelta(days=14)).strftime("%Y-%m-%d"))
logo_var     = tk.StringVar()

default_logo_path = os.path.join(SCRIPT_DIR, "Logo.jpg")
if os.path.exists(default_logo_path):
    logo_var.set(default_logo_path)

# LEFT
left_frame = tk.Frame(root, bg="#2c2c2c")
left_frame.pack(side="left", padx=10, pady=10, fill="y")

def update_file_list(folder_path):
    file_list.delete(0, tk.END)
    try:
        for f in os.listdir(folder_path):
            if f.lower().endswith(".xlsx"):
                file_list.insert(tk.END, f)
    except Exception:
        pass

def select_folder():
    p = filedialog.askdirectory()
    if p:
        folder_var.set(p); update_file_list(p)

ttk.Label(left_frame, text="Select folder:").grid(row=0, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=folder_var, width=50).grid(row=0, column=1)
ttk.Button(left_frame, text="Browse", command=select_folder).grid(row=0, column=2)

ttk.Label(left_frame, text="Client name:").grid(row=1, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=customer_var).grid(row=1, column=1)
ttk.Label(left_frame, text="Offer number:").grid(row=2, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=offer_var).grid(row=2, column=1)
ttk.Button(left_frame, text="Get number", command=lambda: offer_var.set(get_next_offer_number())).grid(row=2, column=2)
ttk.Label(left_frame, text="Offer date:").grid(row=3, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=date_var).grid(row=3, column=1)
ttk.Label(left_frame, text="Validity period:").grid(row=4, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=validity_var).grid(row=4, column=1)

def upload_logo():
    p = filedialog.askopenfilename(filetypes=[("Image files", "*.png;*.jpg;*.jpeg")])
    if p: logo_var.set(p)

ttk.Label(left_frame, text="Load logo:").grid(row=5, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=logo_var, width=50).grid(row=5, column=1)
ttk.Button(left_frame, text="Browse", command=upload_logo).grid(row=5, column=2)

ttk.Label(left_frame, text="Contact details:").grid(row=6, column=0, sticky="ne")
contact_text = tk.Text(left_frame, height=5, width=50, bg="#3c3c3c", fg="white", insertbackground="white")
contact_text.grid(row=6, column=1)
contact_text.insert(tk.INSERT,
    "LP KONSTAL Sp. z o.o.\n59-800 Lubań\nPisarzowice 203B\nNIP: 1010004508\n\n"
    "Laser Team\nMateusz Brzostek M. +48 537 883 393\n"
    "Artur Jednoróg M. +48 515 803 333\nE. laser@konstal.com"
)

ttk.Label(left_frame, text="Preceding text:").grid(row=7, column=0, sticky="ne")
preceding_text_var = tk.Text(left_frame, height=5, width=50, bg="#3c3c3c", fg="white", insertbackground="white")
preceding_text_var.grid(row=7, column=1)
preceding_text_var.insert(tk.INSERT,
    "Szanowni Państwo,\n\n"
    "dziękujemy za przesłanie zapytania ofertowego dotyczącego usługi cięcia laserem blach. "
    "Z przyjemnością przedstawiamy przygotowaną dla Państwa ofertę..."
)

ttk.Label(left_frame, text="").grid(row=8, column=0, pady=10)
ttk.Label(left_frame, text="Finishing text:").grid(row=9, column=0, sticky="ne")
finishing_text_var = tk.Text(left_frame, height=10, width=50, bg="#3c3c3c", fg="white", insertbackground="white")
finishing_text_var.grid(row=9, column=1)
finishing_text_var.insert(tk.INSERT, "Wyłączenia odpowiedzialności \r\nDokumentacja techniczna\r\nRealizacja zamówienia odbywa się wyłącznie na podstawie dokumentacji technicznej dostarczonej przez Klienta. Odpowiedzialność za jej kompletność, poprawność oraz zgodność z założeniami projektowymi leży wyłącznie po stronie Zleceniodawcy. Wszelkie błędy, niejasności, czy niezgodności w przesłanych plikach uniemożliwiające prawidłowe wykonanie wyrobu, nie mogą stanowić podstawy do roszczeń wobec naszej firmy.\r\n\r\nMateriał powierzone i dostarczany przez Klienta\r\nNie ponosimy odpowiedzialności za uszkodzenia, błędy obróbki, zmiany struktury, odkształcenia ani inne wady powstałe w wyniku specyficznych właściwości materiału powierzonego przez Klienta, jego niejednorodności, błędnej deklaracji gatunku, braku wymaganych atestów czy oznaczeń partii. Klient zobowiązany jest dostarczyć materiał zgodny ze specyfikacją oraz wolny od wad fizycznych i chemicznych, mogących negatywnie wpływać na proces cięcia i jakość finalnego wyrobu.\r\n\r\nDostawcy materiałów\r\nNasza firma dołoży wszelkich starań w zakresie selekcji i zakupów materiałów wyłącznie od sprawdzonych dostawców. Zastrzegamy sobie jednak, że odpowiedzialność za parametry, właściwości lub wady ukryte materiału ogranicza się wyłącznie do zakresu wynikającego z dokumentacji danego producenta lub certyfikatu jakości — zgodnie z obowiązującym prawem oraz praktyką rynku stalowego.\r\n\r\nOgraniczenie odpowiedzialności prawnej\r\nOdpowiadamy wyłącznie za zgodność wykonanych prac z przesłaną dokumentacją oraz z obowiązującymi normami i przepisami prawa. Nie ponosimy odpowiedzialności za ewentualne szkody pośrednie, utracone korzyści, koszty produkcji, opóźnienia wynikające z przerw w dostawie materiałów, siły wyższej, zdarzeń losowych czy skutków niezastosowania się Klienta do obowiązujących przepisów i wymogów technicznych.\r\n\r\nPrzepisy prawa i gwarancje\r\nWszelkie realizacje podlegają przepisom prawa polskiego, normom branżowym oraz ustaleniom indywidualnym zawartym w zamówieniu. Ewentualna odpowiedzialność spółki ogranicza się do wartości usługi, a w szczególnych wypadkach — do ponownego wykonania usługi lub zwrotu jej kosztu. Nie udzielamy gwarancji na materiały powierzone, a zakres gwarancji na produkty wykonane z własnych materiałów jest określony indywidualnie w ofercie i na fakturze.\r\n\r\nMamy nadzieję, że powyższe wyjaśnienia pozwolą na jasne i czytelne określenie zasad współpracy oraz przyczynią się do pomyślnej realizacji Państwa zamówienia. Zapraszamy do zapoznania się ze szczegółami przygotowanej oferty oraz kontaktu w przypadku pytań lub wątpliwości.\r\n\r\nZ wyrazami szacunku,\r\nLaserTeam")

ttk.Label(left_frame, text="Read files:").grid(row=10, column=0, sticky="ne")
file_list = tk.Listbox(left_frame, height=5, width=50, bg="#3c3c3c", fg="white")
file_list.grid(row=10, column=1)

def open_selected_file(event=None):
    sel = file_list.curselection()
    if sel:
        f = os.path.join(folder_var.get(), file_list.get(sel[0]))
        try:
            os.startfile(f)
        except Exception:
            pass

file_list.bind('<Double-Button-1>', open_selected_file)

buttons_frame = tk.Frame(left_frame, bg="#2c2c2c")
buttons_frame.grid(row=11, column=1, sticky="s")

# RIGHT
right_frame = tk.Frame(root, bg="#2c2c2c")
right_frame.pack(side="right", padx=10, pady=10, fill="both", expand=True)

right_paned = tk.PanedWindow(right_frame, orient=tk.VERTICAL, bg="#2c2c2c", sashrelief="raised", borderwidth=1)
right_paned.pack(fill="both", expand=True)

panel_a = tk.PanedWindow(right_paned, orient=tk.VERTICAL, bg="#2c2c2c", sashrelief="raised", borderwidth=1)

# --- PANEL 1 ---
subpanel1 = tk.LabelFrame(panel_a, text="PANEL 1 — PREVIEW", bg="#2c2c2c", fg="white")
columns = ("1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11")
tree = ttk.Treeview(subpanel1, columns=columns, show="tree headings")
tree.column("#0", width=150, minwidth=100, stretch=tk.NO)  # Increased width for better thumbnail visibility
tree.heading("1", text="Nr");     tree.column("1", minwidth=50,  width=50,  stretch=tk.NO)
tree.heading("2", text="SubNr");  tree.column("2", minwidth=50,  width=50,  stretch=tk.NO)
tree.heading("3", text="Name");  tree.column("3", minwidth=150, width=400, stretch=tk.NO)
tree.heading("4", text="Material"); tree.column("4", minwidth=50, width=80, stretch=tk.NO)
tree.heading("5", text="Thickness");  tree.column("5", minwidth=50, width=80, stretch=tk.NO, anchor="e")
tree.heading("6", text="Quantity");    tree.column("6", minwidth=50, width=80, stretch=tk.NO, anchor="e")
tree.heading("7", text="L+M Cost");    tree.column("7", minwidth=50, width=100, stretch=tk.NO, anchor="e")
tree.heading("8", text="Bending/pc."); tree.column("8", minwidth=50, width=100, stretch=tk.NO, anchor="e")
tree.heading("9", text="Additional/pc."); tree.column("9", minwidth=50, width=120, stretch=tk.NO, anchor="e")
tree.heading("10", text="Weight"); tree.column("10", minwidth=50, width=80, stretch=tk.NO, anchor="e")
tree.heading("11", text="Cutting length"); tree.column("11", minwidth=50, width=120, stretch=tk.NO, anchor="e")

# Add scrollbar for treeview
scrollbar = ttk.Scrollbar(subpanel1, orient="vertical", command=tree.yview)
tree.configure(yscrollcommand=scrollbar.set)
tree.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")

def edit_cell(event):
    item = tree.identify_row(event.y)
    column = tree.identify_column(event.x)
    if not item or not column:
        return
    col_index = int(column[1:]) - 1
    if col_index in [5, 6, 7, 8]:
        x, y, w, h = tree.bbox(item, column)
        e = tk.Entry(subpanel1, bg="#3c3c3c", fg="white", insertbackground="white")
        e.place(x=x, y=y, width=w, height=h)
        e.insert(0, tree.item(item, 'values')[col_index])
        e.focus()
        def save_edit(_):
            vals = list(tree.item(item, 'values'))
            vals[col_index] = e.get()
            tree.item(item, values=vals)
            e.destroy()
            update_total()  # Recalculate total after edit
        e.bind("<Return>", save_edit); e.bind("<FocusOut>", save_edit)

tree.bind("<Double-1>", edit_cell)
panel_a.add(subpanel1, minsize=220)

# --- PANEL 2 ---
subpanel2 = tk.LabelFrame(panel_a, text="PANEL 2 – FIXED COSTS AND CALCULATION", bg="#2c2c2c", fg="white")

# Left column - Fixed costs
ttk.Label(subpanel2, text="Operational costs per sheet:").grid(row=0, column=0, sticky="w", padx=(5,10))
op_cost_entry = ttk.Entry(subpanel2)
op_cost_entry.grid(row=0, column=1, padx=(0,20))
op_cost_entry.insert(tk.INSERT, "40,00")

ttk.Label(subpanel2, text="Technology per order:").grid(row=1, column=0, sticky="w", padx=(5,10))
tech_order_entry = ttk.Entry(subpanel2)
tech_order_entry.grid(row=1, column=1, padx=(0,20))
tech_order_entry.insert(tk.INSERT, "50,00")

ttk.Label(subpanel2, text="Additional costs for the order:").grid(row=2, column=0, sticky="w", padx=(5,10))
add_order_cost_entry = ttk.Entry(subpanel2)
add_order_cost_entry.grid(row=2, column=1, padx=(0,20))
add_order_cost_entry.insert(tk.INSERT, "0,00")

# Right column - Cutting time calculation header
ttk.Label(subpanel2, text="CUTTING TIME CALCULATION", font=("Arial", 10, "bold")).grid(row=0, column=2, columnspan=2, pady=(0, 5), padx=(20,5))

ttk.Label(subpanel2, text="O₂ cutting rate [PLN/h]:").grid(row=1, column=2, sticky="w", padx=(20,10))
oxygen_rate_entry = ttk.Entry(subpanel2)
oxygen_rate_entry.grid(row=1, column=3, padx=(0,5))
oxygen_rate_entry.insert(tk.INSERT, "300,00")

ttk.Label(subpanel2, text="N₂ cutting rate [PLN/h]:").grid(row=2, column=2, sticky="w", padx=(20,10))
nitrogen_rate_entry = ttk.Entry(subpanel2)
nitrogen_rate_entry.grid(row=2, column=3, padx=(0,5))
nitrogen_rate_entry.insert(tk.INSERT, "350,00")

# Separator
ttk.Label(subpanel2, text="").grid(row=3, column=0, columnspan=4, pady=5)

# Time and cost displays - spanning both column groups
ttk.Label(subpanel2, text="O₂ cutting time [h]:").grid(row=4, column=0, sticky="w", padx=(5,10))
oxygen_time_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
oxygen_time_label.grid(row=4, column=1, sticky="ew", padx=(0,20))

ttk.Label(subpanel2, text="N₂ cutting time [h]:").grid(row=4, column=2, sticky="w", padx=(20,10))
nitrogen_time_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
nitrogen_time_label.grid(row=4, column=3, sticky="ew", padx=(0,5))

ttk.Label(subpanel2, text="O₂ cutting cost [PLN]:").grid(row=5, column=0, sticky="w", padx=(5,10))
oxygen_cost_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
oxygen_cost_label.grid(row=5, column=1, sticky="ew", padx=(0,20))

ttk.Label(subpanel2, text="N₂ cutting cost [PLN]:").grid(row=5, column=2, sticky="w", padx=(20,10))
nitrogen_cost_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
nitrogen_cost_label.grid(row=5, column=3, sticky="ew", padx=(0,5))

# Separator
ttk.Label(subpanel2, text="").grid(row=6, column=0, columnspan=4, pady=5)

# Summary section header
ttk.Label(subpanel2, text="COST SUMMARY", font=("Arial", 10, "bold")).grid(row=7, column=0, columnspan=4, pady=(5, 5))

# Summary fields - arranged in two columns
ttk.Label(subpanel2, text="Material cost [PLN]:").grid(row=8, column=0, sticky="w", padx=(5,10))
material_cost_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
material_cost_label.grid(row=8, column=1, sticky="ew", padx=(0,20))

ttk.Label(subpanel2, text="Total cutting cost [PLN]:").grid(row=8, column=2, sticky="w", padx=(20,10))
total_cutting_cost_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
total_cutting_cost_label.grid(row=8, column=3, sticky="ew", padx=(0,5))

ttk.Label(subpanel2, text="Operational costs [PLN]:").grid(row=9, column=0, sticky="w", padx=(5,10))
operational_cost_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
operational_cost_label.grid(row=9, column=1, sticky="ew", padx=(0,20))

# Separator before total
ttk.Label(subpanel2, text="").grid(row=10, column=0, columnspan=4, pady=5)

# Total sum - spanning columns for emphasis
ttk.Label(subpanel2, text="TOTAL OF ALL COSTS [PLN]:").grid(row=11, column=0, columnspan=2, sticky="w", padx=(5,10))
total_all_costs_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=30, font=("Arial", 11, "bold"))
total_all_costs_label.grid(row=11, column=2, columnspan=2, sticky="ew", padx=(20,5))

# Modify PANEL 2 section - add these elements after the total_all_costs_label (around row 11-12):

# Make total costs editable with new label
ttk.Label(subpanel2, text="TOTAL FOR CORRECTION [PLN]:").grid(row=13, column=0, columnspan=2, sticky="w", padx=(5,10))
total_all_costs_entry = ttk.Entry(subpanel2, width=30, font=("Arial", 11, "bold"))
total_all_costs_entry.grid(row=13, column=2, columnspan=2, sticky="ew", padx=(20,5))
total_all_costs_entry.insert(tk.INSERT, "0,00")

# Add update button
update_prices_button = ttk.Button(subpanel2, text="UPDATE PRICES BASED ON TIME", 
                                  command=lambda: update_prices_based_on_time())
update_prices_button.grid(row=14, column=0, columnspan=4, pady=(10, 5))

# Add event handlers with Enter key support
total_all_costs_entry.bind('<FocusOut>', lambda e: validate_total_entry() if all_parts else None)

# Configure column weights for proper resizing
subpanel2.grid_columnconfigure(1, weight=1)
subpanel2.grid_columnconfigure(3, weight=1)

subpanel2.update_idletasks()
panel2_height = subpanel2.winfo_reqheight() + 20
panel_a.add(subpanel2, height=panel2_height, minsize=panel2_height)

# Add event handlers for automatic recalculation
oxygen_rate_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
nitrogen_rate_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
op_cost_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
tech_order_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
add_order_cost_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
# --- PANEL 3 ---
subpanel3 = tk.LabelFrame(panel_a, text="PANEL 3 — PRICE LISTS AND TESTS", bg="#2c2c2c", fg="white", padx=6, pady=6)

mat_frame = tk.LabelFrame(subpanel3, text="Material price list (PLN/kg)", bg="#2c2c2c", fg="white")
mat_frame.grid(row=0, column=0, sticky="nwe", padx=4, pady=4)
def _update_led(canvas, ok): canvas.delete("all"); canvas.create_oval(2,2,18,18, fill=("green" if ok else "red"))
btn_load_mat = ttk.Button(mat_frame, text="Load material price list", command=lambda: load_material_prices(preview=True))
btn_load_mat.grid(row=0, column=0, sticky="w")
material_led = tk.Canvas(mat_frame, width=20, height=20, bg="#2c2c2c", highlightthickness=0); material_led.grid(row=0, column=1, padx=(6,0)); _update_led(material_led, False)
ttk.Label(mat_frame, text="Material:").grid(row=1, column=0, sticky="e", pady=(6,0))
material_var = tk.StringVar(); material_cb = ttk.Combobox(mat_frame, textvariable=material_var, width=22, state="readonly"); material_cb.grid(row=1, column=1, sticky="w", pady=(6,0))
ttk.Label(mat_frame, text="Thickness [mm]:").grid(row=2, column=0, sticky="e")
thickness_mat_var = tk.StringVar(); thickness_mat_cb = ttk.Combobox(mat_frame, textvariable=thickness_mat_var, width=12, state="readonly"); thickness_mat_cb.grid(row=2, column=1, sticky="w")
btn_find_mat = ttk.Button(mat_frame, text="Find material price", command=lambda: ui_find_material_price()); btn_find_mat.grid(row=3, column=0, columnspan=2, pady=4, sticky="we")
material_result_label = ttk.Label(mat_frame, text="Material Result: —"); material_result_label.grid(row=4, column=0, columnspan=2, sticky="w")

cut_frame = tk.LabelFrame(subpanel3, text="Cutting price list (PLN/m)", bg="#2c2c2c", fg="white")
cut_frame.grid(row=0, column=1, sticky="nwe", padx=4, pady=4)
btn_load_cut = ttk.Button(cut_frame, text="Load cutting price list", command=lambda: load_cutting_prices(preview=True))
btn_load_cut.grid(row=0, column=0, sticky="w")
cutting_led = tk.Canvas(cut_frame, width=20, height=20, bg="#2c2c2c", highlightthickness=0); cutting_led.grid(row=0, column=1, padx=(6,0)); _update_led(cutting_led, False)
ttk.Label(cut_frame, text="Material:").grid(row=1, column=0, sticky="e", pady=(6,0))
material_cut_var = tk.StringVar(); material_cut_cb = ttk.Combobox(cut_frame, textvariable=material_cut_var, width=22, state="readonly"); material_cut_cb.grid(row=1, column=1, sticky="w", pady=(6,0))
ttk.Label(cut_frame, text="Thickness [mm]:").grid(row=2, column=0, sticky="e")
thickness_cut_var = tk.StringVar(); thickness_cut_cb = ttk.Combobox(cut_frame, textvariable=thickness_cut_var, width=12, state="readonly"); thickness_cut_cb.grid(row=2, column=1, sticky="w")
ttk.Label(cut_frame, text="Gas:").grid(row=3, column=0, sticky="e")
gas_var = tk.StringVar(); gas_cb = ttk.Combobox(cut_frame, textvariable=gas_var, width=12, state="readonly"); gas_cb.grid(row=3, column=1, sticky="w")
btn_find_cut = ttk.Button(cut_frame, text="Find cutting price", command=lambda: ui_find_cutting_price()); btn_find_cut.grid(row=4, column=0, columnspan=2, pady=4, sticky="we")
cutting_result_label = ttk.Label(cut_frame, text="Cutting Result: —"); cutting_result_label.grid(row=5, column=0, columnspan=2, sticky="w")

btn_load_both = ttk.Button(subpanel3, text="Load both price lists and refresh lists",
                          command=lambda: (load_material_prices(True), load_cutting_prices(True)))
btn_load_both.grid(row=1, column=0, columnspan=2, sticky="we", padx=4, pady=(2,6))
subpanel3.grid_columnconfigure(0, weight=1); subpanel3.grid_columnconfigure(1, weight=1)

panel_a.add(subpanel3, minsize=200)
right_paned.add(panel_a)

def update_cost_calculations():
    """Update all cost calculation displays in Panel 2"""
    global oxygen_cutting_time, nitrogen_cutting_time, total_material_cost
    
    # Get rates from entries
    oxygen_rate = _parse_float(oxygen_rate_entry.get()) or 0.0
    nitrogen_rate = _parse_float(nitrogen_rate_entry.get()) or 0.0
    op_cost_per_sheet = _parse_float(op_cost_entry.get()) or 0.0
    tech_per_order = _parse_float(tech_order_entry.get()) or 0.0
    add_costs_order = _parse_float(add_order_cost_entry.get()) or 0.0
    
    # Calculate cutting costs
    oxygen_cost = oxygen_cutting_time * oxygen_rate
    nitrogen_cost = nitrogen_cutting_time * nitrogen_rate
    total_cutting_cost = oxygen_cost + nitrogen_cost
    
    # Calculate operational costs
    operational_costs = (total_sheets * op_cost_per_sheet) + tech_per_order + add_costs_order
    
    # Calculate base total
    base_total = total_material_cost + total_cutting_cost + operational_costs
    
    # Update display labels
    oxygen_time_label.config(text=f"{oxygen_cutting_time:.2f}".replace('.', ','))
    nitrogen_time_label.config(text=f"{nitrogen_cutting_time:.2f}".replace('.', ','))
    oxygen_cost_label.config(text=format_pln(oxygen_cost))
    nitrogen_cost_label.config(text=format_pln(nitrogen_cost))
    material_cost_label.config(text=format_pln(total_material_cost))
    total_cutting_cost_label.config(text=format_pln(total_cutting_cost))
    operational_cost_label.config(text=format_pln(operational_costs))
    
    # Update the editable total field
    total_all_costs_entry.delete(0, tk.END)
    
def validate_total_entry():
    """Validate and format the manually entered total"""
    try:
        # Get text from the input field
        value_str = total_all_costs_entry.get().strip()
        
        # Remove spaces and replace commas with dots as decimal separator
        value_str = value_str.replace(' ', '').replace(',', '.')
        
        # Remove everything except digits and dot to avoid errors with other characters
        value_str = ''.join(c for c in value_str if c.isdigit() or c == '.')
        
        # Convert to float if the string is not empty
        if value_str:
            value = float(value_str)
            if value is not None:
                total_all_costs_entry.delete(0, tk.END)
                total_all_costs_entry.insert(0, format_pln(value))
    except ValueError:
        pass  # Ignore errors if conversion fails

def update_prices_based_on_time():
    """Update unit prices in treeview based on time calculations and proportional distribution"""
    global all_parts, total_row_iid
    
    if not all_parts:
        messagebox.showwarning("Warning", "No data to update. Perform analysis first.")
        return
    
    # Get the target total from the editable field
    value_str = total_all_costs_entry.get().strip()
    value_str = value_str.replace(' ', '').replace(',', '.')  # Remove spaces and replace commas with dots
    value_str = ''.join(c for c in value_str if c.isdigit() or c == '.')  # Keep only digits and dots
    target_total = _parse_float(value_str) if value_str else None
    if not target_total or target_total <= 0:
        messagebox.showerror("Error", "Invalid total costs.")
        return
    
    # Calculate current total from treeview (excluding total row)
    current_total = 0.0
    items_data = []
    
    for item in tree.get_children():
        if item == total_row_iid:
            continue
        vals = tree.item(item, 'values')
        qty = _parse_float(vals[5]) or 0
        cost = _parse_float(vals[6]) or 0
        bending = _parse_float(vals[7]) or 0
        additional = _parse_float(vals[8]) or 0
        
        item_total = (cost + bending + additional) * qty
        current_total += item_total
        
        items_data.append({
            'item': item,
            'qty': qty,
            'cost': cost,
            'bending': bending,
            'additional': additional,
            'item_total': item_total,
            'proportion': 0.0
        })
    
    if current_total <= 0:
        messagebox.showerror("Error", "No costs to recalculate.")
        return
    
    # Calculate proportions for each item
    for item_data in items_data:
        item_data['proportion'] = item_data['item_total'] / current_total
    
    # Apply proportional distribution of the new total
    new_grand_total = 0.0
    for idx, item_data in enumerate(items_data):
        # Calculate new item total based on proportion
        new_item_total = target_total * item_data['proportion']
        
        # Calculate new unit cost (preserving bending and additional costs)
        if item_data['qty'] > 0:
            new_unit_cost = (new_item_total / item_data['qty']) - item_data['bending'] - item_data['additional']
            
            # Ensure non-negative cost
            new_unit_cost = max(0, new_unit_cost)
            
            # Update tree item
            vals = list(tree.item(item_data['item'], 'values'))
            vals[6] = format_pln(new_unit_cost)
            tree.item(item_data['item'], values=vals)
            
            # Update all_parts array
            if idx < len(all_parts):
                all_parts[idx]['cost_per_unit'] = new_unit_cost
            
            new_grand_total += (new_unit_cost + item_data['bending'] + item_data['additional']) * item_data['qty']
    
    # Update the total row
    tree.set(total_row_iid, column="7", value=format_pln(new_grand_total))
    SetTotalPricePerOrder(new_grand_total)
    
    messagebox.showinfo("Success", f"Prices have been updated proportionally.\n"
                                  f"Old sum: {format_pln(current_total)}\n"
                                  f"New sum: {format_pln(new_grand_total)}")

def update_total():
    """Recalculate total after manual edits in treeview"""
    global total_row_iid
    total = 0.0
    for item in tree.get_children():
        if item == total_row_iid:
            continue
        vals = tree.item(item, 'values')
        qty = _parse_float(vals[5]) or 0
        cost = _parse_float(vals[6]) or 0
        bending = _parse_float(vals[7]) or 0
        additional = _parse_float(vals[8]) or 0
        total += (cost + bending + additional) * qty
    tree.set(total_row_iid, column="7", value=format_pln(total))
    SetTotalPricePerOrder(total)

# ---- Price list loaders ----
def _tree_preview_clear_and_headers(headers):
    for item in tree.get_children():
        tree.delete(item)
    tree.insert('', 'end', values=(0, '', ' | '.join(headers), '', '', '', '', '', ''))

def load_material_prices(preview=False):
    global material_prices, _mat_set, _thk_set
    material_prices.clear(); _mat_set.clear(); _thk_set.clear()
    try:
        if not os.path.exists(MATERIALS_FILE):
            raise FileNotFoundError(f"File not found: {MATERIALS_FILE}")
        wb = load_workbook(MATERIALS_FILE, data_only=True)
        sheet = wb.active
        headers = [str(c.value).strip().lower() if c.value is not None else "" for c in next(sheet.iter_rows(min_row=1, max_row=1))]
        need = ("material", "thickness", "price")
        idx = {n: headers.index(n) for n in need if n in headers}
        if not set(need).issubset(idx):
            raise ValueError("Missing required columns: material, thickness, price")
        if preview: _tree_preview_clear_and_headers(["materials prices.xlsx → material/thickness/price"])
        for row in sheet.iter_rows(min_row=2, values_only=True):
            mat = _norm_s(row[idx["material"]]); thk = _parse_float(row[idx["thickness"]]); prc = _parse_float(row[idx["price"]])
            if mat and thk is not None and prc is not None:
                material_prices[(mat, thk)] = prc; _mat_set.add(mat); _thk_set.add(thk)
                if preview: tree.insert('', 'end', values=("", "", f"{mat} @ {thk:.2f} mm → {format_pln(prc)} PLN/kg", "", "", "", "", "", ""))
        mats_sorted = sorted(_mat_set); thk_sorted = [f"{t:.2f}".rstrip("0").rstrip(".") for t in sorted(_thk_set)]
        material_cb["values"] = mats_sorted; material_cut_cb["values"] = mats_sorted
        thickness_mat_cb["values"] = thk_sorted
        if not thickness_cut_cb["values"]: thickness_cut_cb["values"] = thk_sorted
        _update_led(material_led, len(material_prices) > 0)
    except Exception as e:
        _update_led(material_led, False); messagebox.showerror("Error", f"Loading material prices:\n{e}")

def load_cutting_prices(preview=False):
    global cutting_prices, _mat_set, _thk_set, _gas_set
    cutting_prices.clear(); _gas_set.clear()
    try:
        if not os.path.exists(CUTTING_FILE):
            raise FileNotFoundError(f"File not found: {CUTTING_FILE}")
        wb = load_workbook(CUTTING_FILE, data_only=True)
        sheet = wb.active
        headers = [str(c.value).strip().lower() if c.value is not None else "" for c in next(sheet.iter_rows(min_row=1, max_row=1))]
        need = ("thickness", "material", "gas", "price")
        idx = {n: headers.index(n) for n in need if n in headers}
        if not set(need).issubset(idx):
            raise ValueError("Missing required columns: thickness, material, gas, price")
        if preview: _tree_preview_clear_and_headers(["cutting prices.xlsx → thickness/material/gas/price"])
        for row in sheet.iter_rows(min_row=2, values_only=True):
            thk = _parse_float(row[idx["thickness"]]); mat = _norm_s(row[idx["material"]]); gas = _norm_s(row[idx["gas"]]); prc = _parse_float(row[idx["price"]])
            if thk is not None and mat and gas and prc is not None:
                cutting_prices[(thk, mat, gas)] = prc; _mat_set.add(mat); _thk_set.add(thk); _gas_set.add(gas)
                if preview: tree.insert('', 'end', values=("", "", f"{thk:.2f} mm / {mat} / {gas} → {format_pln(prc)} PLN/m", "", "", "", "", "", ""))
        mats_sorted = sorted(_mat_set); thk_sorted = [f"{t:.2f}".rstrip("0").rstrip(".") for t in sorted(_thk_set)]; gas_sorted = sorted(_gas_set)
        material_cut_cb["values"] = mats_sorted
        if not material_cb["values"]: material_cb["values"] = mats_sorted
        thickness_cut_cb["values"] = thk_sorted
        if not thickness_mat_cb["values"]: thickness_mat_cb["values"] = thk_sorted
        gas_cb["values"] = gas_sorted
        _update_led(cutting_led, len(cutting_prices) > 0)
    except Exception as e:
        _update_led(cutting_led, False); messagebox.showerror("Error", f"Loading cutting prices:\n{e}")

# ---- UI tests (Panel 3) ----
def ui_find_material_price():
    mat = _norm_s(material_var.get()); thk = _parse_float(thickness_mat_var.get())
    if not mat or thk is None:
        messagebox.showerror("Error", "Fill in Material and Thickness (mm)."); return
    price = material_prices.get((mat, thk))
    material_result_label.config(text="Material Result: not found" if price is None else f"Material Result: {format_pln(price)} PLN/kg")

def ui_find_cutting_price():
    mat = _norm_s(material_cut_var.get()); thk = _parse_float(thickness_cut_var.get()); gas = _norm_s(gas_var.get())
    if not mat or thk is None or not gas:
        messagebox.showerror("Error", "Fill in Material, Thickness (mm) and Gas."); return
    price = cutting_prices.get((thk, mat, gas))
    cutting_result_label.config(text="Cutting Result: not found" if price is None else f"Cutting Result: {format_pln(price)} PLN/m")

# ---- Folder analysis ----
last_groups = []; last_total_cost = 0.0; last_folder_path = ""

def _ensure_cenniki_loaded():
    ok = True
    if not material_prices:
        try: load_material_prices(preview=False)
        except Exception: ok = False
    if not cutting_prices:
        try: load_cutting_prices(preview=False)
        except Exception: ok = False
    return ok


def get_total_cut_length(ws, text="Total") -> float:
    """
    Searches in column A for the first cell containing 'text' (case-insensitive),
    then reads the value from column H (8) in the same row.
    Returns float; handles Polish format '312,51'.
    """
    # We only iterate over column A — it's fast and simple
    for cell in ws['A']:
        val = cell.value
        if val and str(text).lower() in str(val).lower():
            raw = ws.cell(row=cell.row, column=8).value  # col. H
            # openpyxl usually returns float for numeric cells;
            # if it's a string in '312,51' format, replace comma:
            if isinstance(raw, (int, float)):
                return float(raw)
            try:
                return float(str(raw).replace(" ", "").replace("\xa0", "").replace(",", "."))
            except Exception:
                return 0.0
    raise ValueError(f"No row with text '{text}' found in column A")


def parse_duration_to_hours(value) -> float:
    """
    Converts '1h26min21s', '1h26m21s', '86min', '90s', '1:26:21', '1:26' etc. to float hours.
    Returns 0.0 if unable to parse.
    """
    if value is None:
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)

    s = str(value).strip().lower().replace(" ", "")

    # H:M:S or H:M
    if ":" in s:
        parts = s.split(":")
        try:
            def to_f(x): return float(x.replace(",", "."))
            if len(parts) == 3:
                h, m, sec = parts
            elif len(parts) == 2:
                h, m = parts; sec = "0"
            else:
                h, m, sec = parts[0], "0", "0"
            return to_f(h) + to_f(m)/60.0 + to_f(sec)/3600.0
        except Exception:
            pass

    def take(pattern):
        m = re.search(pattern, s)
        return float(m.group(1).replace(",", ".")) if m else 0.0

    h = take(r"(\d+(?:[.,]\d+)?)h")
    m = take(r"(\d+(?:[.,]\d+)?)(?:m|min)")
    sec = take(r"(\d+(?:[.,]\d+)?)s")

    if h == 0.0 and m == 0.0 and sec == 0.0:
        # Just '123' → treat as seconds
        try:
            seconds = float(s.replace(",", "."))
            return seconds / 3600.0
        except Exception:
            return 0.0

    return h + m/60.0 + sec/3600.0


# ---- Dynamic pricing based on thickness, length and time ----
_THK_SPEED_MPM = [
    (1.0, 18.0),
    (1.5, 18.0),
    (2.0, 14.0),
    (3.0, 4.0),
    (5.0, 3.5),
    (6.0, 3.0),
    (8.0, 2.7),
    (10.0, 2.1),
    (12.0, 2.1),
    (15.0, 2.1),
]

def _interp(x, pts):
    """Linear interpolation over sorted (x, y) points. Clamps outside the range."""
    pts = sorted(pts, key=lambda p: p[0])
    if x <= pts[0][0]:
        return pts[0][1]
    if x >= pts[-1][0]:
        return pts[-1][1]
    for (x0, y0), (x1, y1) in zip(pts[:-1], pts[1:]):
        if x0 <= x <= x1:
            if x1 == x0:
                return y0
            t = (x - x0) / (x1 - x0)
            return y0 + t * (y1 - y0)
    return pts[-1][1]

def get_speed_mpm(thickness_mm: float) -> float:
    """Return cutting speed [m/min] for a given thickness using piecewise-linear interpolation."""
    if thickness_mm is None:
        return 0.0
    return float(_interp(float(thickness_mm), _THK_SPEED_MPM))

def get_time_thresholds_minutes(thickness_mm: float):
    """
    Returns (t_min, t_neutral) in minutes for the given thickness.
    Anchors: 1 mm → (1, 45), 15 mm → (5, 90). Linearly and clamped.
    """
    thk = max(1.0, min(float(thickness_mm or 1.0), 15.0))
    t_min = 1.0 + 4.0 * (thk - 1.0) / 14.0       # 1 → 5
    t_neutral = 45.0 + 45.0 * (thk - 1.0) / 14.0 # 45 → 90
    return (t_min, t_neutral)

def compute_effective_minutes(thickness_mm: float, total_cut_length_m: float, cut_time_hours: float) -> float:
    """
    Effective order time [min] based on path length and given cut_time.
    Calculate time from length: L / v(thk), take maximum with the cell value.
    """
    speed = get_speed_mpm(thickness_mm)  # m/min
    t_from_len = 0.0
    if total_cut_length_m is not None and speed > 0:
        try:
            t_from_len = float(total_cut_length_m) / float(speed)
        except Exception:
            t_from_len = 0.0
    t_from_cell = float(cut_time_hours or 0.0) * 60.0
    if t_from_len <= 0:
        return t_from_cell
    if t_from_cell <= 0:
        return t_from_len
    return max(t_from_len, t_from_cell)

def compute_boost_factor(thickness_mm: float, effective_minutes: float, max_boost: float = 3.5) -> float:
    """
    Multiplier in the range [1.0, max_boost].
    = max_boost when time ≤ t_min; = 1.0 when time ≥ t_neutral; linear between.
    """
    t_min, t_neutral = get_time_thresholds_minutes(thickness_mm)
    if effective_minutes <= t_min:
        return float(max_boost)
    if effective_minutes >= t_neutral:
        return 1.0
    alpha = (t_neutral - effective_minutes) / (t_neutral - t_min)  # (0,1)
    return 1.0 + (max_boost - 1.0) * alpha

def apply_dynamic_pricing(price_per_kg: float, rate_per_m: float, thickness_mm: float, total_cut_length_m: float, cut_time_hours: float):
    """
    Rules:
    - material: MINIMUM +7% above purchase price,
    - material and cutting: additional 'boost' multiplier depending on time (1.0 → 3.5 = +250%) according to thresholds,
    - material: max(1.07, boost) × purchase price,
      cutting:   boost × rate per meter.
    Returns (material_price_adj, cutting_rate_adj, debug_dict).
    """
    try:
        base_material = float(price_per_kg or 0.0)
        base_rate = float(rate_per_m or 0.0)
    except Exception:
        base_material, base_rate = 0.0, 0.0

    eff_minutes = compute_effective_minutes(thickness_mm, total_cut_length_m, cut_time_hours)
    boost = compute_boost_factor(thickness_mm, eff_minutes, max_boost=3.5)

    material_factor = max(1.07, boost)  # minimally +7%
    cutting_factor = boost

    mat_adj = base_material * material_factor
    rate_adj = base_rate * cutting_factor

    dbg = {
        'eff_minutes': eff_minutes,
        't_min_t_neutral': get_time_thresholds_minutes(thickness_mm),
        'boost': boost,
        'material_factor': material_factor,
        'cutting_factor': cutting_factor,
    }
    return mat_adj, rate_adj, dbg


# References to PhotoImage to prevent images from disappearing (GC)
thumbnail_imgs = []

def analyze_xlsx_folder():
    global all_parts, last_groups, last_total_cost, last_folder_path, total_sheets, total_parts_qty, total_row_iid
    global oxygen_cutting_time, nitrogen_cutting_time, total_material_cost
    
    # Initialize cutting time accumulators
    oxygen_cutting_time = 0.0
    nitrogen_cutting_time = 0.0
    total_material_cost = 0.0

    for item in tree.get_children():
        tree.delete(item)
    thumbnail_imgs.clear()
    all_parts = []
    folder_path = folder_var.get()
    if not folder_path:
        messagebox.showerror("Error", "Please select a folder."); return
    try:
        files = [f for f in os.listdir(folder_path) if f.lower().endswith(".xlsx")]
    except Exception:
        files = []
    if not files:
        messagebox.showerror("Error", "No .xlsx files in the selected folder."); return
    if not _ensure_cenniki_loaded():
        messagebox.showwarning("Warning", "Price lists not loaded — using 0.00, check Panel 3.")

    global op_cost_per_sheet, tech_per_order, add_costs_order
    op_cost_per_sheet = _parse_float(op_cost_entry.get()) or 0.0
    tech_per_order = _parse_float(tech_order_entry.get()) or 0.0
    add_costs_order = _parse_float(add_order_cost_entry.get()) or 0.0

    total_sheets = 0
    total_parts_qty = 0
    groups = []
    subnr = 0

    thumbnails = {}  # To store raw image data for each file if needed

    for fname in files:
        path = os.path.join(folder_path, fname)
        try:
            wb = load_workbook(path, data_only=True)
            if "All Task List" not in wb.sheetnames:
                raise KeyError("No 'All Task List' sheet")
            all_task = wb["All Task List"]
            all_part_list = wb["All Parts List"]
            file_thumbnails = {}
            for img in all_part_list._images:
                row = img.anchor._from.row + 1
                col = img.anchor._from.col + 1
                if col == 2:  # Column B
                    img_data = img._data()
                    file_thumbnails[row] = img_data
            cut_time = parse_duration_to_hours(all_task['F4'].value)
            total_cut_length = get_total_cut_length(all_task, "Total")
            material_name = all_task["B4"].value
            thickness_raw = all_task["C4"].value
            gas_raw = all_task["E4"].value
            mat_norm = _norm_s(material_name)
            thk_val = _parse_float(thickness_raw)
            gas_key = _map_gas_to_key(gas_raw)

            # Add accumulation of cutting time by gas type:
            if gas_key == "O":
                oxygen_cutting_time += cut_time
            elif gas_key == "N":
                nitrogen_cutting_time += cut_time

            if not mat_norm:
                raise ValueError("All Task List!B4 (Material) — no value")
            if thk_val is None:
                raise ValueError("All Task List!C4 (Thickness(mm)) — no number")
            if not gas_key:
                raise ValueError("All Task List!E4 (Gas) — unsupported gas type")

            base_price_per_kg = material_prices.get((mat_norm, thk_val), 0.0)
            base_rate_per_cut_length = cutting_prices.get((thk_val, mat_norm, gas_key), 0.0)

            price_per_kg, rate_per_cut_length, _dpdbg = apply_dynamic_pricing(
                base_price_per_kg, base_rate_per_cut_length, thk_val, total_cut_length, cut_time
            )

            if "Cost List" not in wb.sheetnames:
                raise KeyError("No 'Cost List' sheet")
            cost_sheet = wb["Cost List"]

            util_row = None
            for r in range(1, cost_sheet.max_row + 1):
                for c in range(1, cost_sheet.max_column + 1):
                    v = cost_sheet.cell(row=r, column=c).value
                    if v and "Average utilization:" in str(v):
                        util_row = r
                        break
                if util_row:
                    break
            if util_row is None:
                raise ValueError("Not found 'Average utilization:'")
            util_str = cost_sheet.cell(row=util_row, column=11).value
            util_val = _parse_float(str(util_str).replace("%", "")) if util_str is not None else None
            utilization_rate = (util_val / 100.0) if (util_val is not None) else 0.0
            if utilization_rate <= 0 or utilization_rate > 1:
                messagebox.showwarning("Warning", f"Average utilization out of range ({utilization_rate}).")

            mat_price_row = None
            for r in range(1, cost_sheet.max_row + 1):
                v = cost_sheet.cell(row=r, column=1).value
                if v and "Material Price" in str(v):
                    mat_price_row = r
                    break
            if mat_price_row is None:
                raise ValueError("No 'Material Price' row")

            def parse_num(cellv):
                if cellv is None:
                    return 0.0
                s = str(cellv).strip()
                s = ''.join(ch for ch in s.split()[0] if ch.isdigit() or ch in ('.', ','))
                s = s.replace(",", ".")
                try:
                    return float(s or "0.0")
                except Exception:
                    return 0.0

            rate_per_contour = parse_num(cost_sheet.cell(row=mat_price_row, column=7).value)  # G
            rate_per_marking_length = parse_num(cost_sheet.cell(row=mat_price_row, column=9).value)  # I
            rate_per_defilm_length = parse_num(cost_sheet.cell(row=mat_price_row, column=10).value)  # J

            r_idx = 8
            while all_task.cell(row=r_idx, column=4).value is not None:
                v = all_task.cell(row=r_idx, column=4).value
                if isinstance(v, (int, float)):
                    total_sheets += int(v)
                r_idx += 1

            start_row = None
            for r in range(1, cost_sheet.max_row + 1):
                a_val = cost_sheet.cell(row=r, column=1).value
                if a_val and isinstance(a_val, (int, float)):
                    start_row = r
                    break
            if start_row is None:
                raise ValueError("No starting row found (col. A — ID)")

            parts_for_group = []
            subnr += 1
            lp = 0
            row = start_row
            while row <= cost_sheet.max_row and isinstance(cost_sheet.cell(row=row, column=1).value, (int, float)):
                lp += 1
                part_name = cost_sheet.cell(row=row, column=2).value
                part_qty = cost_sheet.cell(row=row, column=5).value or 0
                weight = parse_num(cost_sheet.cell(row=row, column=6).value)
                contours_qty = parse_num(cost_sheet.cell(row=row, column=7).value)
                cut_length = parse_num(cost_sheet.cell(row=row, column=8).value)
                marking_length = parse_num(cost_sheet.cell(row=row, column=9).value)
                defilm_length = parse_num(cost_sheet.cell(row=row, column=10).value)

                adj_weight = (weight / utilization_rate) if utilization_rate > 0 else weight

                # Base costs
                base_material_cost = adj_weight * base_price_per_kg
                base_cut_cost = cut_length * base_rate_per_cut_length
                base_total_part = base_material_cost + contours_qty * rate_per_contour + base_cut_cost + marking_length * rate_per_marking_length + defilm_length * rate_per_defilm_length

                # Dynamic costs
                material_cost = adj_weight * price_per_kg
                cut_cost = cut_length * rate_per_cut_length
                total_part = material_cost + contours_qty * rate_per_contour + cut_cost + marking_length * rate_per_marking_length + defilm_length * rate_per_defilm_length

                thumbnail_data = None
                all_parts_row = 2 + lp  # Assuming row 3 is first part
                if all_parts_row in file_thumbnails:
                    thumbnail_data = file_thumbnails[all_parts_row]

                thumbnail_photo = None
                if thumbnail_data:
                    try:
                        pil_img = Image.open(io.BytesIO(thumbnail_data))
                        max_w, max_h = 140, 70
                        w, h = pil_img.size
                        ratio = min(max_w / w, max_h / h, 1.0)
                        new_w = int(w * ratio)
                        new_h = int(h * ratio)
                        pil_img = pil_img.resize((new_w, new_h), Image.LANCZOS)
                        thumbnail_photo = ImageTk.PhotoImage(pil_img)
                        thumbnail_imgs.append(thumbnail_photo)
                    except Exception:
                        pass

                all_parts.append({
                    'id': lp,
                    'subnr': subnr,
                    'name': part_name,
                    'material': material_name,
                    'thickness': thk_val,
                    'qty': int(part_qty) if isinstance(part_qty, (int, float)) else 0,
                    'cost_per_unit': float(f"{total_part:.2f}"),
                    'base_cost_per_unit': float(f"{base_total_part:.2f}"),
                    'bending_per_unit': 0.0,
                    'additional_per_unit': 0.0,
                    'raw_weight': weight,
                    'contours_qty': contours_qty,
                    'cut_length': cut_length,
                    'marking_length': marking_length,
                    'defilm_length': defilm_length,
                    'adj_weight': adj_weight,
                    'base_price_per_kg': base_price_per_kg,
                    'base_rate_per_cut_length': base_rate_per_cut_length,
                    'base_cut_cost': base_cut_cost,
                    'rate_per_contour': rate_per_contour,
                    'rate_per_marking_length': rate_per_marking_length,
                    'rate_per_defilm_length': rate_per_defilm_length,
                    'thumb_data': thumbnail_data,  # Raw bytes
                })

                parts_for_group.append((part_name, float(f"{total_part:.2f}"),
                                        int(part_qty) if isinstance(part_qty, (int, float)) else 0))
                total_parts_qty += int(part_qty) if isinstance(part_qty, (int, float)) else 0
                row += 1

            groups.append((material_name, thk_val, parts_for_group))

        except Exception as e:
            messagebox.showerror("Error", f"Error processing file {fname}: {e}")
            return

    # distribution of overheads per piece
    if total_parts_qty > 0:
        extra_per_part = (tech_per_order + add_costs_order) / total_parts_qty
        op_cost_per_part = (total_sheets * op_cost_per_sheet) / total_parts_qty
    else:
        extra_per_part = 0.0
        op_cost_per_part = 0.0

    for p in all_parts:
        p['cost_per_unit'] += extra_per_part + op_cost_per_part
        p['base_cost_per_unit'] += extra_per_part + op_cost_per_part
        p['cost_per_unit'] = float(f"{p['cost_per_unit']:.2f}")
        p['base_cost_per_unit'] = float(f"{p['base_cost_per_unit']:.2f}")

   # Calculate material costs for all parts
    for p in all_parts:
        material_cost_per_part = p['adj_weight'] * p.get('base_price_per_kg', 0.0)
        total_material_cost += material_cost_per_part * p['qty']
    
    # Update Panel 2 display fields
    update_cost_calculations()


    # table
    for i, p in enumerate(all_parts, start=1):
        item_values = (
            i,
            p['subnr'],
            p['name'],
            p['material'],
            f"{p['thickness']}",
            p['qty'],
            format_pln(p['cost_per_unit']),
            "",
            "",
            format_pln(p['adj_weight']),
            format_pln(p['cut_length']),
        )
        opts = {'values': item_values}
        if p['thumb_data']:
            try:
                pil_img = Image.open(io.BytesIO(p['thumb_data']))
                max_w, max_h = 140, 70
                w, h = pil_img.size
                ratio = min(max_w / w, max_h / h, 1.0)
                new_w = int(w * ratio)
                new_h = int(h * ratio)
                pil_img = pil_img.resize((new_w, new_h), Image.LANCZOS)
                thumb = ImageTk.PhotoImage(pil_img)
                thumbnail_imgs.append(thumb)
                opts['image'] = thumb
            except Exception:
                pass

        tree.insert('', 'end', **opts)

    # Add total row to treeview
    total_order = sum(p['cost_per_unit'] * p['qty'] for p in all_parts)
    SetTotalPricePerOrder(total_order)
    total_row_iid = tree.insert('', 'end', values=('', '', 'Total', '', '', '', format_pln(total_order), '', '', '', ''))

   # Create merged groups (this code should already exist in your function)
    total_sum = 0.0
    merged_groups = []
    for (mat_name, thk, parts) in groups:
        adj = []
        for (nm, cost, qty) in parts:
            c = float(f"{cost:.2f}")
            adj.append((nm, c, qty))
            total_sum += c * qty
        merged_groups.append((mat_name, thk, adj))
    
    # NOW you can safely assign these variables
    last_groups = merged_groups
    last_total_cost = total_sum
    last_folder_path = folder_path
    
    messagebox.showinfo("Analysis", "XLSX files analysis completed. Data in Panel 1 filled.")


def get_next_offer_number():
    month_year = datetime.datetime.now().strftime("%m/%Y")
    month_key = datetime.datetime.now().strftime("counter_%Y-%m")
    try:
        response = requests.get(f"https://abacus.jasoncameron.dev/hit/xai_offers/{month_key}")
        if response.status_code == 200:
            counter_value = int(response.json()['value'])
            return f"Laser/{counter_value:04d}/{month_year}"
        else:
            return "Laser/0001/08/2025"  # Fallback
    except Exception:
        return "Laser/0001/08/2025"  # Fallback

# report
def generate_report():
    if not all_parts:
        messagebox.showerror("Error", "No data for the report. First 'Analyze XLSX'.")
        return
    folder_path = folder_var.get().strip() or last_folder_path
    if not folder_path or not os.path.isdir(folder_path):
        messagebox.showerror("Error", "Invalid target folder.")
        return
    customer_name = customer_var.get().strip() or "Client"
    offer_number = offer_var.get().strip()
    if not offer_number:
        offer_number = get_next_offer_number()
        offer_var.set(offer_number)
    offer_date = date_var.get().strip() or datetime.datetime.now().strftime("%Y-%m-%d")
    validity = validity_var.get().strip() or (datetime.datetime.now() + timedelta(days=14)).strftime("%Y-%m-%d")
    logo_path = logo_var.get().strip()
    contact_details = contact_text.get("1.0", tk.END).strip()
    preceding_text = preceding_text_var.get("1.0", tk.END).strip()
    finishing_text = finishing_text_var.get("1.0", tk.END).strip()

    # Create Raporty folder
    raporty_path = os.path.join(folder_path, "Raporty")
    os.makedirs(raporty_path, exist_ok=True)

    # Update all_parts from tree
    tree_items = tree.get_children()
    if len(tree_items) != len(all_parts) + 1:  # +1 for total row
        messagebox.showerror("Error", "Data mismatch between table and parts list.")
        return

    for idx, item in enumerate(tree_items[:-1]):  # Exclude total row
        vals = tree.item(item, 'values')
        if len(vals) < 11:
            vals = vals + ('', '') * (11 - len(vals))
        all_parts[idx]['qty'] = int(vals[5] or 0)
        all_parts[idx]['cost_per_unit'] = _parse_float(vals[6]) or 0.0
        all_parts[idx]['bending_per_unit'] = _parse_float(vals[7]) or 0.0
        all_parts[idx]['additional_per_unit'] = _parse_float(vals[8]) or 0.0


    # Log start
    log_path = os.path.join(raporty_path, "cost_calculation_log.txt")
    with open(log_path, 'w', encoding='utf-8') as log:
        log.write(f"Calculation Log - {datetime.datetime.now()}\n")
        log.write(f"Folder: {folder_path}\n")
        log.write(f"Client: {customer_name}\n")
        log.write("Price sources: materials prices.xlsx, cutting prices.xlsx\n")
        log.write("\nCalculation Details:\n")

    # Generate DOCX
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    if logo_path and os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(3.0))
        except Exception:
            pass
    if contact_details:
        p = doc.add_paragraph(contact_details)
        for r in p.runs:
            r.bold = False

    doc.add_heading(f"Offer for {customer_name}", level=1)
    p = doc.add_paragraph(f"Offer number: {offer_number}")
    p.runs[0].bold = True
    doc.add_paragraph(f"Offer date: {offer_date}")
    doc.add_paragraph(f"Validity period: {validity}")
    if preceding_text:
        doc.add_paragraph(preceding_text)

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Lp.'
    hdr[1].text = 'Miniatura'
    hdr[2].text = 'Part name'
    hdr[3].text = 'Quantity'
    hdr[4].text = 'Net weight'
    hdr[5].text = 'Cost (PLN)'
    hdr[6].text = 'Total (PLN)'
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '006995')
        tcPr.append(shd)
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True
        run.font.size = Pt(11)

    total = 0.0
    lp = 1
    for mat_name, thk, parts in last_groups:
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = f"Material: {mat_name}, Thickness: {thk} mm"
        row[2].merge(row[6])
        run = row[2].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.italic = True
        for nm, cost_per_unit, qty in parts:
            part = next(p for p in all_parts if p['name'] == nm)
            r = table.add_row().cells
            r[0].text = str(lp)
            # Embed graphic in column 2 (Miniatura)
            if part['thumb_data']:
                try:
                    run = r[1].add_paragraph().add_run()
                    run.add_picture(io.BytesIO(part['thumb_data']))
                except Exception:
                    pass
            r[2].text = str(nm) if nm else "No name"
            r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r[3].paragraphs[0].add_run(f"{int(part['qty'])}  ").font.size = Pt(10)
            r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r[4].paragraphs[0].add_run(f"{format_pln(part['raw_weight'])}  ").font.size = Pt(10)
            r[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r[5].paragraphs[0].add_run(f"{format_pln(part['cost_per_unit'] + part['bending_per_unit'] + part['additional_per_unit'])}  ").font.size = Pt(10)
            row_total = (part['cost_per_unit'] + part['bending_per_unit'] + part['additional_per_unit']) * part['qty']
            r[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r[6].paragraphs[0].add_run(f"{format_pln(row_total)}  ").font.size = Pt(10)
            total += row_total
            lp += 1

    srow = table.add_row().cells
    srow[1].text = ""
    srow[2].text = "Total"
    srow[4].text = ""
    srow[6].text = format_pln(total)
    srow[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in srow[2].paragraphs[0].runs:
        r.bold = True
    for r in srow[6].paragraphs[0].runs:
        r.bold = True

    widths = [Cm(1), Cm(2), Cm(6), Cm(2), Cm(2), Cm(3), Cm(3)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    p = doc.add_paragraph(f"Total cost: {format_pln(total)} PLN")
    p.paragraph_format.space_before = Pt(12)
    for r in p.runs:
        r.font.size = Pt(14)
    if finishing_text:
        pf = doc.add_paragraph(finishing_text)
        for r in pf.runs:
            r.font.size = Pt(9)

    current_date = datetime.datetime.now().strftime("%Y%m%d")
    # Keep the '/' replacement for filename
    fname = f"Oferta_{sanitize_filename(customer_name) or 'Klient'}_{current_date}_{offer_number.replace('/', '-')}.docx"
    full = os.path.join(raporty_path, fname)
    try:
        doc.save(full)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save DOCX file:\n{e}")
        return


    # Generate enhanced cost report XLSX
    cost_wb = Workbook()
    
    # Sheet 1: Detailed Cost Breakdown with Thumbnail as column 2
    detail_ws = cost_wb.active
    detail_ws.title = "Detailed calculation"
    
    # Headers with Miniatura as second column
    headers = [
        "ID", "Miniatura", "Part name", "Material", "Thickness [mm]", "Quantity [pcs]",
        "Unit weight [kg]", "Adjusted weight [kg]", "Cutting length [m]",
        "Number of contours", "Marking length [m]", "Defilm length [m]",
        "Material price [PLN/kg]", "Cutting rate [PLN/m]", 
        "Material cost [PLN]", "Cutting cost [PLN]", "Contours cost [PLN]",
        "Marking cost [PLN]", "Defilm cost [PLN]",
        "Operational cost [PLN]", "Technology cost [PLN]",
        "Bending [PLN]", "Additional costs [PLN]",
        "Unit cost [PLN]", "Total cost [PLN]"
    ]
    
    for col, header in enumerate(headers, 1):
        cell = detail_ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    # Calculate overhead distribution
    if total_parts_qty > 0:
        extra_per_part = (tech_per_order + add_costs_order) / total_parts_qty
        op_cost_per_part = (total_sheets * op_cost_per_sheet) / total_parts_qty
    else:
        extra_per_part = 0.0
        op_cost_per_part = 0.0
    
    # Data for corrected charts
    cost_components = {
        'Material': 0.0,
        'Laser cutting': 0.0,
        'Contours': 0.0,
        'Marking': 0.0,
        'Defilm': 0.0,
        'Operational costs': 0.0,
        'Technology': 0.0,
        'Bending': 0.0,
        'Additional costs': 0.0
    }
    
    row_num = 2
    for part in all_parts:
        # Calculate individual cost components
        mat_cost = part['adj_weight'] * part.get('base_price_per_kg', 0.0)
        cut_cost = part.get('cut_length', 0.0) * part.get('base_rate_per_cut_length', 0.0)
        contour_cost = part.get('contours_qty', 0.0) * part.get('rate_per_contour', 0.0)
        marking_cost = part.get('marking_length', 0.0) * part.get('rate_per_marking_length', 0.0)
        defilm_cost = part.get('defilm_length', 0.0) * part.get('rate_per_defilm_length', 0.0)
        bending_cost = part.get('bending_per_unit', 0.0)
        
        # Accumulate for charts (multiply by quantity for total costs)
        cost_components['Material'] += mat_cost * part['qty']
        cost_components['Laser cutting'] += cut_cost * part['qty']
        cost_components['Contours'] += contour_cost * part['qty']
        cost_components['Marking'] += marking_cost * part['qty']
        cost_components['Defilm'] += defilm_cost * part['qty']
        cost_components['Operational costs'] += op_cost_per_part * part['qty']
        cost_components['Technology'] += extra_per_part * part['qty']
        cost_components['Bending'] += bending_cost * part['qty']
        cost_components['Additional costs'] += part.get('additional_per_unit', 0.0) * part['qty']
        
        unit_cost = (mat_cost + cut_cost + contour_cost + marking_cost + defilm_cost + 
                    op_cost_per_part + extra_per_part + bending_cost + part.get('additional_per_unit', 0.0))
        total_part_cost = unit_cost * part['qty']
        
        # Write row data with proper column order (Miniatura as column 2)
        row_data = [
            part['id'],
            '',  # Placeholder for thumbnail
            part['name'],
            part['material'],
            part['thickness'],
            part['qty'],
            f"{part.get('raw_weight', 0.0):.3f}",
            f"{part.get('adj_weight', 0.0):.3f}",
            f"{part.get('cut_length', 0.0):.2f}",
            part.get('contours_qty', 0),
            f"{part.get('marking_length', 0.0):.2f}",
            f"{part.get('defilm_length', 0.0):.2f}",
            f"{part.get('base_price_per_kg', 0.0):.2f}",
            f"{part.get('base_rate_per_cut_length', 0.0):.2f}",
            f"{mat_cost:.2f}",
            f"{cut_cost:.2f}",
            f"{contour_cost:.2f}",
            f"{marking_cost:.2f}",
            f"{defilm_cost:.2f}",
            f"{op_cost_per_part:.2f}",
            f"{extra_per_part:.2f}",
            f"{bending_cost:.2f}",
            f"{part.get('additional_per_unit', 0.0):.2f}",
            f"{unit_cost:.2f}",
            f"{total_part_cost:.2f}"
        ]
        
        for col, value in enumerate(row_data, 1):
            cell = detail_ws.cell(row=row_num, column=col, value=value)
            if col >= 4:  # Numeric columns
                cell.alignment = Alignment(horizontal="right")
        
        # Add thumbnail in column 2 (B)
        if part.get('thumb_data'):
            try:
                img = OpenpyxlImage(io.BytesIO(part['thumb_data']))
                img.width = 60
                img.height = 40
                detail_ws.add_image(img, f'B{row_num}')
                detail_ws.row_dimensions[row_num].height = 45
            except Exception:
                pass
        
        row_num += 1
    
    # Add totals row
    total_row = row_num
    detail_ws.cell(row=total_row, column=3, value="TOTAL SUM").font = Font(bold=True)
    detail_ws.cell(row=total_row, column=25, value=f"{sum(cost_components.values()):.2f}").font = Font(bold=True)
    
    # Autofit columns
    for column in detail_ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = min(max_length + 2, 30)
        detail_ws.column_dimensions[column_letter].width = adjusted_width
    
    # Set specific width for thumbnail column
    detail_ws.column_dimensions['B'].width = 12
    
    # Sheet 2: Charts and Analysis with corrected data
    chart_ws = cost_wb.create_sheet("Charts and Analysis")

    # Title
    chart_ws['A1'] = "FINANCIAL ANALYSIS OF THE ORDER"
    chart_ws['A1'].font = Font(bold=True, size=16)
    chart_ws.merge_cells('A1:D1')
    
    # Cost breakdown table
    chart_ws['A3'] = "Cost component"
    chart_ws['B3'] = "Value [PLN]"
    chart_ws['C3'] = "Share [%]"
    
    for cell in ['A3', 'B3', 'C3']:
        chart_ws[cell].font = Font(bold=True)
        chart_ws[cell].fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    
    # Filter out zero-value components and calculate percentages
    total_costs = sum(cost_components.values())
    active_components = [(k, v) for k, v in cost_components.items() if v > 0]
    
    row = 4
    for name, value in active_components:
        chart_ws.cell(row=row, column=1, value=name)
        chart_ws.cell(row=row, column=2, value=round(value, 2))
        percentage = (value / total_costs * 100) if total_costs > 0 else 0
        chart_ws.cell(row=row, column=3, value=round(percentage, 1))
        row += 1
    
    # Total row
    chart_ws.cell(row=row, column=1, value="TOTAL").font = Font(bold=True)
    chart_ws.cell(row=row, column=2, value=round(total_costs, 2)).font = Font(bold=True)
    chart_ws.cell(row=row, column=3, value=100.0).font = Font(bold=True)
    

    # Autofit columns for "Charts and Analysis" sheet
    for column_cells in chart_ws.columns:
        max_length = 0
        column_letter = None
    
        for cell in column_cells:
            # Skip merged cells and get column letter from first regular cell
            if hasattr(cell, 'column_letter'):
                if column_letter is None:
                    column_letter = cell.column_letter
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
    
        # Only set width if we found a valid column letter
        if column_letter:
            adjusted_width = min(max_length + 2, 40)  # Cap at 40 to avoid overly wide columns
            chart_ws.column_dimensions[column_letter].width = adjusted_width

    # Set minimum widths for specific columns
    chart_ws.column_dimensions['A'].width = max(chart_ws.column_dimensions['A'].width, 20)  # Component names
    chart_ws.column_dimensions['B'].width = max(chart_ws.column_dimensions['B'].width, 15)  # Values
    chart_ws.column_dimensions['C'].width = max(chart_ws.column_dimensions['C'].width, 12)  # Percentages

    # Pie Chart - Corrected with proper data reference
    if len(active_components) > 0:
        pie = PieChart()
        pie.title = "Cost structure (%)"
        pie.width = 20
        pie.height = 15
    
        # Data range (values)
        data = Reference(chart_ws, min_col=3, min_row=4, max_row=3+len(active_components))
        # Categories (labels)
        categories = Reference(chart_ws, min_col=1, min_row=4, max_row=3+len(active_components))
    
        pie.add_data(data, titles_from_data=False)
        pie.set_categories(categories)
    
        # Data labels showing percentages
        pie.dataLabels = DataLabelList()
        pie.dataLabels.showPercent = True
        pie.dataLabels.showCatName = True
        pie.dataLabels.showSerName = False  # This removes "Serie1" from labels
        pie.dataLabels.showVal = False  # Don't show raw values
        chart_ws.add_chart(pie, "E3")
        
    
    # Financial Result Summary
    chart_ws['A' + str(row + 3)] = "FINANCIAL RESULT"
    chart_ws['A' + str(row + 3)].font = Font(bold=True, size=12)
    
    client_price = total_price_per_order 
    margin = client_price - total_costs
    margin_percent = (margin / total_costs * 100) if total_costs > 0 else 0
    
    chart_ws['A' + str(row + 5)] = "Total costs:"
    chart_ws['B' + str(row + 5)] = f"{total_costs:.2f} PLN"
    chart_ws['A' + str(row + 6)] = "Price for client:"
    chart_ws['B' + str(row + 6)] = f"{client_price:.2f} PLN"
    chart_ws['A' + str(row + 7)] = "Margin:"
    chart_ws['B' + str(row + 7)] = f"{margin:.2f} PLN ({margin_percent:.1f}%)"
    chart_ws['B' + str(row + 7)].font = Font(bold=True, color="008000")
    
    # Save the cost report
    cost_wb.save(os.path.join(raporty_path, "Cost report.xlsx"))
    
    # Generate enhanced client report with professional styling
    client_wb = Workbook()
    client_ws = client_wb.active
    client_ws.title = "Offer for client"
    
    # Add header with company info and logo space
    client_ws.merge_cells('A1:I1')
    client_ws['A1'] = "LP KONSTAL Sp. z o.o."
    client_ws['A1'].font = Font(bold=True, size=16, color="366092")
    client_ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
    
    # Add offer details
    client_ws.merge_cells('A3:D3')
    client_ws['A3'] = f"OFFER NO: {offer_number}"
    client_ws['A3'].font = Font(bold=True, size=14)
    
    client_ws.merge_cells('F3:I3')
    client_ws['F3'] = f"Date: {datetime.datetime.now().strftime('%d.%m.%Y')}"
    client_ws['F3'].alignment = Alignment(horizontal="right")
    
    client_ws.merge_cells('A4:D4')
    client_ws['A4'] = f"For: {customer_name}"
    client_ws['A4'].font = Font(size=12)
    
    client_ws.merge_cells('F4:I4')
    client_ws['F4'] = f"Valid until: {validity}"
    client_ws['F4'].alignment = Alignment(horizontal="right")
    
    # Add a separator row
    client_ws.merge_cells('A6:I6')
    client_ws['A6'].fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    client_ws.row_dimensions[6].height = 3
    
    # Column headers for client report
    headers = [
        "ID", "Miniatura", "Part name", "Material", 
        "Thickness [mm]", "Unit weight [kg]", 
        "Quantity [pcs]", "Unit cost [PLN]", "Total cost [PLN]"
    ]
    
    header_row = 8
    for col, header in enumerate(headers, 1):
        cell = client_ws.cell(row=header_row, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )  # <- This closing parenthesis was missing
    
    # Add data rows with alternating colors
    data_start_row = header_row + 1
    client_total = 0.0
    
    for idx, part in enumerate(all_parts):
        row_num = data_start_row + idx
        
        # Alternate row colors for better readability
        fill_color = "F2F2F2" if idx % 2 == 0 else "FFFFFF"
        
        unit_total = part['cost_per_unit'] + part['bending_per_unit'] + part['additional_per_unit']
        total_part = unit_total * part['qty']
        client_total += total_part
        
        # ID
        cell = client_ws.cell(row=row_num, column=1, value=part['id'])
        cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
        cell.border = Border(left=Side(style='thin'), right=Side(style='thin'))
        
        # Miniatura (column 2)
        cell = client_ws.cell(row=row_num, column=2, value='')
        cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
        if part.get('thumb_data'):
            try:
                img = OpenpyxlImage(io.BytesIO(part['thumb_data']))
                img.width = 50
                img.height = 35
                client_ws.add_image(img, f'B{row_num}')
                client_ws.row_dimensions[row_num].height = 40
            except:
                pass
        
        # Other columns
        values = [
            part['name'],
            part['material'],
            part['thickness'],
            f"{part.get('raw_weight', 0.0):.3f}",
            part['qty'],
            f"{unit_total:.2f}",
            f"{total_part:.2f}"
        ]
        
        for col, value in enumerate(values, 3):
            cell = client_ws.cell(row=row_num, column=col, value=value)
            cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
            if col in [5, 6, 8, 9]:  # Numeric columns
                cell.alignment = Alignment(horizontal="right")
            cell.border = Border(left=Side(style='thin'), right=Side(style='thin'))
    
    # Total row
    total_row = data_start_row + len(all_parts)
    client_ws.merge_cells(f'A{total_row}:F{total_row}')
    cell = client_ws.cell(row=total_row, column=1, value="TOTAL")
    cell.font = Font(bold=True, size=12)
    cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    cell.alignment = Alignment(horizontal="right")
    
    for col in range(7, 9):
        cell = client_ws.cell(row=total_row, column=col, value="")
        cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    
    cell = client_ws.cell(row=total_row, column=9, value=f"{client_total:.2f}")
    cell.font = Font(bold=True, size=12)
    cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    cell.alignment = Alignment(horizontal="right")
    cell.border = Border(top=Side(style='double'), bottom=Side(style='double'))
    
    # Add closing text with disclaimers
    disclaimer_start = total_row + 3
    client_ws.merge_cells(f'A{disclaimer_start}:I{disclaimer_start}')
    client_ws[f'A{disclaimer_start}'] = "IMPLEMENTATION CONDITIONS"
    client_ws[f'A{disclaimer_start}'].font = Font(bold=True, size=11)
    
    disclaimers = [
        "1. Order implementation follows offer acceptance and submission of technical documentation.",
        "2. Delivery time: to be agreed, standard 5-7 working days.",
        "3. Prices do not include transportation costs.",
        "4. Liability exclusions:",
        "   • For errors in the provided technical documentation",
        "   • For defects in the material provided",
        "   • For damages resulting from force majeure",
        "5. Payment: transfer 14 days from the invoice VAT date."
    ]
    
    for idx, text in enumerate(disclaimers):
        row = disclaimer_start + idx + 1
        client_ws.merge_cells(f'A{row}:I{row}')
        client_ws[f'A{row}'] = text
        client_ws[f'A{row}'].font = Font(size=9)
        client_ws[f'A{row}'].alignment = Alignment(wrap_text=True)
    
    # Footer with contact info
    footer_row = disclaimer_start + len(disclaimers) + 3
    client_ws.merge_cells(f'A{footer_row}:I{footer_row}')
    client_ws[f'A{footer_row}'] = "Laser Team | Tel: +48 537 883 393 | Email: laser@konstal.com"
    client_ws[f'A{footer_row}'].font = Font(size=10, italic=True)
    client_ws[f'A{footer_row}'].alignment = Alignment(horizontal="center")
    
    # Set column widths
    column_widths = {
        'A': 8,   # ID
        'B': 12,  # Miniatura
        'C': 35,  # Part name
        'D': 15,  # Material
        'E': 12,  # Thickness
        'F': 18,  # Weight
        'G': 10,  # Quantity
        'H': 18,  # Unit cost
        'I': 18   # Total cost
    }
    
    for col, width in column_widths.items():
        client_ws.column_dimensions[col].width = width
    
    # Add print settings
    client_ws.page_setup.orientation = 'landscape'
    client_ws.page_setup.fitToWidth = 1
    client_ws.page_setup.fitToHeight = 0
    
    # Save the client report
    client_wb.save(os.path.join(raporty_path, "Client report.xlsx"))
    
    # Generate DOCX (keep existing code for Word document)
    # ... [existing DOCX generation code] ...
    
    messagebox.showinfo("Success", "Reports generated in the Raporty folder.")
    


# left buttons
ttk.Button(buttons_frame, text="Analyze XLSX", command=analyze_xlsx_folder).pack(side="left", padx=5)
ttk.Button(buttons_frame, text="Generate report", command=generate_report).pack(side="left")

# ---- sash setup ----
def set_sash_positions(attempt=1):
    try:
        root.update_idletasks()
        panes = panel_a.panes()
        sash_count = max(len(panes) - 1, 0)
        if sash_count == 0:
            root.after(60, lambda: set_sash_positions(attempt+1)); return

        h = panel_a.winfo_height()
        if h < 400 and attempt < 10:
            # window is still expanding — try later
            root.after(80, lambda: set_sash_positions(attempt+1)); return

        # layout: Panel1 ~ 50% height, Panel2 ~ 20%, Panel3 the rest (minsize protects against 0px)
        y1 = max(220, int(h * 0.50))
        try: panel_a.sash_place(0, 0, y1)
        except Exception: pass

        if sash_count >= 2:
            y2 = min(h - 200, y1 + panel2_height)
            try: panel_a.sash_place(1, 0, y2)
            except Exception: pass
    except tk.TclError:
        if attempt < 10:
            root.after(80, lambda: set_sash_positions(attempt+1))

root.after_idle(set_sash_positions)

# run
root.geometry("1280x800")
root.mainloop()