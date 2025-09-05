#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
wycena.py - Main GUI script for analyzing XLSX files and generating cost reports.
CORRECTED VERSION with Fixed Excel Generation and Polish Number Format

Key corrections:
- Fixed Excel file corruption issues
- Polish number notation (comma as decimal separator)
- Using Excel formulas instead of static values
- Enhanced logging with detailed cost breakdowns
- Proper image handling in Excel files
"""

import os
import sys
import datetime
from datetime import timedelta
import re
from tkinter import ttk
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from openpyxl import load_workbook
from openpyxl import Workbook
from openpyxl.drawing.image import Image as OpenpyxlImage
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import PieChart, BarChart, Reference
from openpyxl.chart.label import DataLabelList
from openpyxl.utils.units import pixels_to_EMU
from openpyxl.drawing.spreadsheet_drawing import AnchorMarker, OneCellAnchor
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, RGBColor, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from PIL import Image, ImageTk
import locale
import requests
from openpyxl.drawing.spreadsheet_drawing import AnchorMarker, TwoCellAnchor
from openpyxl.utils import get_column_letter
from openpyxl.utils.units import pixels_to_EMU
import base64, json



# ========= SAVE / LOAD PROJECT (one-file JSON with base64 images) =========
def _b64_encode(b: bytes) -> str:
    return base64.b64encode(b).decode("ascii") if b is not None else ""

def _b64_decode(s: str) -> bytes:
    try:
        return base64.b64decode(s.encode("ascii")) if s else b""
    except Exception:
        return b""

def save_project_ui():
    """Ask for file and save current state (parts, calculations, margins, texts, pictures)."""

    folder_path = folder_var.get()
    if not folder_path:
        analysis_logger.log("No folder selected", "ERROR")
        messagebox.showerror("Error", "Please select a folder.")
        return

    current_date = datetime.datetime.now().strftime("%Y%m%d")

    customer_name = customer_var.get().strip() or "Client"
    offer_number = offer_var.get().strip()
    if not offer_number:
        offer_number = get_next_offer_number()
        offer_var.set(offer_number)

    fname = f"Oferta_{sanitize_filename(customer_name) or 'Klient'}_{current_date}_{offer_number.replace('/', '-')}.lpf"
    full_name = os.path.join(folder_path, fname)

    path = filedialog.asksaveasfilename(
        title="Save project",
        defaultextension=".lpf",
        filetypes=[("Laser Project File", "*.lpf"), ("JSON", "*.json"), ("All files", "*.*")],
        initialdir=folder_path,   # domyślny katalog
        initialfile=fname         # domyślna nazwa pliku
    )

    if not path:
        return

    try:
        # collect parts from tree + all_parts (thumbs)
        parts_payload = []
        tree_items = list(tree.get_children())
        for i, iid in enumerate(tree_items):
            # pomijamy wiersz TOTAL jeśli jest
            if iid == total_row_iid:
                continue
            vals = list(tree.item(iid, "values"))
            # Safeguard dla indeksów
            vals += [""] * (11 - len(vals))
            # domyślne dane z 'all_parts'
            thumb_b64 = ""
            if i < len(all_parts):
                thumb_b64 = _b64_encode(all_parts[i].get("thumb_data") or b"")
            parts_payload.append({
                "values": vals,                # kolumny TreeView
                "thumb_b64": thumb_b64,        # miniatura (base64)
                "cost_per_unit": all_parts[i].get("cost_per_unit") if i < len(all_parts) else None,
                "qty": all_parts[i].get("qty") if i < len(all_parts) else None,
                "bending_per_unit": all_parts[i].get("bending_per_unit") if i < len(all_parts) else None,
                "additional_per_unit": all_parts[i].get("additional_per_unit") if i < len(all_parts) else None,
            })

        payload = {
            "meta": {
                "saved_at": datetime.datetime.now().isoformat(),
                "app": "wycena.py",
            },
            "header": {
                "folder": folder_var.get(),
                "customer": customer_var.get(),
                "offer": offer_var.get(),
                "date": date_var.get(),
                "validity": validity_var.get(),
                "logo": logo_var.get(),
            },
            "texts": {
                "contact": contact_text.get("1.0", "end-1c"),
                "preceding": preceding_text_var.get("1.0", "end-1c"),
                "finishing": finishing_text_var.get("1.0", "end-1c"),
            },
            "fixed_costs": {
                "op_cost_per_sheet": op_cost_entry.get(),
                "tech_order": tech_order_entry.get(),
                "add_order": add_order_cost_entry.get(),
            },
            "rates": {
                "O_rate": oxygen_rate_entry.get(),
                "N_rate": nitrogen_rate_entry.get(),
                "ALN_rate": al_nitrogen_rate_entry.get(),
                "O_rate_tkw": oxygen_rate_entry_TKW.get(),
                "N_rate_tkw": nitrogen_rate_entry_TKW.get(),
                "ALN_rate_tkw": al_nitrogen_rate_entry_TKW.get(),
                "bend_percent_tkw": bending_percent_entry_TKW.get(),
            },
            "margins": {
                "material": material_margin_var.get(),
                "cutting": cutting_margin_var.get(),
                "min_area": min_area_var.get(),
                "max_area": max_area_var.get(),
                "min_cut_len": min_cutting_var.get(),
                "max_cut_len": max_cutting_var.get(),
            },
            "calculated": {
                "oxygen_cutting_time": oxygen_cutting_time,
                "nitrogen_cutting_time": nitrogen_cutting_time,
                "aluminum_nitrogen_cutting_time": aluminum_nitrogen_cutting_time,
                "total_material_cost": total_material_cost,
                "total_price_per_order": float(total_price_per_order) if total_price_per_order else 0.0,
                "labels": {
                    "oxygen_time": oxygen_time_label.cget("text"),
                    "nitrogen_time": nitrogen_time_label.cget("text"),
                    "oxygen_cost": oxygen_cost_label.cget("text"),
                    "nitrogen_cost": nitrogen_cost_label.cget("text"),
                    "material_cost": material_cost_label.cget("text"),
                    "total_cutting_cost": total_cutting_cost_label.cget("text"),
                    "operational_cost": operational_cost_label.cget("text"),
                    "total_all_costs": total_all_costs_label.cget("text"),
                    "total_for_correction": total_all_costs_entry.get(),
                }
            },
            "parts": parts_payload
        }

        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)

        analysis_logger.log(f"Project saved: {os.path.basename(path)}", "SUCCESS")
        messagebox.showinfo("Saved", f"Project saved to:\n{path}")

    except Exception as e:
        analysis_logger.log(f"Save failed: {e}", "ERROR")
        messagebox.showerror("Error", f"Save failed:\n{e}")

def load_project_ui():
    """Open project file and restore full UI (parts, images, fields, labels)."""
    path = filedialog.askopenfilename(
        title="Open project",
        filetypes=[("Laser Project File", "*.lpf *.json"), ("All files", "*.*")]
    )
    if not path:
        return

    try:
        with open(path, "r", encoding="utf-8") as f:
            payload = json.load(f)
    except Exception as e:
        messagebox.showerror("Error", f"Cannot open file:\n{e}")
        return

    try:
        # header
        folder_var.set(payload["header"].get("folder", ""))
        customer_var.set(payload["header"].get("customer", ""))
        offer_var.set(payload["header"].get("offer", ""))
        date_var.set(payload["header"].get("date", ""))
        validity_var.set(payload["header"].get("validity", ""))
        logo_var.set(payload["header"].get("logo", ""))

        # texts
        contact_text.delete("1.0", "end"); contact_text.insert("1.0", payload["texts"].get("contact", ""))
        preceding_text_var.delete("1.0", "end"); preceding_text_var.insert("1.0", payload["texts"].get("preceding", ""))
        finishing_text_var.delete("1.0", "end"); finishing_text_var.insert("1.0", payload["texts"].get("finishing", ""))

        # fixed costs
        op_cost_entry.delete(0, "end"); op_cost_entry.insert(0, payload["fixed_costs"].get("op_cost_per_sheet", "0,00"))
        tech_order_entry.delete(0, "end"); tech_order_entry.insert(0, payload["fixed_costs"].get("tech_order", "0,00"))
        add_order_cost_entry.delete(0, "end"); add_order_cost_entry.insert(0, payload["fixed_costs"].get("add_order", "0,00"))

        # rates (left + right column)
        oxygen_rate_entry.delete(0,"end"); oxygen_rate_entry.insert(0, payload["rates"].get("O_rate", "350,00"))
        nitrogen_rate_entry.delete(0,"end"); nitrogen_rate_entry.insert(0, payload["rates"].get("N_rate", "550,00"))
        al_nitrogen_rate_entry.delete(0,"end"); al_nitrogen_rate_entry.insert(0, payload["rates"].get("ALN_rate", "650,00"))
        oxygen_rate_entry_TKW.delete(0,"end"); oxygen_rate_entry_TKW.insert(0, payload["rates"].get("O_rate_tkw", "262,50"))
        nitrogen_rate_entry_TKW.delete(0,"end"); nitrogen_rate_entry_TKW.insert(0, payload["rates"].get("N_rate_tkw", "412,50"))
        al_nitrogen_rate_entry_TKW.delete(0,"end"); al_nitrogen_rate_entry_TKW.insert(0, payload["rates"].get("ALN_rate_tkw", "487,50"))
        bending_percent_entry_TKW.delete(0,"end"); bending_percent_entry_TKW.insert(0, payload["rates"].get("bend_percent_tkw", "75,00"))

        # margins + ranges
        material_margin_var.set(payload["margins"].get("material", "0,00"))
        cutting_margin_var.set(payload["margins"].get("cutting", "0,00"))
        min_area_var.set(payload["margins"].get("min_area", "0,00"))
        max_area_var.set(payload["margins"].get("max_area", "1,00"))
        min_cutting_var.set(payload["margins"].get("min_cut_len", "0,00"))
        max_cutting_var.set(payload["margins"].get("max_cut_len", "5000,00"))

        # calculated (restore labels, times, totals)
        global oxygen_cutting_time, nitrogen_cutting_time, aluminum_nitrogen_cutting_time, total_material_cost
        global total_row_iid, all_parts, total_price_per_order

        calc = payload.get("calculated", {})
        oxygen_cutting_time = float(calc.get("oxygen_cutting_time", 0.0))
        nitrogen_cutting_time = float(calc.get("nitrogen_cutting_time", 0.0))
        aluminum_nitrogen_cutting_time = float(calc.get("aluminum_nitrogen_cutting_time", 0.0))
        total_material_cost = float(calc.get("total_material_cost", 0.0))
        total_price_per_order = float(calc.get("total_price_per_order", 0.0))

        # Przywracamy napisy (opcjonalnie — i tak przeliczymy dalej)
        labels = calc.get("labels", {})
        oxygen_time_label.config(text=labels.get("oxygen_time", "0,00"))
        nitrogen_time_label.config(text=labels.get("nitrogen_time", "0,00"))
        oxygen_cost_label.config(text=labels.get("oxygen_cost", "0,00"))
        nitrogen_cost_label.config(text=labels.get("nitrogen_cost", "0,00"))
        material_cost_label.config(text=labels.get("material_cost", "0,00"))
        total_cutting_cost_label.config(text=labels.get("total_cutting_cost", "0,00"))
        operational_cost_label.config(text=labels.get("operational_cost", "0,00"))
        total_all_costs_label.config(text=labels.get("total_all_costs", "0,00"))
        total_all_costs_entry.delete(0,"end"); total_all_costs_entry.insert(0, labels.get("total_for_correction", "0,00"))

        # Rebuild tree + all_parts + thumbnails
        for item in tree.get_children():
            tree.delete(item)
        thumbnail_imgs.clear()
        all_parts = []

        for p in payload.get("parts", []):
            vals = p.get("values", [""]*11)
            # wstaw do tree
            iid = tree.insert('', 'end', values=vals)
            # odtwórz miniaturę
            b = _b64_decode(p.get("thumb_b64", ""))
            if b:
                try:
                    img = ImageTk.PhotoImage(Image.open(io.BytesIO(b)).resize((80, 80)))
                    tree.item(iid, image=img)
                    thumbnail_imgs.append(img)  # prevent GC
                except Exception:
                    pass
            # odtwórz all_parts (tyle ile potrzebujemy do przeliczeń i zapisu)
            all_parts.append({
                "thumb_data": b if b else None,
                "cost_per_unit": p.get("cost_per_unit"),
                "qty": p.get("qty"),
                "bending_per_unit": p.get("bending_per_unit"),
                "additional_per_unit": p.get("additional_per_unit"),
                # można dodać inne pola według potrzeb analizy
            })

        # Dodaj / przelicz wiersz sumy
        total_row_iid = None
        update_total()
        update_cost_calculations()

        analysis_logger.log(f"Project loaded: {os.path.basename(path)}", "SUCCESS")
        messagebox.showinfo("Loaded", f"Project loaded:\n{path}")

    except Exception as e:
        analysis_logger.log(f"Load failed: {e}", "ERROR")
        messagebox.showerror("Error", f"Load failed:\n{e}")


def SetTotalPricePerOrder(value):
    global total_price_per_order
    total_price_per_order = value   

try:
    locale.setlocale(locale.LC_ALL, 'pl_PL.UTF-8')
except Exception:
    pass

def format_pln(value):
    """Formats the value to PLN with a comma and thousands grouping."""
    try:
        s = locale.format_string('%.2f', float(value), grouping=True)
        return s.replace('.', ',')
    except Exception:
        try:
            return f"{float(value):.2f}".replace('.', ',')
        except Exception:
            return "0,00"

def format_excel_number(value):
    """Format number for Excel with Polish notation (comma as decimal separator)"""
    try:
        return f"{float(value):.2f}".replace('.', ',')
    except:
        return "0,00"

def sanitize_filename(name):
    """Sanitizes the file name by replacing disallowed characters."""
    for ch in r'< > : " / \ | ? *':
        name = name.replace(ch, '_')
    return name

def _norm_s(s):
    """Normalizes the string to uppercase, removes spaces."""
    return (str(s).strip().upper() if s is not None else "")

def _parse_float(val):
    """Parses the value to float, handles commas."""
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip().replace(" ", "").replace("\xa0", "").replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None

def _map_gas_to_key(gas_raw: str) -> str:
    """Maps the gas name to the key 'N' or 'O'."""
    g = _norm_s(gas_raw)
    if g in {"NITROGEN", "AZOT", "氮气", "N", "高压氮气"}:
        return "N"
    if g in {"OXYGEN", "TLEN", "氧气", "O"}:
        return "O"
    return ""

# ---- Analysis Logger Class ----
class AnalysisLogger:
    def __init__(self, text_widget):
        self.text_widget = text_widget
        self.phase_counter = 0
        
    def clear(self):
        """Clear the log widget"""
        self.text_widget.config(state=tk.NORMAL)
        self.text_widget.delete(1.0, tk.END)
        self.text_widget.config(state=tk.DISABLED)
        self.phase_counter = 0
        
    def log(self, message, level="INFO"):
        """Log a message with specified level (INFO, WARNING, ERROR, SUCCESS, PHASE)"""
        self.text_widget.config(state=tk.NORMAL)
        
        timestamp = datetime.datetime.now().strftime("%H:%M:%S")
        
        # Format and add message based on level
        if level == "PHASE":
            self.phase_counter += 1
            formatted_message = f"\n[{timestamp}] ===== PHASE {self.phase_counter}: {message} =====\n"
            self.text_widget.insert(tk.END, formatted_message, "phase")
        elif level == "ERROR":
            formatted_message = f"[{timestamp}] ❌ ERROR: {message}\n"
            self.text_widget.insert(tk.END, formatted_message, "error")
        elif level == "WARNING":
            formatted_message = f"[{timestamp}] ⚠️ WARNING: {message}\n"
            self.text_widget.insert(tk.END, formatted_message, "warning")
        elif level == "SUCCESS":
            formatted_message = f"[{timestamp}] ✅ SUCCESS: {message}\n"
            self.text_widget.insert(tk.END, formatted_message, "success")
        else:  # INFO
            formatted_message = f"[{timestamp}] ℹ️ {message}\n"
            self.text_widget.insert(tk.END, formatted_message, "info")
        
        # Auto-scroll to bottom
        self.text_widget.see(tk.END)
        self.text_widget.config(state=tk.DISABLED)
        
        # Force update to show log immediately
        self.text_widget.update_idletasks()

# ---- price lists ----
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MATERIALS_FILE = os.path.join(SCRIPT_DIR, "materials prices.xlsx")
CUTTING_FILE   = os.path.join(SCRIPT_DIR, "cutting prices.xlsx")
material_prices = {}  # (MAT, THK)-> PLN/kg
cutting_prices  = {}  # (THK, MAT, GAS)-> PLN/m
_mat_set, _thk_set, _gas_set = set(), set(), set()

# Global variables for shared data
all_parts = []
last_groups = []
last_total_cost = 0.0
last_folder_path = ""
total_sheets = 0
total_parts_qty = 0
total_row_iid = None
total_price_per_order = 0.0
analysis_logger = None  # Will be initialized after GUI creation

# Global variables for margin calculations
file_margins = []  # List of margin data for each file
avg_material_margin = 0.0
avg_cutting_margin = 0.0

# Global variables for cutting time calculations
oxygen_cutting_time = 0.0
nitrogen_cutting_time = 0.0
aluminum_nitrogen_cutting_time = 0.0
total_material_cost = 0.0

# ---- GUI ----
root = tk.Tk()
root.title("Cost Report Generator – Enhanced with Dynamic Margins and TKW Analysis")
root.configure(bg="#2c2c2c")  # Dark background for modern look

# Use a modern ttk theme
style = ttk.Style(root)
style.theme_use('clam')  # Or 'alt', 'default', etc. for better visuals
style.configure("TLabel", foreground="white", background="#2c2c2c", font=("Arial", 10))
style.configure("TEntry", fieldbackground="#3c3c3c", foreground="white")
style.configure("TButton", background="#4c4c4c", foreground="white", borderwidth=0)
style.configure("TCombobox", fieldbackground="#3c3c3c", foreground="white")
style.configure("Treeview", background="#3c3c3c", foreground="white", fieldbackground="#3c3c3c", rowheight=80)  # Increased row height for better thumbnails
style.configure("Treeview.Heading", background="#4c4c4c", foreground="white")
style.map("TButton", background=[('active', '#5c5c5c')])
style.map("Treeview", background=[('selected', '#5c5c5c')])

folder_var   = tk.StringVar()
customer_var = tk.StringVar()
offer_var    = tk.StringVar()
date_var     = tk.StringVar(value=datetime.datetime.now().strftime("%Y-%m-%d"))
validity_var = tk.StringVar(value=(datetime.datetime.now() + timedelta(days=14)).strftime("%Y-%m-%d"))
logo_var     = tk.StringVar()

# Margin calculation variables - will be updated with calculated averages
material_margin_var = tk.StringVar(value="0,00")  # Will be auto-populated with calculated average
cutting_margin_var = tk.StringVar(value="0,00")   # Will be auto-populated with calculated average
min_area_var = tk.StringVar(value="0,00")          # Minimum area for 250% margin (m²)
max_area_var = tk.StringVar(value="1,00")          # Maximum area for 0% margin (m²)
min_cutting_var = tk.StringVar(value="0,00")       # Minimum cutting length for 200% margin (mm)
max_cutting_var = tk.StringVar(value="5000,00")    # Maximum cutting length for 0% margin (mm)

default_logo_path = os.path.join(SCRIPT_DIR, "Logo.jpg")
if os.path.exists(default_logo_path):
    logo_var.set(default_logo_path)

# LEFT
left_frame = tk.Frame(root, bg="#2c2c2c")
left_frame.pack(side="left", padx=10, pady=10, fill="y")

# Add Analysis Log Panel at the top of left frame
log_frame = tk.LabelFrame(left_frame, text="ANALYSIS LOG", bg="#2c2c2c", fg="white", padx=5, pady=5)
log_frame.grid(row=0, column=0, columnspan=3, sticky="ew", pady=(0, 10))

# Create scrolled text widget for log
log_text = scrolledtext.ScrolledText(
    log_frame, 
    height=12, 
    width=60, 
    bg="#1c1c1c", 
    fg="white", 
    insertbackground="white",
    wrap=tk.WORD,
    state=tk.DISABLED
)
log_text.pack(fill="both", expand=True)

# Configure text tags for different log levels
log_text.tag_configure("phase", foreground="#00BFFF", font=("Arial", 10, "bold"))
log_text.tag_configure("error", foreground="#FF4444", font=("Arial", 10, "bold"))
log_text.tag_configure("warning", foreground="#FFA500", font=("Arial", 10))
log_text.tag_configure("success", foreground="#00FF00", font=("Arial", 10))
log_text.tag_configure("info", foreground="#FFFFFF", font=("Arial", 10))

# Initialize the logger
analysis_logger = AnalysisLogger(log_text)

# Clear log button
ttk.Button(log_frame, text="Clear Log", command=analysis_logger.clear).pack(pady=(5, 0))

def update_file_list(folder_path):
    file_list.delete(0, tk.END)
    try:
        xlsx_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".xlsx")]
        for f in xlsx_files:
            file_list.insert(tk.END, f)
        analysis_logger.log(f"Found {len(xlsx_files)} XLSX files in folder", "INFO")
    except Exception as e:
        analysis_logger.log(f"Failed to read folder: {str(e)}", "ERROR")

def select_folder():
    p = filedialog.askdirectory()
    if p:
        folder_var.set(p)
        update_file_list(p)
        analysis_logger.log(f"Selected folder: {p}", "SUCCESS")

# Shift all row numbers down by 1 to account for the log panel
ttk.Label(left_frame, text="Select folder:").grid(row=1, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=folder_var, width=50).grid(row=1, column=1)
ttk.Button(left_frame, text="Browse", command=select_folder).grid(row=1, column=2)

ttk.Label(left_frame, text="Client name:").grid(row=2, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=customer_var).grid(row=2, column=1)
ttk.Label(left_frame, text="Offer number:").grid(row=3, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=offer_var).grid(row=3, column=1)
ttk.Button(left_frame, text="Get number", command=lambda: offer_var.set(get_next_offer_number())).grid(row=3, column=2)
ttk.Label(left_frame, text="Offer date:").grid(row=4, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=date_var).grid(row=4, column=1)
ttk.Label(left_frame, text="Validity period:").grid(row=5, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=validity_var).grid(row=5, column=1)

def upload_logo():
    p = filedialog.askopenfilename(filetypes=[("Image files", "*.png;*.jpg;*.jpeg")])
    if p: 
        logo_var.set(p)
        analysis_logger.log(f"Logo selected: {os.path.basename(p)}", "INFO")

ttk.Label(left_frame, text="Load logo:").grid(row=6, column=0, sticky="e")
ttk.Entry(left_frame, textvariable=logo_var, width=50).grid(row=6, column=1)
ttk.Button(left_frame, text="Browse", command=upload_logo).grid(row=6, column=2)

ttk.Label(left_frame, text="Contact details:").grid(row=7, column=0, sticky="ne")
contact_text = tk.Text(left_frame, height=5, width=50, bg="#3c3c3c", fg="white", insertbackground="white")
contact_text.grid(row=7, column=1)
contact_text.insert(tk.INSERT,
    "LP KONSTAL Sp. z o.o.\n59-800 Lubań\nPisarzowice 203B\nNIP: 1010004508\n\n"
    "Laser Team\nMateusz Brzostek M. +48 537 883 393\n"
    "Artur Jednoróg M. +48 515 803 333\nE. laser@konstal.com"
)

# Add Margin Calculation Panel
margin_frame = tk.LabelFrame(left_frame, text="DYNAMIC MARGIN SETTINGS", bg="#2c2c2c", fg="yellow", padx=5, pady=5)
margin_frame.grid(row=8, column=0, columnspan=3, sticky="ew", pady=(10, 5))

# Material margin settings
ttk.Label(margin_frame, text="Proposed margin for material [%]:").grid(row=0, column=0, sticky="e")
ttk.Entry(margin_frame, textvariable=material_margin_var, width=15).grid(row=0, column=1, sticky="w", padx=(5,10))

ttk.Label(margin_frame, text="Proposed margin for cutting [%]:").grid(row=0, column=2, sticky="e")
ttk.Entry(margin_frame, textvariable=cutting_margin_var, width=15).grid(row=0, column=3, sticky="w", padx=(5,0))

# Limit values for material area
ttk.Label(margin_frame, text="Min area for 250% margin [m²]:").grid(row=1, column=0, sticky="e")
ttk.Entry(margin_frame, textvariable=min_area_var, width=15).grid(row=1, column=1, sticky="w", padx=(5,10))

ttk.Label(margin_frame, text="Max area for 0% margin [m²]:").grid(row=1, column=2, sticky="e")
ttk.Entry(margin_frame, textvariable=max_area_var, width=15).grid(row=1, column=3, sticky="w", padx=(5,0))

# Limit values for cutting length
ttk.Label(margin_frame, text="Min length for 200% margin [mm]:").grid(row=2, column=0, sticky="e")
ttk.Entry(margin_frame, textvariable=min_cutting_var, width=15).grid(row=2, column=1, sticky="w", padx=(5,10))

ttk.Label(margin_frame, text="Max length for 0% margin [mm]:").grid(row=2, column=2, sticky="e")
ttk.Entry(margin_frame, textvariable=max_cutting_var, width=15).grid(row=2, column=3, sticky="w", padx=(5,0))

# Display calculated averages
avg_display_frame = tk.Frame(margin_frame, bg="#2c2c2c")
avg_display_frame.grid(row=3, column=0, columnspan=4, pady=(10,5), sticky="ew")

ttk.Label(avg_display_frame, text="Calculated avg material margin:", font=("Arial", 9, "bold")).pack(side="left")
avg_material_label = ttk.Label(avg_display_frame, text="0.00%", foreground="lime", font=("Arial", 9, "bold"))
avg_material_label.pack(side="left", padx=(5,20))

ttk.Label(avg_display_frame, text="Calculated avg cutting margin:", font=("Arial", 9, "bold")).pack(side="left")
avg_cutting_label = ttk.Label(avg_display_frame, text="0.00%", foreground="lime", font=("Arial", 9, "bold"))
avg_cutting_label.pack(side="left", padx=(5,0))

ttk.Label(left_frame, text="Preceding text:").grid(row=9, column=0, sticky="ne")
preceding_text_var = tk.Text(left_frame, height=5, width=50, bg="#3c3c3c", fg="white", insertbackground="white")
preceding_text_var.grid(row=9, column=1)
preceding_text_var.insert(tk.INSERT,
    "Szanowni Państwo,\n\n"
    "dziękujemy za przesłanie zapytania ofertowego dotyczącego usługi cięcia laserem blach. "
    "Z przyjemnością przedstawiamy przygotowaną dla Państwa ofertę.")

ttk.Label(left_frame, text="").grid(row=10, column=0, pady=10)
ttk.Label(left_frame, text="Finishing text:").grid(row=11, column=0, sticky="ne")
finishing_text_var = tk.Text(left_frame, height=10, width=50, bg="#3c3c3c", fg="white", insertbackground="white")
finishing_text_var.grid(row=11, column=1)
finishing_text_var.insert(tk.INSERT, "Wyłączenia odpowiedzialności \r\nDokumentacja techniczna\r\nRealizacja zamówienia odbywa się wyłącznie na podstawie dokumentacji technicznej dostarczonej przez Klienta. Odpowiedzialność za jej kompletność, poprawność oraz zgodność z założeniami projektowymi leży wyłącznie po stronie Zleceniodawcy. Wszelkie błędy, niejasności, czy niezgodności w przesłanych plikach uniemożliwiające prawidłowe wykonanie wyrobu, nie mogą stanowić podstawy do roszczeń wobec naszej firmy.\r\n\r\nMateriał powierzone i dostarczany przez Klienta\r\nNie ponosimy odpowiedzialności za uszkodzenia, błędy obróbki, zmiany struktury, odkształcenia ani inne wady powstałe w wyniku specyficznych właściwości materiału powierzonego przez Klienta, jego niejednorodności, błędnej deklaracji gatunku, braku wymaganych atestów czy oznaczeń partii. Klient zobowiązany jest dostarczyć materiał zgodny ze specyfikacją oraz wolny od wad fizycznych i chemicznych, mogących negatywnie wpływać na proces cięcia i jakość finalnego wyrobu.\r\n\r\nDostawcy materiałów\r\nNasza firma dołoży wszelkich starań w zakresie selekcji i zakupów materiałów wyłącznie od sprawdzonych dostawców. Zastrzegamy sobie jednak, że odpowiedzialność za parametry, właściwości lub wady ukryte materiału ogranicza się wyłącznie do zakresu wynikającego z dokumentacji danego producenta lub certyfikatu jakości – zgodnie z obowiązującym prawem oraz praktyką rynku stalowego.\r\n\r\nOgraniczenie odpowiedzialności prawnej\r\nOdpowiadamy wyłącznie za zgodność wykonanych prac z przesłaną dokumentacją oraz z obowiązującymi normami i przepisami prawa. Nie ponosimy odpowiedzialności za ewentualne szkody pośrednie, utracone korzyści, koszty produkcji, opóźnienia wynikające z przerw w dostawie materiałów, siły wyższej, zdarzeń losowych czy skutków niezastosowania się Klienta do obowiązujących przepisów i wymogów technicznych.\r\n\r\nPrzepisy prawa i gwarancje\r\nWszelkie realizacje podlegają przepisom prawa polskiego, normom branżowym oraz ustaleniom indywidualnym zawartym w zamówieniu. Ewentualna odpowiedzialność spółki ogranicza się do wartości usługi, a w szczególnych wypadkach – do ponownego wykonania usługi lub zwrotu jej kosztu. Nie udzielamy gwarancji na materiały powierzone, a zakres gwarancji na produkty wykonane z własnych materiałów jest określony indywidualnie w ofercie i na fakturze.\r\n\r\nMamy nadzieję, że powyższe wyjaśnienia pozwolą na jasne i czytelne określenie zasad współpracy oraz przyczynią się do pomyślnej realizacji Państwa zamówienia. Zapraszamy do zapoznania się ze szczegółami przygotowanej oferty oraz kontaktu w przypadku pytań lub wątpliwości.\r\n\r\nZ wyrazami szacunku,\r\nLaserTeam")

ttk.Label(left_frame, text="Read files:").grid(row=12, column=0, sticky="ne")
file_list = tk.Listbox(left_frame, height=5, width=50, bg="#3c3c3c", fg="white")
file_list.grid(row=12, column=1)

def open_selected_file(event=None):
    sel = file_list.curselection()
    if sel:
        f = os.path.join(folder_var.get(), file_list.get(sel[0]))
        try:
            os.startfile(f)
            analysis_logger.log(f"Opened file: {file_list.get(sel[0])}", "INFO")
        except Exception as e:
            analysis_logger.log(f"Failed to open file: {str(e)}", "ERROR")

file_list.bind('<Double-Button-1>', open_selected_file)

buttons_frame = tk.Frame(left_frame, bg="#2c2c2c")
buttons_frame.grid(row=13, column=1, sticky="s")

# RIGHT
right_frame = tk.Frame(root, bg="#2c2c2c")
right_frame.pack(side="right", padx=10, pady=10, fill="both", expand=True)

right_paned = tk.PanedWindow(right_frame, orient=tk.VERTICAL, bg="#2c2c2c", sashrelief="raised", borderwidth=1)
right_paned.pack(fill="both", expand=True)

panel_a = tk.PanedWindow(right_paned, orient=tk.VERTICAL, bg="#2c2c2c", sashrelief="raised", borderwidth=1)

# --- PANEL 1 ---
subpanel1 = tk.LabelFrame(panel_a, text="PANEL 1 – PREVIEW", bg="#2c2c2c", fg="white")
columns = ("1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11")
tree = ttk.Treeview(subpanel1, columns=columns, show="tree headings")
tree.column("#0", width=150, minwidth=100, stretch=tk.NO)  # Increased width for better thumbnail visibility
tree.heading("1", text="Nr");     tree.column("1", minwidth=50,  width=50,  stretch=tk.NO)
tree.heading("2", text="SubNr");  tree.column("2", minwidth=50,  width=50,  stretch=tk.NO)
tree.heading("3", text="Name");  tree.column("3", minwidth=150, width=400, stretch=tk.NO)
tree.heading("4", text="Material"); tree.column("4", minwidth=50, width=80, stretch=tk.NO)
tree.heading("5", text="Thickness");  tree.column("5", minwidth=50, width=80, stretch=tk.NO, anchor="e")
tree.heading("6", text="Quantity");    tree.column("6", minwidth=50, width=80, stretch=tk.NO, anchor="e")
tree.heading("7", text="L+M Cost");    tree.column("7", minwidth=50, width=100, stretch=tk.NO, anchor="e")
tree.heading("8", text="Bending/pc."); tree.column("8", minwidth=50, width=100, stretch=tk.NO, anchor="e")
tree.heading("9", text="Additional/pc."); tree.column("9", minwidth=50, width=120, stretch=tk.NO, anchor="e")
tree.heading("10", text="Weight"); tree.column("10", minwidth=50, width=80, stretch=tk.NO, anchor="e")
tree.heading("11", text="Cutting length"); tree.column("11", minwidth=50, width=120, stretch=tk.NO, anchor="e")

# Add scrollbar for treeview
scrollbar = ttk.Scrollbar(subpanel1, orient="vertical", command=tree.yview)
tree.configure(yscrollcommand=scrollbar.set)
tree.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")


# ========= BOTTOM-LEFT buttons (add to existing buttons_frame) =========
ttk.Button(buttons_frame, text="Save Project", command=save_project_ui).grid(row=0, column=0, padx=5, pady=5, sticky="we")
ttk.Button(buttons_frame, text="Load Project", command=load_project_ui).grid(row=0, column=1, padx=5, pady=5, sticky="we")
buttons_frame.grid_columnconfigure(0, weight=1)
buttons_frame.grid_columnconfigure(1, weight=1)

def edit_cell(event):
    item = tree.identify_row(event.y)
    column = tree.identify_column(event.x)
    if not item or not column:
        return
    col_index = int(column[1:]) - 1
    if col_index in [5, 6, 7, 8]:
        x, y, w, h = tree.bbox(item, column)
        e = tk.Entry(subpanel1, bg="#3c3c3c", fg="white", insertbackground="white")
        e.place(x=x, y=y, width=w, height=h)
        e.insert(0, tree.item(item, 'values')[col_index])
        e.focus()
        def save_edit(_):
            vals = list(tree.item(item, 'values'))
            vals[col_index] = e.get()
            tree.item(item, values=vals)
            e.destroy()
            update_total()  # Recalculate total after edit
        e.bind("<Return>", save_edit); e.bind("<FocusOut>", save_edit)

tree.bind("<Double-1>", edit_cell)
panel_a.add(subpanel1, minsize=220)

# --- PANEL 2 ---
subpanel2 = tk.LabelFrame(panel_a, text="PANEL 2 – FIXED COSTS AND CALCULATION", bg="#2c2c2c", fg="white")

# Left column - Fixed costs
ttk.Label(subpanel2, text="Operational costs per sheet:").grid(row=0, column=0, sticky="w", padx=(5,10))
op_cost_entry = ttk.Entry(subpanel2)
op_cost_entry.grid(row=0, column=1, padx=(0,20))
op_cost_entry.insert(tk.INSERT, "40,00")

ttk.Label(subpanel2, text="Technology per order:").grid(row=1, column=0, sticky="w", padx=(5,10))
tech_order_entry = ttk.Entry(subpanel2)
tech_order_entry.grid(row=1, column=1, padx=(0,20))
tech_order_entry.insert(tk.INSERT, "50,00")

ttk.Label(subpanel2, text="Additional costs for the order:").grid(row=2, column=0, sticky="w", padx=(5,10))
add_order_cost_entry = ttk.Entry(subpanel2)
add_order_cost_entry.grid(row=2, column=1, padx=(0,20))
add_order_cost_entry.insert(tk.INSERT, "0,00")
# --- LEFT: Cutting time calculation (cols 2-3) ---
ttk.Label(subpanel2, text="CUTTING TIME CALCULATION",
          font=("Arial", 10, "bold")).grid(row=0, column=2, columnspan=2, pady=(0, 5), padx=(20, 5), sticky="w")

ttk.Label(subpanel2, text="O₂ cutting rate [PLN/h]:").grid(row=1, column=2, sticky="w", padx=(20,10))
oxygen_rate_entry = ttk.Entry(subpanel2)
oxygen_rate_entry.grid(row=1, column=3, padx=(0,5), sticky="we")
oxygen_rate_entry.insert(tk.INSERT, "350,00")

ttk.Label(subpanel2, text="N₂ cutting rate [PLN/h]:").grid(row=2, column=2, sticky="w", padx=(20,10))
nitrogen_rate_entry = ttk.Entry(subpanel2)
nitrogen_rate_entry.grid(row=2, column=3, padx=(0,5), sticky="we")
nitrogen_rate_entry.insert(tk.INSERT, "550,00")

ttk.Label(subpanel2, text="AL N₂ cutting rate [PLN/h]:").grid(row=3, column=2, sticky="w", padx=(20,10))
al_nitrogen_rate_entry = ttk.Entry(subpanel2)
al_nitrogen_rate_entry.grid(row=3, column=3, padx=(0,5), sticky="we")
al_nitrogen_rate_entry.insert(tk.INSERT, "650,00")


# --- RIGHT: Real TKW (cols 4-5) ---
ttk.Label(subpanel2, text="Real TKW",
          font=("Arial", 10, "bold")).grid(row=0, column=4, columnspan=2, pady=(0, 5), padx=(20, 5), sticky="w")

ttk.Label(subpanel2, text="O₂ cutting rate TKW [PLN/h]:").grid(row=1, column=4, sticky="w", padx=(20,10))
oxygen_rate_entry_TKW = ttk.Entry(subpanel2)
oxygen_rate_entry_TKW.grid(row=1, column=5, padx=(0,5), sticky="we")
oxygen_rate_entry_TKW.insert(tk.INSERT, "262,50")

ttk.Label(subpanel2, text="N₂ cutting rate TKW [PLN/h]:").grid(row=2, column=4, sticky="w", padx=(20,10))
nitrogen_rate_entry_TKW = ttk.Entry(subpanel2)
nitrogen_rate_entry_TKW.grid(row=2, column=5, padx=(0,5), sticky="we")
nitrogen_rate_entry_TKW.insert(tk.INSERT, "412,50")

ttk.Label(subpanel2, text="AL N₂ cutting rate TKW [PLN/h]:").grid(row=3, column=4, sticky="w", padx=(20,10))
al_nitrogen_rate_entry_TKW = ttk.Entry(subpanel2)
al_nitrogen_rate_entry_TKW.grid(row=3, column=5, padx=(0,5), sticky="we")
al_nitrogen_rate_entry_TKW.insert(tk.INSERT, "487,50")

ttk.Label(subpanel2, text="Bending percent TKW in cust.price [%]:").grid(row=4, column=4, sticky="w", padx=(20,10))
bending_percent_entry_TKW = ttk.Entry(subpanel2)
bending_percent_entry_TKW.grid(row=4, column=5, padx=(0,5), sticky="we")
bending_percent_entry_TKW.insert(tk.INSERT, "75,00")

# (opcjonalnie) dopasowanie szerokości kolumn dla lepszego layoutu
subpanel2.grid_columnconfigure(3, weight=1)
subpanel2.grid_columnconfigure(5, weight=1)

# Separator
ttk.Label(subpanel2, text="").grid(row=3, column=0, columnspan=4, pady=5)

# Time and cost displays - spanning both column groups
ttk.Label(subpanel2, text="O₂ cutting time [h]:").grid(row=4, column=0, sticky="w", padx=(5,10))
oxygen_time_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
oxygen_time_label.grid(row=4, column=1, sticky="ew", padx=(0,20))

ttk.Label(subpanel2, text="N₂ cutting time [h]:").grid(row=4, column=2, sticky="w", padx=(20,10))
nitrogen_time_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
nitrogen_time_label.grid(row=4, column=3, sticky="ew", padx=(0,5))

ttk.Label(subpanel2, text="O₂ cutting cost [PLN]:").grid(row=5, column=0, sticky="w", padx=(5,10))
oxygen_cost_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
oxygen_cost_label.grid(row=5, column=1, sticky="ew", padx=(0,20))

ttk.Label(subpanel2, text="N₂ cutting cost [PLN]:").grid(row=5, column=2, sticky="w", padx=(20,10))
nitrogen_cost_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
nitrogen_cost_label.grid(row=5, column=3, sticky="ew", padx=(0,5))

# Summary section
ttk.Label(subpanel2, text="").grid(row=6, column=0, columnspan=4, pady=5)
ttk.Label(subpanel2, text="COST SUMMARY", font=("Arial", 10, "bold")).grid(row=7, column=0, columnspan=4, pady=(5, 5))

ttk.Label(subpanel2, text="Material cost [PLN]:").grid(row=8, column=0, sticky="w", padx=(5,10))
material_cost_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
material_cost_label.grid(row=8, column=1, sticky="ew", padx=(0,20))

ttk.Label(subpanel2, text="Total cutting cost [PLN]:").grid(row=8, column=2, sticky="w", padx=(20,10))
total_cutting_cost_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
total_cutting_cost_label.grid(row=8, column=3, sticky="ew", padx=(0,5))

ttk.Label(subpanel2, text="Operational costs [PLN]:").grid(row=9, column=0, sticky="w", padx=(5,10))
operational_cost_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=20)
operational_cost_label.grid(row=9, column=1, sticky="ew", padx=(0,20))

ttk.Label(subpanel2, text="").grid(row=10, column=0, columnspan=4, pady=5)

ttk.Label(subpanel2, text="TOTAL OF ALL COSTS [PLN]:").grid(row=11, column=0, columnspan=2, sticky="w", padx=(5,10))
total_all_costs_label = ttk.Label(subpanel2, text="0,00", relief="sunken", anchor="e", width=30, font=("Arial", 11, "bold"))
total_all_costs_label.grid(row=11, column=2, columnspan=2, sticky="ew", padx=(20,5))

ttk.Label(subpanel2, text="TOTAL FOR CORRECTION [PLN]:").grid(row=13, column=0, columnspan=2, sticky="w", padx=(5,10))
total_all_costs_entry = ttk.Entry(subpanel2, width=30, font=("Arial", 11, "bold"))
total_all_costs_entry.grid(row=13, column=2, columnspan=2, sticky="ew", padx=(20,5))
total_all_costs_entry.insert(tk.INSERT, "0,00")

update_prices_button = ttk.Button(subpanel2, text="UPDATE PRICES BASED ON TIME", 
                                  command=lambda: update_prices_based_on_time())
update_prices_button.grid(row=14, column=0, columnspan=4, pady=(10, 5))

# Add margin update button - USER MUST CLICK THIS TO APPLY MARGINS
update_margins_button = ttk.Button(subpanel2, text="UPDATE WITH DYNAMIC MARGINS", 
                                   command=lambda: update_with_margins())
update_margins_button.grid(row=15, column=0, columnspan=4, pady=(5, 10))

# Make button more prominent with styling
style.configure("Accent.TButton", background="#FF6B35", foreground="white", font=("Arial", 10, "bold"))
style.map("Accent.TButton", background=[('active', '#FF8C5A')])
update_margins_button.configure(style="Accent.TButton")

# Configure column weights for proper resizing
subpanel2.grid_columnconfigure(1, weight=1)
subpanel2.grid_columnconfigure(3, weight=1)

# Add a frame to hold the buttons and ensure they're visible
button_frame = tk.Frame(subpanel2, bg="#2c2c2c")
button_frame.grid(row=16, column=0, columnspan=4, pady=(10, 10), sticky="ew")

# Add visual separator
ttk.Separator(button_frame, orient='horizontal').pack(fill='x', pady=(0, 10))

# Add label to make button purpose clear
ttk.Label(button_frame, text="APPLY MARGINS TO PRICES:", font=("Arial", 10, "bold")).pack()

# The UPDATE WITH DYNAMIC MARGINS button - clearly visible
apply_margins_btn = tk.Button(
    button_frame, 
    text="UPDATE WITH DYNAMIC MARGINS", 
    command=lambda: update_with_margins(),
    bg="#4CAF50",
    fg="white",
    font=("Arial", 11, "bold"),
    relief="raised",
    bd=3,
    cursor="hand2",
    padx=20,
    pady=10
)
apply_margins_btn.pack(pady=(5, 0))

# Add helper text
ttk.Label(button_frame, text="Click above to apply the proposed margins to all prices", 
         font=("Arial", 9, "italic")).pack(pady=(2, 0))

subpanel2.update_idletasks()
panel2_height = subpanel2.winfo_reqheight() + 50  # Added extra height
panel_a.add(subpanel2, height=panel2_height, minsize=panel2_height)

# Add event handlers for automatic recalculation
oxygen_rate_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
nitrogen_rate_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
al_nitrogen_rate_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
op_cost_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
tech_order_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
add_order_cost_entry.bind('<FocusOut>', lambda e: update_cost_calculations() if all_parts else None)
total_all_costs_entry.bind('<FocusOut>', lambda e: validate_total_entry() if all_parts else None)

# --- PANEL 3 ---
subpanel3 = tk.LabelFrame(panel_a, text="PANEL 3 – PRICE LISTS AND TESTS", bg="#2c2c2c", fg="white", padx=6, pady=6)

mat_frame = tk.LabelFrame(subpanel3, text="Material price list (PLN/kg)", bg="#2c2c2c", fg="white")
mat_frame.grid(row=0, column=0, sticky="nwe", padx=4, pady=4)
def _update_led(canvas, ok): canvas.delete("all"); canvas.create_oval(2,2,18,18, fill=("green" if ok else "red"))
btn_load_mat = ttk.Button(mat_frame, text="Load material price list", command=lambda: load_material_prices(preview=True))
btn_load_mat.grid(row=0, column=0, sticky="w")
material_led = tk.Canvas(mat_frame, width=20, height=20, bg="#2c2c2c", highlightthickness=0); material_led.grid(row=0, column=1, padx=(6,0)); _update_led(material_led, False)
ttk.Label(mat_frame, text="Material:").grid(row=1, column=0, sticky="e", pady=(6,0))
material_var = tk.StringVar(); material_cb = ttk.Combobox(mat_frame, textvariable=material_var, width=22, state="readonly"); material_cb.grid(row=1, column=1, sticky="w", pady=(6,0))
ttk.Label(mat_frame, text="Thickness [mm]:").grid(row=2, column=0, sticky="e")
thickness_mat_var = tk.StringVar(); thickness_mat_cb = ttk.Combobox(mat_frame, textvariable=thickness_mat_var, width=12, state="readonly"); thickness_mat_cb.grid(row=2, column=1, sticky="w")
btn_find_mat = ttk.Button(mat_frame, text="Find material price", command=lambda: ui_find_material_price()); btn_find_mat.grid(row=3, column=0, columnspan=2, pady=4, sticky="we")
material_result_label = ttk.Label(mat_frame, text="Material Result: —"); material_result_label.grid(row=4, column=0, columnspan=2, sticky="w")

cut_frame = tk.LabelFrame(subpanel3, text="Cutting price list (PLN/m)", bg="#2c2c2c", fg="white")
cut_frame.grid(row=0, column=1, sticky="nwe", padx=4, pady=4)
btn_load_cut = ttk.Button(cut_frame, text="Load cutting price list", command=lambda: load_cutting_prices(preview=True))
btn_load_cut.grid(row=0, column=0, sticky="w")
cutting_led = tk.Canvas(cut_frame, width=20, height=20, bg="#2c2c2c", highlightthickness=0); cutting_led.grid(row=0, column=1, padx=(6,0)); _update_led(cutting_led, False)
ttk.Label(cut_frame, text="Material:").grid(row=1, column=0, sticky="e", pady=(6,0))
material_cut_var = tk.StringVar(); material_cut_cb = ttk.Combobox(cut_frame, textvariable=material_cut_var, width=22, state="readonly"); material_cut_cb.grid(row=1, column=1, sticky="w", pady=(6,0))
ttk.Label(cut_frame, text="Thickness [mm]:").grid(row=2, column=0, sticky="e")
thickness_cut_var = tk.StringVar(); thickness_cut_cb = ttk.Combobox(cut_frame, textvariable=thickness_cut_var, width=12, state="readonly"); thickness_cut_cb.grid(row=2, column=1, sticky="w")
ttk.Label(cut_frame, text="Gas:").grid(row=3, column=0, sticky="e")
gas_var = tk.StringVar(); gas_cb = ttk.Combobox(cut_frame, textvariable=gas_var, width=12, state="readonly"); gas_cb.grid(row=3, column=1, sticky="w")
btn_find_cut = ttk.Button(cut_frame, text="Find cutting price", command=lambda: ui_find_cutting_price()); btn_find_cut.grid(row=4, column=0, columnspan=2, pady=4, sticky="we")
cutting_result_label = ttk.Label(cut_frame, text="Cutting Result: —"); cutting_result_label.grid(row=5, column=0, columnspan=2, sticky="w")

btn_load_both = ttk.Button(subpanel3, text="Load both price lists and refresh lists",
                          command=lambda: (load_material_prices(True), load_cutting_prices(True)))
btn_load_both.grid(row=1, column=0, columnspan=2, sticky="we", padx=4, pady=(2,6))
subpanel3.grid_columnconfigure(0, weight=1); subpanel3.grid_columnconfigure(1, weight=1)

panel_a.add(subpanel3, minsize=200)
right_paned.add(panel_a)

# CORRECTED Dynamic margin calculation functions
def calculate_material_margin(plate_area_m2):
    """Calculate material margin based on plate area using linear interpolation
    250% to 0% for areas from 0 to 1m²"""
    min_area = _parse_float(min_area_var.get()) or 0.00  # Default 0 m²
    max_area = _parse_float(max_area_var.get()) or 1.0   # Default 1.0 m²
    
    if plate_area_m2 <= min_area:
        return 250.0  # 250% margin for zero or very small areas
    elif plate_area_m2 >= max_area:
        return 0.0    # 0% margin for areas >= 1m²
    else:
        # Linear interpolation between 250% and 0%
        ratio = (plate_area_m2 - min_area) / (max_area - min_area)
        return 250.0 * (1.0 - ratio)

def calculate_cutting_margin(cutting_length_mm):
    """Calculate cutting margin based on cutting length using linear interpolation
    200% to 0% for lengths from 0 to 5000mm"""
    min_length = _parse_float(min_cutting_var.get()) or 0.0     # Default 0mm
    max_length = _parse_float(max_cutting_var.get()) or 5000.0  # Default 5000mm
    
    if cutting_length_mm <= min_length:
        return 200.0  # 200% margin for zero or very short cutting
    elif cutting_length_mm >= max_length:
        return 0.0    # 0% margin for lengths >= 5000mm
    else:
        # Linear interpolation between 200% and 0%
        ratio = (cutting_length_mm - min_length) / (max_length - min_length)
        return 200.0 * (1.0 - ratio)

def parse_plate_size(plate_size_str):
    """Parse plate size string like '500*300' or '500x300' to get area in m²"""
    if not plate_size_str:
        return 0.0
    
    try:
        # Handle different separators
        size_str = str(plate_size_str).strip().replace('x', '*').replace('X', '*')
        if '*' in size_str:
            dimensions = size_str.split('*')
            if len(dimensions) >= 2:
                width_mm = _parse_float(dimensions[0]) or 0.0
                height_mm = _parse_float(dimensions[1]) or 0.0
                # Convert mm² to m²
                return (width_mm * height_mm) / 1000000.0
    except Exception as e:
        analysis_logger.log(f"Error parsing plate size '{plate_size_str}': {e}", "WARNING")
    
    return 0.0

# USER-TRIGGERED FUNCTION TO APPLY MARGINS
def update_with_margins():
    """Update all costs with dynamic margins - USER MUST CLICK BUTTON TO TRIGGER THIS"""
    global all_parts, total_row_iid, avg_material_margin, avg_cutting_margin
    
    if not all_parts:
        messagebox.showwarning("Warning", "No data to update. Perform analysis first.")
        return
    
    analysis_logger.log("USER REQUESTED: APPLYING DYNAMIC MARGINS", "PHASE")
    
    try:
        # Get proposed margins from input fields (user may have modified them)
        proposed_material = _parse_float(material_margin_var.get()) or 0.0
        proposed_cutting = _parse_float(cutting_margin_var.get()) or 0.0
        
        analysis_logger.log(f"Applying user-selected margins: Material {proposed_material}%, Cutting {proposed_cutting}%", "INFO")
        
        total_new_cost = 0.0
        
        for i, part in enumerate(all_parts):
            # Get current values from treeview
            item_iid = list(tree.get_children())[i] if i < len(tree.get_children()) else None
            if not item_iid or item_iid == total_row_iid:
                continue
            
            # Get current tree values in case user edited them
            vals = tree.item(item_iid, 'values')
            current_qty = int(vals[5] or 0)
            current_bending = _parse_float(vals[7]) or 0.0
            current_additional = _parse_float(vals[8]) or 0.0
            
            # Calculate base costs with mandatory 7% minimum margin for material
            base_material_cost = part.get('adj_weight', 0.0) * part.get('base_price_per_kg', 0.0) * 1.07
            base_cut_cost = part.get('cut_length', 0.0) * part.get('base_rate_per_cut_length', 0.0)
            
            # Apply user-selected margins
            material_cost_with_margin = base_material_cost * (1.0 + proposed_material / 100.0)
            cutting_cost_with_margin = base_cut_cost * (1.0 + proposed_cutting / 100.0)
            
            # Add other costs
            contour_cost = part.get('contours_qty', 0.0) * part.get('rate_per_contour', 0.0)
            marking_cost = part.get('marking_length', 0.0) * part.get('rate_per_marking_length', 0.0)
            defilm_cost = part.get('defilm_length', 0.0) * part.get('rate_per_defilm_length', 0.0)
            
            # Calculate overhead per part
            op_cost_per_sheet = _parse_float(op_cost_entry.get()) or 0.0
            tech_per_order = _parse_float(tech_order_entry.get()) or 0.0
            add_costs_order = _parse_float(add_order_cost_entry.get()) or 0.0
            
            if total_parts_qty > 0:
                extra_per_part = (tech_per_order + add_costs_order) / total_parts_qty
                op_cost_per_part = (total_sheets * op_cost_per_sheet) / total_parts_qty
            else:
                extra_per_part = 0.0
                op_cost_per_part = 0.0
            
            # Final unit cost
            new_unit_cost = (material_cost_with_margin + cutting_cost_with_margin + contour_cost + 
                           marking_cost + defilm_cost + extra_per_part + op_cost_per_part)
            
            # Update part data
            part['cost_per_unit'] = round(new_unit_cost, 2)
            part['qty'] = current_qty
            part['bending_per_unit'] = current_bending
            part['additional_per_unit'] = current_additional
            
            # Update treeview
            new_vals = list(vals)
            new_vals[6] = format_pln(new_unit_cost)  # L+M Cost column
            tree.item(item_iid, values=new_vals)
            
            # Calculate total for this part
            total_new_cost += (new_unit_cost + current_bending + current_additional) * current_qty
        
        # Update total row
        if total_row_iid:
            tree.set(total_row_iid, column="7", value=format_pln(total_new_cost))
            SetTotalPricePerOrder(total_new_cost)
        
        # Update cost calculations
        update_cost_calculations()
        
        analysis_logger.log(f"Successfully applied user margins. New total: {format_pln(total_new_cost)} PLN", "SUCCESS")
        messagebox.showinfo("Success", f"Margins applied successfully!\n"
                                      f"Material margin: {proposed_material}%\n"
                                      f"Cutting margin: {proposed_cutting}%\n"
                                      f"New total: {format_pln(total_new_cost)} PLN")
        
    except Exception as e:
        analysis_logger.log(f"Error applying margins: {str(e)}", "ERROR")
        messagebox.showerror("Error", f"Failed to apply margins: {str(e)}")

def update_prices_based_on_time():
    """Update unit prices in treeview based on time calculations and proportional distribution"""
    global all_parts, total_row_iid
    
    if not all_parts:
        messagebox.showwarning("Warning", "No data to update. Perform analysis first.")
        return
    
    # Get the target total from the editable field
    value_str = total_all_costs_entry.get().strip()
    value_str = value_str.replace(' ', '').replace(',', '.')
    value_str = ''.join(c for c in value_str if c.isdigit() or c == '.')
    target_total = _parse_float(value_str) if value_str else None
    if not target_total or target_total <= 0:
        messagebox.showerror("Error", "Invalid total costs.")
        return
    
    # Calculate current total from treeview (excluding total row)
    current_total = 0.0
    items_data = []
    
    for item in tree.get_children():
        if item == total_row_iid:
            continue
        vals = tree.item(item, 'values')
        qty = _parse_float(vals[5]) or 0
        cost = _parse_float(vals[6]) or 0
        bending = _parse_float(vals[7]) or 0
        additional = _parse_float(vals[8]) or 0
        
        item_total = (cost + bending + additional) * qty
        current_total += item_total
        
        items_data.append({
            'item': item,
            'qty': qty,
            'cost': cost,
            'bending': bending,
            'additional': additional,
            'item_total': item_total,
            'proportion': 0.0
        })
    
    if current_total <= 0:
        messagebox.showerror("Error", "No costs to recalculate.")
        return
    
    # Calculate proportions for each item
    for item_data in items_data:
        item_data['proportion'] = item_data['item_total'] / current_total
    
    # Apply proportional distribution of the new total
    new_grand_total = 0.0
    for idx, item_data in enumerate(items_data):
        # Calculate new item total based on proportion
        new_item_total = target_total * item_data['proportion']
        
        # Calculate new unit cost (preserving bending and additional costs)
        if item_data['qty'] > 0:
            new_unit_cost = (new_item_total / item_data['qty']) - item_data['bending'] - item_data['additional']
            
            # Ensure non-negative cost
            new_unit_cost = max(0, new_unit_cost)
            
            # Update tree item
            vals = list(tree.item(item_data['item'], 'values'))
            vals[6] = format_pln(new_unit_cost)
            tree.item(item_data['item'], values=vals)
            
            # Update all_parts array
            if idx < len(all_parts):
                all_parts[idx]['cost_per_unit'] = new_unit_cost
            
            new_grand_total += (new_unit_cost + item_data['bending'] + item_data['additional']) * item_data['qty']
    
    # Update the total row
    tree.set(total_row_iid, column="7", value=format_pln(new_grand_total))
    SetTotalPricePerOrder(new_grand_total)
    
    messagebox.showinfo("Success", f"Prices have been updated proportionally.\n"
                                  f"Old sum: {format_pln(current_total)}\n"
                                  f"New sum: {format_pln(new_grand_total)}")

def validate_total_entry():
    """Validate and format the manually entered total"""
    try:
        value_str = total_all_costs_entry.get().strip()
        value_str = value_str.replace(' ', '').replace(',', '.')
        value_str = ''.join(c for c in value_str if c.isdigit() or c == '.')
        
        if value_str:
            value = float(value_str)
            if value is not None:
                total_all_costs_entry.delete(0, tk.END)
                total_all_costs_entry.insert(0, format_pln(value))
    except ValueError:
        pass

# ---- Price list loaders ----
def _tree_preview_clear_and_headers(headers):
    for item in tree.get_children():
        tree.delete(item)
    tree.insert('', 'end', values=(0, '', ' | '.join(headers), '', '', '', '', '', ''))

def load_material_prices(preview=False):
    global material_prices, _mat_set, _thk_set
    material_prices.clear(); _mat_set.clear(); _thk_set.clear()
    try:
        if not os.path.exists(MATERIALS_FILE):
            raise FileNotFoundError(f"File not found: {MATERIALS_FILE}")
        wb = load_workbook(MATERIALS_FILE, data_only=True)
        sheet = wb.active
        headers = [str(c.value).strip().lower() if c.value is not None else "" for c in next(sheet.iter_rows(min_row=1, max_row=1))]
        need = ("material", "thickness", "price")
        idx = {n: headers.index(n) for n in need if n in headers}
        if not set(need).issubset(idx):
            raise ValueError("Missing required columns: material, thickness, price")
        if preview: _tree_preview_clear_and_headers(["materials prices.xlsx → material/thickness/price"])
        for row in sheet.iter_rows(min_row=2, values_only=True):
            mat = _norm_s(row[idx["material"]]); thk = _parse_float(row[idx["thickness"]]); prc = _parse_float(row[idx["price"]])
            if mat and thk is not None and prc is not None:
                material_prices[(mat, thk)] = prc; _mat_set.add(mat); _thk_set.add(thk)
                if preview: tree.insert('', 'end', values=("", "", f"{mat} @ {thk:.2f} mm → {format_pln(prc)} PLN/kg", "", "", "", "", "", ""))
        mats_sorted = sorted(_mat_set); thk_sorted = [f"{t:.2f}".rstrip("0").rstrip(".") for t in sorted(_thk_set)]
        material_cb["values"] = mats_sorted; material_cut_cb["values"] = mats_sorted
        thickness_mat_cb["values"] = thk_sorted
        if not thickness_cut_cb["values"]: thickness_cut_cb["values"] = thk_sorted
        _update_led(material_led, len(material_prices) > 0)
    except Exception as e:
        _update_led(material_led, False); messagebox.showerror("Error", f"Loading material prices:\n{e}")

def load_cutting_prices(preview=False):
    global cutting_prices, _mat_set, _thk_set, _gas_set
    cutting_prices.clear(); _gas_set.clear()
    try:
        if not os.path.exists(CUTTING_FILE):
            raise FileNotFoundError(f"File not found: {CUTTING_FILE}")
        wb = load_workbook(CUTTING_FILE, data_only=True)
        sheet = wb.active
        headers = [str(c.value).strip().lower() if c.value is not None else "" for c in next(sheet.iter_rows(min_row=1, max_row=1))]
        need = ("thickness", "material", "gas", "price")
        idx = {n: headers.index(n) for n in need if n in headers}
        if not set(need).issubset(idx):
            raise ValueError("Missing required columns: thickness, material, gas, price")
        if preview: _tree_preview_clear_and_headers(["cutting prices.xlsx → thickness/material/gas/price"])
        for row in sheet.iter_rows(min_row=2, values_only=True):
            thk = _parse_float(row[idx["thickness"]]); mat = _norm_s(row[idx["material"]]); gas = _norm_s(row[idx["gas"]]); prc = _parse_float(row[idx["price"]])
            if thk is not None and mat and gas and prc is not None:
                cutting_prices[(thk, mat, gas)] = prc; _mat_set.add(mat); _thk_set.add(thk); _gas_set.add(gas)
                if preview: tree.insert('', 'end', values=("", "", f"{thk:.2f} mm / {mat} / {gas} → {format_pln(prc)} PLN/m", "", "", "", "", "", ""))
        mats_sorted = sorted(_mat_set); thk_sorted = [f"{t:.2f}".rstrip("0").rstrip(".") for t in sorted(_thk_set)]; gas_sorted = sorted(_gas_set)
        material_cut_cb["values"] = mats_sorted
        if not material_cb["values"]: material_cb["values"] = mats_sorted
        thickness_cut_cb["values"] = thk_sorted
        if not thickness_mat_cb["values"]: thickness_mat_cb["values"] = thk_sorted
        gas_cb["values"] = gas_sorted
        _update_led(cutting_led, len(cutting_prices) > 0)
    except Exception as e:
        _update_led(cutting_led, False); messagebox.showerror("Error", f"Loading cutting prices:\n{e}")

# ---- UI tests (Panel 3) ----
def ui_find_material_price():
    mat = _norm_s(material_var.get()); thk = _parse_float(thickness_mat_var.get())
    if not mat or thk is None:
        messagebox.showerror("Error", "Fill in Material and Thickness (mm)."); return
    price = material_prices.get((mat, thk))
    material_result_label.config(text="Material Result: not found" if price is None else f"Material Result: {format_pln(price)} PLN/kg")

def ui_find_cutting_price():
    mat = _norm_s(material_cut_var.get()); thk = _parse_float(thickness_cut_var.get()); gas = _norm_s(gas_var.get())
    if not mat or thk is None or not gas:
        messagebox.showerror("Error", "Fill in Material, Thickness (mm) and Gas."); return
    price = cutting_prices.get((thk, mat, gas))
    cutting_result_label.config(text="Cutting Result: not found" if price is None else f"Cutting Result: {format_pln(price)} PLN/m")

# ---- Folder analysis ----
last_groups = []; last_total_cost = 0.0; last_folder_path = ""

def _ensure_cenniki_loaded():
    ok = True
    if not material_prices:
        try: load_material_prices(preview=False)
        except Exception: ok = False
    if not cutting_prices:
        try: load_cutting_prices(preview=False)
        except Exception: ok = False
    return ok

def get_total_cut_length(ws, text="Total") -> float:
    """
    Searches in column A for the first cell containing 'text' (case-insensitive),
    then reads the value from column H (8) in the same row.
    Returns float; handles Polish format '312,51'.
    """
    for cell in ws['A']:
        val = cell.value
        if val and str(text).lower() in str(val).lower():
            raw = ws.cell(row=cell.row, column=8).value  # col. H
            if isinstance(raw, (int, float)):
                return float(raw)
            try:
                return float(str(raw).replace(" ", "").replace("\xa0", "").replace(",", "."))
            except Exception:
                return 0.0
    raise ValueError(f"No row with text '{text}' found in column A")

def parse_duration_to_hours(value) -> float:
    """
    Converts '1h26min21s', '1h26m21s', '86min', '90s', '1:26:21', '1:26' etc. to float hours.
    Returns 0.0 if unable to parse.
    """
    if value is None:
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)

    s = str(value).strip().lower().replace(" ", "")

    # H:M:S or H:M
    if ":" in s:
        parts = s.split(":")
        try:
            def to_f(x): return float(x.replace(",", "."))
            if len(parts) == 3:
                h, m, sec = parts
            elif len(parts) == 2:
                h, m = parts; sec = "0"
            else:
                h, m, sec = parts[0], "0", "0"
            return to_f(h) + to_f(m)/60.0 + to_f(sec)/3600.0
        except Exception:
            pass

    def take(pattern):
        m = re.search(pattern, s)
        return float(m.group(1).replace(",", ".")) if m else 0.0

    h = take(r"(\d+(?:[.,]\d+)?)h")
    m = take(r"(\d+(?:[.,]\d+)?)(?:m|min)")
    sec = take(r"(\d+(?:[.,]\d+)?)s")

    if h == 0.0 and m == 0.0 and sec == 0.0:
        # Just '123' → treat as seconds
        try:
            seconds = float(s.replace(",", "."))
            return seconds / 3600.0
        except Exception:
            return 0.0

    return h + m/60.0 + sec/3600.0

def _interp(x, pts):
    """Linear interpolation over sorted (x, y) points. Clamps outside the range."""
    pts = sorted(pts, key=lambda p: p[0])
    if x <= pts[0][0]:
        return pts[0][1]
    if x >= pts[-1][0]:
        return pts[-1][1]
    for (x0, y0), (x1, y1) in zip(pts[:-1], pts[1:]):
        if x0 <= x <= x1:
            if x1 == x0:
                return y0
            t = (x - x0) / (x1 - x0)
            return y0 + t * (y1 - y0)
    return pts[-1][1]

def update_cost_calculations():
    """Update all cost calculation displays in Panel 2"""
    global oxygen_cutting_time, nitrogen_cutting_time, aluminum_nitrogen_cutting_time, total_material_cost
    global oxygen_rate_entry_tkw, nitrogen_rate_entry_tkw, al_nitrogen_rate_entry_tkw, bending_percent_entry_tkw

    # Get rates from entries
    oxygen_rate = _parse_float(oxygen_rate_entry.get()) or 350.0
    nitrogen_rate = _parse_float(nitrogen_rate_entry.get()) or 550.0
    al_nitrogen_rate = _parse_float(al_nitrogen_rate_entry.get()) or 650.0
    op_cost_per_sheet = _parse_float(op_cost_entry.get()) or 40.0
    tech_per_order = _parse_float(tech_order_entry.get()) or 0.0
    add_costs_order = _parse_float(add_order_cost_entry.get()) or 0.0
    
    oxygen_rate_tkw = _parse_float(oxygen_rate_entry_TKW.get()) or 262.50
    nitrogen_rate_tkw = _parse_float(nitrogen_rate_entry_TKW.get()) or 412.50
    al_nitrogen_rate_tkw = _parse_float(al_nitrogen_rate_entry_TKW.get()) or 487.50
    bending_percent_entry_tkw = _parse_float(bending_percent_entry_TKW.get()) or 75.0

    # Calculate cutting costs
    oxygen_cost = oxygen_cutting_time * oxygen_rate
    nitrogen_cost = nitrogen_cutting_time * nitrogen_rate + aluminum_nitrogen_cutting_time * al_nitrogen_rate
    total_cutting_cost = oxygen_cost + nitrogen_cost
    
    # Calculate operational costs
    operational_costs = (total_sheets * op_cost_per_sheet) + tech_per_order + add_costs_order
    
    # Calculate total
    total_all_costs = total_material_cost + total_cutting_cost + operational_costs
    
    # Update display labels
    oxygen_time_label.config(text=f"{oxygen_cutting_time:.2f}".replace('.', ','))
    nitrogen_time_label.config(text=f"{nitrogen_cutting_time + aluminum_nitrogen_cutting_time:.2f}".replace('.', ','))
    oxygen_cost_label.config(text=format_pln(oxygen_cost))
    nitrogen_cost_label.config(text=format_pln(nitrogen_cost))
    material_cost_label.config(text=format_pln(total_material_cost))
    total_cutting_cost_label.config(text=format_pln(total_cutting_cost))
    operational_cost_label.config(text=format_pln(operational_costs))
    total_all_costs_label.config(text=format_pln(total_all_costs))

def update_total():
    """Update total in the tree view"""
    global total_row_iid
    
    total = 0.0
    for item in tree.get_children():
        if item == total_row_iid:
            continue
        vals = tree.item(item, 'values')
        qty = _parse_float(vals[5]) or 0
        cost = _parse_float(vals[6]) or 0
        bending = _parse_float(vals[7]) or 0
        additional = _parse_float(vals[8]) or 0
        total += (cost + bending + additional) * qty
    
    if total_row_iid:
        tree.set(total_row_iid, column="7", value=format_pln(total))
        SetTotalPricePerOrder(total)

# References to PhotoImage to prevent images from disappearing (GC)
thumbnail_imgs = []

def analyze_xlsx_folder():
    """ANALYZE WITHOUT APPLYING MARGINS - ONLY 7% MATERIAL MARGIN IS AUTOMATIC"""
    global all_parts, last_groups, last_total_cost, last_folder_path, total_sheets, total_parts_qty, total_row_iid
    global oxygen_cutting_time, nitrogen_cutting_time, aluminum_nitrogen_cutting_time, total_material_cost
    global file_margins, avg_material_margin, avg_cutting_margin
    
    # Clear log and start analysis
    analysis_logger.clear()
    analysis_logger.log("STARTING XLSX FOLDER ANALYSIS (BASE PRICES + 7% MATERIAL MARGIN ONLY)", "PHASE")
    
    # Initialize cutting time accumulators
    oxygen_cutting_time = 0.0
    nitrogen_cutting_time = 0.0
    aluminum_nitrogen_cutting_time = 0.0
    total_material_cost = 0.0
    
    # Initialize margin tracking
    file_margins = []

    for item in tree.get_children():
        tree.delete(item)
    thumbnail_imgs.clear()
    all_parts = []
    
    folder_path = folder_var.get()
    if not folder_path:
        analysis_logger.log("No folder selected", "ERROR")
        messagebox.showerror("Error", "Please select a folder.")
        return
    
    analysis_logger.log(f"Analyzing folder: {folder_path}", "INFO")
    
    try:
        files = [f for f in os.listdir(folder_path) if f.lower().endswith(".xlsx")]
        analysis_logger.log(f"Found {len(files)} XLSX files", "INFO")
    except Exception as e:
        analysis_logger.log(f"Failed to list folder contents: {str(e)}", "ERROR")
        files = []
    
    if not files:
        analysis_logger.log("No XLSX files found in the selected folder", "ERROR")
        messagebox.showerror("Error", "No .xlsx files in the selected folder.")
        return
    
    # Check price lists
    analysis_logger.log("CHECKING PRICE LISTS", "PHASE")
    if not _ensure_cenniki_loaded():
        analysis_logger.log("Price lists not loaded - calculations will use 0.00 values", "WARNING")
        messagebox.showwarning("Warning", "Price lists not loaded – using 0.00, check Panel 3.")

    global op_cost_per_sheet, tech_per_order, add_costs_order
    op_cost_per_sheet = _parse_float(op_cost_entry.get()) or 0.0
    tech_per_order = _parse_float(tech_order_entry.get()) or 0.0
    add_costs_order = _parse_float(add_order_cost_entry.get()) or 0.0
    
    analysis_logger.log(f"Fixed costs: Op/sheet={op_cost_per_sheet:.2f}, Tech/order={tech_per_order:.2f}, Add={add_costs_order:.2f}", "INFO")

    total_sheets = 0
    total_parts_qty = 0
    groups = []
    subnr = 0

    # Process each file
    analysis_logger.log("PROCESSING FILES AND CALCULATING MARGIN SUGGESTIONS", "PHASE")
    
    for file_idx, fname in enumerate(files, 1):
        analysis_logger.log(f"Processing file {file_idx}/{len(files)}: {fname}", "INFO")
        path = os.path.join(folder_path, fname)
        
        file_material_margins = []
        file_cutting_margins = []
        
        try:
            wb = load_workbook(path, data_only=True)
            
            # Check for required sheets
            if "All Task List" not in wb.sheetnames:
                analysis_logger.log(f"Missing 'All Task List' sheet in {fname}", "ERROR")
                raise KeyError("No 'All Task List' sheet")
            
            all_task = wb["All Task List"]
            
            if "All Parts List" not in wb.sheetnames:
                analysis_logger.log(f"Missing 'All Parts List' sheet in {fname}", "WARNING")
                all_part_list = None
                file_thumbnails = {}
            else:
                all_part_list = wb["All Parts List"]
                file_thumbnails = {}
                thumbnail_count = 0
                for img in all_part_list._images:
                    row = img.anchor._from.row + 1
                    col = img.anchor._from.col + 1
                    if col == 2:  # Column B
                        img_data = img._data()
                        file_thumbnails[row] = img_data
                        thumbnail_count += 1
                if thumbnail_count > 0:
                    analysis_logger.log(f"Found {thumbnail_count} thumbnails", "INFO")
            
            # Parse key data
            cut_time = parse_duration_to_hours(all_task['F4'].value)
            if cut_time == 0:
                analysis_logger.log(f"Cut time is zero or invalid in {fname}", "WARNING")
            
            total_cut_length = get_total_cut_length(all_task, "Total")
            if total_cut_length == 0:
                analysis_logger.log(f"Total cut length is zero in {fname}", "WARNING")
            
            material_name = all_task["B4"].value
            thickness_raw = all_task["C4"].value
            gas_raw = all_task["E4"].value
            
            mat_norm = _norm_s(material_name)
            thk_val = _parse_float(thickness_raw)
            gas_key = _map_gas_to_key(gas_raw)

            # Validate critical data
            if not mat_norm:
                analysis_logger.log(f"Material name missing in All Task List!B4", "ERROR")
                raise ValueError("All Task List!B4 (Material) – no value")
            if thk_val is None:
                analysis_logger.log(f"Thickness value invalid in All Task List!C4", "ERROR")
                raise ValueError("All Task List!C4 (Thickness(mm)) – no number")
            if not gas_key:
                analysis_logger.log(f"Gas type '{gas_raw}' not recognized (should be O/N)", "ERROR")
                raise ValueError("All Task List!E4 (Gas) – unsupported gas type")
            
            analysis_logger.log(f"Material: {mat_norm}, Thickness: {thk_val}mm, Gas: {gas_key}", "INFO")

            # Analyze plate sizes for margin CALCULATION (not application)
            analysis_logger.log("Calculating suggested margins (will not be applied automatically)", "INFO")
            
            # Find header row
            plate_size_col = None
            sheets_qty_col = None
            
            for col in range(1, all_task.max_column + 1):
                header_val = all_task.cell(row=7, column=col).value
                if header_val and "Plate Size" in str(header_val):
                    plate_size_col = col
                    analysis_logger.log(f"Found 'Plate Size' in column {col}", "INFO")
                if header_val and ("Sheets" in str(header_val) or col == 4):
                    sheets_qty_col = col
                    analysis_logger.log(f"Found 'Sheets' quantity in column {col}", "INFO")
            
            if not plate_size_col:
                analysis_logger.log("No 'Plate Size(mm*mm)' column found in row 7", "WARNING")
                plate_size_col = 3  # Default to column C
            
            if not sheets_qty_col:
                sheets_qty_col = 4  # Default to column D
            
            # Process data rows (starting from row 8)
            row_idx = 8
            file_total_area = 0.0
            file_total_cutting = 0.0
            row_count = 0
            
            while all_task.cell(row=row_idx, column=sheets_qty_col).value is not None:
                try:
                    # Get plate size and calculate area
                    plate_size_str = all_task.cell(row=row_idx, column=plate_size_col).value
                    plate_area_m2 = parse_plate_size(plate_size_str)
                    
                    # Get number of sheets
                    sheets_qty = _parse_float(all_task.cell(row=row_idx, column=sheets_qty_col).value) or 0
                    
                    # Calculate cutting length for this row (from column H)
                    row_cutting_length = _parse_float(all_task.cell(row=row_idx, column=8).value)*1000 or 0.0
                    
                    # Calculate margins for SUGGESTION ONLY
                    material_margin = calculate_material_margin(plate_area_m2)
                    cutting_margin = calculate_cutting_margin(row_cutting_length)
                    
                    # Accumulate weighted averages
                    file_material_margins.append((material_margin, sheets_qty))
                    file_cutting_margins.append((cutting_margin, row_cutting_length))
                    
                    file_total_area += plate_area_m2 * sheets_qty
                    file_total_cutting += row_cutting_length
                    row_count += 1
                    
                    analysis_logger.log(f"Row {row_idx}: Calculated suggested margins - "
                                      f"Material {material_margin:.1f}%, Cutting {cutting_margin:.1f}%", "INFO")
                    
                except Exception as e:
                    analysis_logger.log(f"Error processing row {row_idx}: {e}", "WARNING")
                
                row_idx += 1
            
            # Calculate average margins for SUGGESTION
            avg_file_material_margin = 0.0
            avg_file_cutting_margin = 0.0
            
            if file_material_margins:
                total_weight = sum(weight for margin, weight in file_material_margins)
                if total_weight > 0:
                    avg_file_material_margin = sum(margin * weight for margin, weight in file_material_margins) / total_weight
            
            if file_cutting_margins:
                total_length = sum(length for margin, length in file_cutting_margins)
                if total_length > 0:
                    avg_file_cutting_margin = sum(margin * length for margin, length in file_cutting_margins) / total_length
            
            analysis_logger.log(f"File suggested margins: Material {avg_file_material_margin:.1f}%, "
                              f"Cutting {avg_file_cutting_margin:.1f}%", "INFO")

            # Accumulate cutting time by gas type
            if gas_key == "O":
                oxygen_cutting_time += cut_time
                analysis_logger.log(f"Added {cut_time:.2f}h to O₂ cutting time", "INFO")
            elif gas_key == "N":
                if 'AL' in mat_norm:
                    aluminum_nitrogen_cutting_time += cut_time
                    analysis_logger.log(f"Added {cut_time:.2f}h to AL N₂ cutting time", "INFO")
                else:
                    nitrogen_cutting_time += cut_time
                    analysis_logger.log(f"Added {cut_time:.2f}h to N₂ cutting time", "INFO")

            # Look up prices
            base_price_per_kg = material_prices.get((mat_norm, thk_val), 0.0)
            if base_price_per_kg == 0.0:
                analysis_logger.log(f"No material price found for {mat_norm} {thk_val}mm - using 0.00", "WARNING")
                messagebox.showerror("Warning", f"No material price found for {mat_norm} {thk_val}mm - using 0.00")
            else:
                analysis_logger.log(f"Material price found: {base_price_per_kg} PLN/kg", "INFO")
            base_rate_per_cut_length = cutting_prices.get((thk_val, mat_norm, gas_key), 0.0)
            if base_rate_per_cut_length == 0.0:
                analysis_logger.log(f"No cutting price found for {mat_norm} {thk_val}mm with {gas_key} - using 0.00", "WARNING")
                messagebox.showerror("Warning", f"No cutting price found for {mat_norm} {thk_val}mm with {gas_key} - using 0.00")
            else:
                analysis_logger.log(f"Cutting price found: {base_rate_per_cut_length} PLN/m", "INFO")
            # Check Cost List sheet
            if "Cost List" not in wb.sheetnames:
                analysis_logger.log(f"Missing 'Cost List' sheet in {fname}", "ERROR")
                raise KeyError("No 'Cost List' sheet")
                messagebox.showerror("Error", "No 'Cost List' sheet")
            cost_sheet = wb["Cost List"]

            # Find utilization rate
            util_row = None
            for r in range(1, cost_sheet.max_row + 1):
                for c in range(1, cost_sheet.max_column + 1):
                    v = cost_sheet.cell(row=r, column=c).value
                    if v and "Average utilization:" in str(v):
                        util_row = r
                        break
                if util_row:
                    break
            
            if util_row is None:
                analysis_logger.log("'Average utilization:' not found in Cost List", "ERROR")
                raise ValueError("Not found 'Average utilization:'")
            
            util_str = cost_sheet.cell(row=util_row, column=11).value
            util_val = _parse_float(str(util_str).replace("%", "")) if util_str is not None else None
            utilization_rate = (util_val / 100.0) if (util_val is not None) else 0.0
            
            if utilization_rate <= 0 or utilization_rate > 1:
                analysis_logger.log(f"Average utilization out of range: {utilization_rate*100:.1f}%", "WARNING")

            # Find Material Price row
            mat_price_row = None
            for r in range(1, cost_sheet.max_row + 1):
                v = cost_sheet.cell(row=r, column=1).value
                if v and "Material Price" in str(v):
                    mat_price_row = r
                    break
            
            if mat_price_row is None:
                analysis_logger.log("'Material Price' row not found in Cost List", "ERROR")
                raise ValueError("No 'Material Price' row")
                messagebox.showerror("Warning", f"Material Price' row not found in Cost List")
            def parse_num(cellv):
                if cellv is None:
                    return 0.0
                s = str(cellv).strip()
                s = ''.join(ch for ch in s.split()[0] if ch.isdigit() or ch in ('.', ','))
                s = s.replace(",", ".")
                try:
                    return float(s or "0.0")
                except Exception:
                    analysis_logger.log(f"Failed to parse number: {cellv}", "WARNING")
                    return 0.0

            rate_per_contour = parse_num(cost_sheet.cell(row=mat_price_row, column=7).value)
            rate_per_marking_length = parse_num(cost_sheet.cell(row=mat_price_row, column=9).value)
            rate_per_defilm_length = parse_num(cost_sheet.cell(row=mat_price_row, column=10).value)

            # Count sheets
            r_idx = 8
            sheets_in_file = 0
            while all_task.cell(row=r_idx, column=4).value is not None:
                v = all_task.cell(row=r_idx, column=4).value
                if isinstance(v, (int, float)):
                    sheets_in_file += int(v)
                r_idx += 1
            total_sheets += sheets_in_file
            analysis_logger.log(f"Found {sheets_in_file} sheets in file", "INFO")

            # Find parts data starting row
            start_row = None
            for r in range(1, cost_sheet.max_row + 1):
                a_val = cost_sheet.cell(row=r, column=1).value
                if a_val and isinstance(a_val, (int, float)):
                    start_row = r
                    break
            
            if start_row is None:
                analysis_logger.log("No starting row found for parts data (col. A – ID)", "ERROR")
                raise ValueError("No starting row found (col. A – ID)")

            # Process parts
            parts_for_group = []
            subnr += 1
            lp = 0
            row = start_row
            parts_count = 0
            
            while row <= cost_sheet.max_row and isinstance(cost_sheet.cell(row=row, column=1).value, (int, float)):
                lp += 1
                part_name = cost_sheet.cell(row=row, column=2).value
                part_qty = cost_sheet.cell(row=row, column=5).value or 0
                weight = parse_num(cost_sheet.cell(row=row, column=6).value)
                contours_qty = parse_num(cost_sheet.cell(row=row, column=7).value)
                cut_length = parse_num(cost_sheet.cell(row=row, column=8).value)
                marking_length = parse_num(cost_sheet.cell(row=row, column=9).value)
                defilm_length = parse_num(cost_sheet.cell(row=row, column=10).value)

                adj_weight = (weight / utilization_rate) if utilization_rate > 0 else weight

                # Calculate base costs with ONLY 7% material margin (automatic)
                base_material_cost = adj_weight * base_price_per_kg
                base_cut_cost = cut_length * base_rate_per_cut_length
                
                # Apply ONLY mandatory 7% minimum material cost increase
                material_cost = adj_weight * base_price_per_kg * 1.07
                cut_cost = cut_length * base_rate_per_cut_length  # NO margin on cutting
                
                base_total_part = base_material_cost + contours_qty * rate_per_contour + base_cut_cost + marking_length * rate_per_marking_length + defilm_length * rate_per_defilm_length
                total_part = material_cost + contours_qty * rate_per_contour + cut_cost + marking_length * rate_per_marking_length + defilm_length * rate_per_defilm_length

                thumbnail_data = None
                all_parts_row = 2 + lp
                if all_parts_row in file_thumbnails:
                    thumbnail_data = file_thumbnails[all_parts_row]

                all_parts.append({
                    'id': lp,
                    'subnr': subnr,
                    'name': part_name,
                    'material': material_name,
                    'thickness': thk_val,
                    'qty': int(part_qty) if isinstance(part_qty, (int, float)) else 0,
                    'cost_per_unit': float(f"{total_part:.2f}"),
                    'base_cost_per_unit': float(f"{base_total_part:.2f}"),
                    'bending_per_unit': 0.0,
                    'additional_per_unit': 0.0,
                    'raw_weight': weight,
                    'contours_qty': contours_qty,
                    'cut_length': cut_length,
                    'marking_length': marking_length,
                    'defilm_length': defilm_length,
                    'adj_weight': adj_weight,
                    'base_price_per_kg': base_price_per_kg,
                    'base_rate_per_cut_length': base_rate_per_cut_length,
                    'base_cut_cost': base_cut_cost,
                    'rate_per_contour': rate_per_contour,
                    'rate_per_marking_length': rate_per_marking_length,
                    'rate_per_defilm_length': rate_per_defilm_length,
                    'thumb_data': thumbnail_data,
                    'calculated_material_margin': avg_file_material_margin,  # Store for later use
                    'calculated_cutting_margin': avg_file_cutting_margin,    # Store for later use
                    'file_name': fname,
                })

                # Log detailed cost breakdown for this part
                analysis_logger.log(f"Part {lp} ({part_name}) cost breakdown:", "INFO")
                analysis_logger.log(f"  Base material: {base_material_cost:.2f} PLN (weight: {adj_weight:.3f} kg @ {base_price_per_kg:.2f} PLN/kg)", "INFO")
                analysis_logger.log(f"  Material with 7%: {material_cost:.2f} PLN", "INFO")
                analysis_logger.log(f"  Cutting: {cut_cost:.2f} PLN (length: {cut_length:.2f} m @ {base_rate_per_cut_length:.2f} PLN/m)", "INFO")
                analysis_logger.log(f"  Contours: {contours_qty * rate_per_contour:.2f} PLN", "INFO")
                analysis_logger.log(f"  Marking: {marking_length * rate_per_marking_length:.2f} PLN", "INFO")
                analysis_logger.log(f"  Defilm: {defilm_length * rate_per_defilm_length:.2f} PLN", "INFO")
                analysis_logger.log(f"  Total per unit: {total_part:.2f} PLN", "INFO")

                parts_for_group.append((part_name, float(f"{total_part:.2f}"),
                                        int(part_qty) if isinstance(part_qty, (int, float)) else 0))
                total_parts_qty += int(part_qty) if isinstance(part_qty, (int, float)) else 0
                parts_count += 1
                row += 1

            analysis_logger.log(f"Processed {parts_count} parts from {fname} with only 7% material margin", "SUCCESS")
            groups.append((material_name, thk_val, parts_for_group))
            
            # Store file margin data FOR SUGGESTION
            file_margins.append({
                'filename': fname,
                'material_margin': avg_file_material_margin,
                'cutting_margin': avg_file_cutting_margin,
                'total_area': file_total_area,
                'total_cutting': file_total_cutting,
                'row_count': row_count
            })

        except Exception as e:
            analysis_logger.log(f"Critical error processing {fname}: {str(e)}", "ERROR")
            messagebox.showerror("Error", f"Error processing file {fname}: {e}")
            return

    # Calculate overall average margins FOR SUGGESTION
    analysis_logger.log("CALCULATING SUGGESTED MARGINS (NOT APPLIED)", "PHASE")
    
    if file_margins:
        total_material_weight = sum(fm['total_area'] for fm in file_margins)
        total_cutting_length = sum(fm['total_cutting'] for fm in file_margins)
        
        if total_material_weight > 0:
            avg_material_margin = sum(fm['material_margin'] * fm['total_area'] for fm in file_margins) / total_material_weight
        else:
            avg_material_margin = 0.0
        
        if total_cutting_length > 0:
            avg_cutting_margin = sum(fm['cutting_margin'] * fm['total_cutting'] for fm in file_margins) / total_cutting_length
        else:
            avg_cutting_margin = 0.0
        
        analysis_logger.log(f"Suggested margins: Material {avg_material_margin:.1f}%, Cutting {avg_cutting_margin:.1f}%", "SUCCESS")
        
        # Auto-populate proposed margin fields with calculated averages
        material_margin_var.set(format_pln(avg_material_margin))
        cutting_margin_var.set(format_pln(avg_cutting_margin))
        
        # Update GUI display
        avg_material_label.config(text=f"{avg_material_margin:.2f}%")
        avg_cutting_label.config(text=f"{avg_cutting_margin:.2f}%")
    else:
        avg_material_margin = 0.0
        avg_cutting_margin = 0.0
        material_margin_var.set("0,00")
        cutting_margin_var.set("0,00")

    analysis_logger.log("CALCULATING OVERHEAD DISTRIBUTION", "PHASE")
    
    # Distribution of overheads per piece
    if total_parts_qty > 0:
        extra_per_part = (tech_per_order + add_costs_order) / total_parts_qty
        op_cost_per_part = (total_sheets * op_cost_per_sheet) / total_parts_qty
        analysis_logger.log(f"Overhead per part: Tech+Add={extra_per_part:.2f}, Op={op_cost_per_part:.2f}", "INFO")
    else:
        extra_per_part = 0.0
        op_cost_per_part = 0.0
        analysis_logger.log("No parts found - overhead is 0", "WARNING")

    for p in all_parts:
        p['cost_per_unit'] += extra_per_part + op_cost_per_part
        p['base_cost_per_unit'] += extra_per_part + op_cost_per_part
        p['cost_per_unit'] = float(f"{p['cost_per_unit']:.2f}")
        p['base_cost_per_unit'] = float(f"{p['base_cost_per_unit']:.2f}")

    # Calculate material costs
    analysis_logger.log("CALCULATING MATERIAL COSTS", "PHASE")
    for p in all_parts:
        material_cost_per_part = p['adj_weight'] * p.get('base_price_per_kg', 0.0) * 1.07  # Include 7% minimum
        total_material_cost += material_cost_per_part * p['qty']
    
    analysis_logger.log(f"Total material cost (with 7% margin): {format_pln(total_material_cost)} PLN", "INFO")
    
    # Update Panel 2 display fields
    update_cost_calculations()

    # Populate treeview
    analysis_logger.log("POPULATING DATA TABLE", "PHASE")
    for i, p in enumerate(all_parts, start=1):
        item_values = (
            i,
            p['subnr'],
            p['name'],
            p['material'],
            f"{p['thickness']}",
            p['qty'],
            format_pln(p['cost_per_unit']),
            "",
            "",
            format_pln(p['adj_weight']),
            format_pln(p['cut_length']),
        )
        opts = {'values': item_values}
        if p['thumb_data']:
            try:
                pil_img = Image.open(io.BytesIO(p['thumb_data']))
                max_w, max_h = 140, 70
                w, h = pil_img.size
                ratio = min(max_w / w, max_h / h, 1.0)
                new_w = int(w * ratio)
                new_h = int(h * ratio)
                pil_img = pil_img.resize((new_w, new_h), Image.LANCZOS)
                thumb = ImageTk.PhotoImage(pil_img)
                thumbnail_imgs.append(thumb)
                opts['image'] = thumb
            except Exception as e:
                analysis_logger.log(f"Failed to create thumbnail: {str(e)}", "WARNING")

        tree.insert('', 'end', **opts)

    # Add total row
    total_order = sum(p['cost_per_unit'] * p['qty'] for p in all_parts)
    SetTotalPricePerOrder(total_order)
    total_row_iid = tree.insert('', 'end', values=('', '', 'Total', '', '', '', format_pln(total_order), '', '', '', ''))
    
    analysis_logger.log(f"Total order value (base + 7% material): {format_pln(total_order)} PLN", "SUCCESS")

    # Create merged groups
    analysis_logger.log("CREATING MERGED GROUPS", "PHASE")
    total_sum = 0.0
    merged_groups = []
    for (mat_name, thk, parts) in groups:
        adj = []
        for (nm, cost, qty) in parts:
            c = float(f"{cost:.2f}")
            adj.append((nm, c, qty))
            total_sum += c * qty
        merged_groups.append((mat_name, thk, adj))
    
    last_groups = merged_groups
    last_total_cost = total_sum
    last_folder_path = folder_path
    
    # Final summary
    analysis_logger.log("ANALYSIS COMPLETED - BASE PRICES + 7% MATERIAL MARGIN", "PHASE")
    analysis_logger.log(f"Total sheets: {total_sheets}", "INFO")
    analysis_logger.log(f"Total parts quantity: {total_parts_qty}", "INFO")
    analysis_logger.log(f"O₂ cutting time: {oxygen_cutting_time:.2f}h", "INFO")
    analysis_logger.log(f"N₂ cutting time: {nitrogen_cutting_time:.2f}h", "INFO")
    analysis_logger.log(f"AL N₂ cutting time: {aluminum_nitrogen_cutting_time:.2f}h", "INFO")
    analysis_logger.log(f"SUGGESTED material margin: {avg_material_margin:.2f}%", "INFO")
    analysis_logger.log(f"SUGGESTED cutting margin: {avg_cutting_margin:.2f}%", "INFO")
    analysis_logger.log(f"Files processed: {len(files)}", "SUCCESS")
    
    messagebox.showinfo("Analysis Complete", 
                       f"XLSX analysis completed!\n\n"
                       f"Current prices: Base + 7% material margin only\n\n"
                       f"Suggested margins (not applied):\n"
                       f"• Material: {avg_material_margin:.1f}%\n"
                       f"• Cutting: {avg_cutting_margin:.1f}%\n\n"
                       f"To apply these margins, click 'UPDATE WITH DYNAMIC MARGINS'")

def get_next_offer_number():
    month_year = datetime.datetime.now().strftime("%m/%Y")
    month_key = datetime.datetime.now().strftime("counter_%Y-%m")
    try:
        response = requests.get(f"https://abacus.jasoncameron.dev/hit/xai_offers/{month_key}")
        if response.status_code == 200:
            counter_value = int(response.json()['value'])
            return f"Laser/{counter_value:04d}/{month_year}"
        else:
            return "Laser/0001/12/2024"  # Fallback
    except Exception:
        return "Laser/0001/12/2024"  # Fallback

# FULL report generation function WITH Excel reports using FORMULAS and Polish notation
def generate_report():
    """Generate complete reports including DOCX and both Excel files with formulas"""
    if not all_parts:
        messagebox.showerror("Error", "No data for the report. First 'Analyze XLSX'.")
        return
    folder_path = folder_var.get().strip() or last_folder_path
    if not folder_path or not os.path.isdir(folder_path):
        messagebox.showerror("Error", "Invalid target folder.")
        return
    customer_name = customer_var.get().strip() or "Client"
    offer_number = offer_var.get().strip()
    if not offer_number:
        offer_number = get_next_offer_number()
        offer_var.set(offer_number)
    offer_date = date_var.get().strip() or datetime.datetime.now().strftime("%Y-%m-%d")
    validity = validity_var.get().strip() or (datetime.datetime.now() + timedelta(days=14)).strftime("%Y-%m-%d")
    logo_path = logo_var.get().strip()
    contact_details = contact_text.get("1.0", tk.END).strip()
    preceding_text = preceding_text_var.get("1.0", tk.END).strip()
    finishing_text = finishing_text_var.get("1.0", tk.END).strip()

    # Create Raporty folder
    raporty_path = os.path.join(folder_path, "Raporty")
    os.makedirs(raporty_path, exist_ok=True)

    # Update all_parts from tree
    tree_items = tree.get_children()
    if len(tree_items) != len(all_parts) + 1:  # +1 for total row
        messagebox.showerror("Error", "Data mismatch between table and parts list.")
        return

    for idx, item in enumerate(tree_items[:-1]):  # Exclude total row
        vals = tree.item(item, 'values')
        if len(vals) < 11:
            vals = vals + ('', '') * (11 - len(vals))
        all_parts[idx]['qty'] = int(vals[5] or 0)
        all_parts[idx]['cost_per_unit'] = _parse_float(vals[6]) or 0.0
        all_parts[idx]['bending_per_unit'] = _parse_float(vals[7]) or 0.0
        all_parts[idx]['additional_per_unit'] = _parse_float(vals[8]) or 0.0

    # Enhanced log file with detailed cost breakdowns
    log_path = os.path.join(raporty_path, "cost_calculation_log.txt")
    with open(log_path, 'w', encoding='utf-8') as log:
        log.write(f"Enhanced Calculation Log with Dynamic Margins - {datetime.datetime.now()}\n")
        log.write("="*80 + "\n")
        log.write(f"Folder: {folder_path}\n")
        log.write(f"Client: {customer_name}\n")
        log.write(f"Offer: {offer_number}\n")
        log.write("="*80 + "\n\n")
        
        log.write("MARGIN CALCULATION SUMMARY\n")
        log.write("-"*40 + "\n")
        log.write("Price sources: materials prices.xlsx, cutting prices.xlsx\n")
        log.write(f"Average material margin calculated: {avg_material_margin:.2f}%\n")
        log.write(f"Average cutting margin calculated: {avg_cutting_margin:.2f}%\n")
        log.write(f"Applied material margin: {_parse_float(material_margin_var.get()):.2f}%\n")
        log.write(f"Applied cutting margin: {_parse_float(cutting_margin_var.get()):.2f}%\n\n")
        
        log.write("Margin Calculation Parameters:\n")
        log.write(f"Material: 250% to 0% margin for areas {_parse_float(min_area_var.get()):.2f} to {_parse_float(max_area_var.get()):.2f} m²\n")
        log.write(f"Cutting: 200% to 0% margin for lengths {_parse_float(min_cutting_var.get()):.0f} to {_parse_float(max_cutting_var.get()):.0f} mm\n\n")
        
        log.write("FILE-BY-FILE MARGIN ANALYSIS\n")
        log.write("-"*40 + "\n")
        for fm in file_margins:
            log.write(f"\nFile: {fm['filename']}\n")
            log.write(f"  Material margin: {fm['material_margin']:.2f}%\n")
            log.write(f"  Cutting margin: {fm['cutting_margin']:.2f}%\n")
            log.write(f"  Total area: {fm['total_area']:.4f} m²\n")
            log.write(f"  Total cutting: {fm['total_cutting']:.1f} mm\n")
            log.write(f"  Rows processed: {fm['row_count']}\n")
        
        log.write("\n" + "="*80 + "\n")
        log.write("DETAILED PART-BY-PART COST BREAKDOWN\n")
        log.write("="*80 + "\n\n")
        
        # Get overhead values
        op_cost_per_sheet = _parse_float(op_cost_entry.get()) or 0.0
        tech_per_order = _parse_float(tech_order_entry.get()) or 0.0
        add_costs_order = _parse_float(add_order_cost_entry.get()) or 0.0
        
        if total_parts_qty > 0:
            extra_per_part = (tech_per_order + add_costs_order) / total_parts_qty
            op_cost_per_part = (total_sheets * op_cost_per_sheet) / total_parts_qty
        else:
            extra_per_part = 0.0
            op_cost_per_part = 0.0
        
        for part in all_parts:
            log.write(f"Part ID: {part['id']} - {part['name']}\n")
            log.write("-"*60 + "\n")
            log.write(f"  File: {part.get('file_name', 'N/A')}\n")
            log.write(f"  Material: {part['material']} {part['thickness']} mm\n")
            log.write(f"  Quantity: {part['qty']} pcs\n")
            log.write(f"  Raw weight: {part.get('raw_weight', 0.0):.3f} kg\n")
            log.write(f"  Adjusted weight: {part.get('adj_weight', 0.0):.3f} kg\n\n")
            
            log.write("  Cost Components:\n")
            mat_cost = part.get('adj_weight', 0.0) * part.get('base_price_per_kg', 0.0) * 1.07
            log.write(f"    Material cost: {mat_cost:.2f} PLN\n")
            log.write(f"      Weight: {part.get('adj_weight', 0.0):.3f} kg\n")
            log.write(f"      Price: {part.get('base_price_per_kg', 0.0):.2f} PLN/kg\n")
            log.write(f"      With 7% margin: {mat_cost:.2f} PLN\n")
            
            cut_cost = part.get('cut_length', 0.0) * part.get('base_rate_per_cut_length', 0.0)
            log.write(f"    Cutting cost: {cut_cost:.2f} PLN\n")
            log.write(f"      Length: {part.get('cut_length', 0.0):.2f} m\n")
            log.write(f"      Rate: {part.get('base_rate_per_cut_length', 0.0):.2f} PLN/m\n")
            
            contour_cost = part.get('contours_qty', 0.0) * part.get('rate_per_contour', 0.0)
            log.write(f"    Contours: {contour_cost:.2f} PLN\n")
            log.write(f"      Quantity: {part.get('contours_qty', 0.0):.0f}\n")
            log.write(f"      Rate: {part.get('rate_per_contour', 0.0):.2f} PLN/pc\n")
            
            marking_cost = part.get('marking_length', 0.0) * part.get('rate_per_marking_length', 0.0)
            log.write(f"    Marking: {marking_cost:.2f} PLN\n")
            
            defilm_cost = part.get('defilm_length', 0.0) * part.get('rate_per_defilm_length', 0.0)
            log.write(f"    Defilm: {defilm_cost:.2f} PLN\n")
            
            log.write(f"    Operational overhead: {op_cost_per_part:.2f} PLN\n")
            log.write(f"    Technology overhead: {extra_per_part:.2f} PLN\n")
            
            log.write(f"\n  Final unit cost: {part['cost_per_unit']:.2f} PLN\n")
            log.write(f"  Total for {part['qty']} pcs: {part['cost_per_unit'] * part['qty']:.2f} PLN\n")
            log.write("\n")

    # Generate DOCX
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    if logo_path and os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(3.0))
        except Exception:
            pass
    if contact_details:
        p = doc.add_paragraph(contact_details)
        for r in p.runs:
            r.bold = False

    doc.add_heading(f"Offer for {customer_name}", level=1)
    p = doc.add_paragraph(f"Offer number: {offer_number}")
    p.runs[0].bold = True
    doc.add_paragraph(f"Offer date: {offer_date}")
    doc.add_paragraph(f"Validity period: {validity}")
    if preceding_text:
        doc.add_paragraph(preceding_text)

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Lp.'
    hdr[1].text = 'Miniatura'
    hdr[2].text = 'Part name'
    hdr[3].text = 'Quantity'
    hdr[4].text = 'Net weight'
    hdr[5].text = 'Cost (PLN)'
    hdr[6].text = 'Total (PLN)'
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '006995')
        tcPr.append(shd)
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True
        run.font.size = Pt(11)

    total = 0.0
    lp = 1
    for mat_name, thk, parts in last_groups:
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = f"Material: {mat_name}, Thickness: {thk} mm"
        row[2].merge(row[6])
        run = row[2].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.italic = True
        for nm, cost_per_unit, qty in parts:
            part = next(p for p in all_parts if p['name'] == nm)
            r = table.add_row().cells
            r[0].text = str(lp)
            # Embed graphic in column 2 (Miniatura)
            if part['thumb_data']:
                try:
                    run = r[1].add_paragraph().add_run()
                    run.add_picture(io.BytesIO(part['thumb_data']))
                except Exception:
                    pass
            r[2].text = str(nm) if nm else "No name"
            r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r[3].paragraphs[0].add_run(f"{int(part['qty'])}  ").font.size = Pt(10)
            r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r[4].paragraphs[0].add_run(f"{format_pln(part['raw_weight'])}  ").font.size = Pt(10)
            r[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r[5].paragraphs[0].add_run(f"{format_pln(part['cost_per_unit'] + part['bending_per_unit'] + part['additional_per_unit'])}  ").font.size = Pt(10)
            row_total = (part['cost_per_unit'] + part['bending_per_unit'] + part['additional_per_unit']) * part['qty']
            r[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r[6].paragraphs[0].add_run(f"{format_pln(row_total)}  ").font.size = Pt(10)
            total += row_total
            lp += 1

    srow = table.add_row().cells
    srow[1].text = ""
    srow[2].text = "Total"
    srow[4].text = ""
    srow[6].text = format_pln(total)
    srow[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in srow[2].paragraphs[0].runs:
        r.bold = True
    for r in srow[6].paragraphs[0].runs:
        r.bold = True

    widths = [Cm(1), Cm(2), Cm(6), Cm(2), Cm(2), Cm(3), Cm(3)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    p = doc.add_paragraph(f"Total cost: {format_pln(total)} PLN")
    p.paragraph_format.space_before = Pt(12)
    for r in p.runs:
        r.font.size = Pt(14)
    
    if finishing_text:
        pf = doc.add_paragraph(finishing_text)
        for r in pf.runs:
            r.font.size = Pt(9)

    current_date = datetime.datetime.now().strftime("%Y%m%d")
    fname = f"Oferta_{sanitize_filename(customer_name) or 'Klient'}_{current_date}_{offer_number.replace('/', '-')}.docx"
    full = os.path.join(raporty_path, fname)
    try:
        doc.save(full)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save DOCX file:\n{e}")
        return

   # Generate Enhanced Cost Report with Margins.xlsx with FORMULAS and Polish notation
    cost_wb = Workbook()
    
    # Set workbook locale to Polish
    cost_wb.iso_dates = False
    
    # Sheet 1: Detailed Cost Breakdown with Thumbnail as column 2
    detail_ws = cost_wb.active
    detail_ws.title = "Szczegółowa kalkulacja"
    
    # Headers with Miniatura as second column - ALL IN POLISH
    headers = [
        "ID", "Miniatura", "Nazwa części", "Materiał", "Grubość [mm]", "Ilość [szt]",
        "Waga jednostkowa [kg]", "Waga skorygowana [kg]",
        "Cena materiału [PLN/kg]", "Koszt materiału [PLN]",
        "Czas cięcia [h]", "Koszt cięcia [PLN]",
        "Koszt operacyjny [PLN]", "Koszt technologii [PLN]",
        "Gięcie [PLN]", "Koszty dodatkowe [PLN]",
        "Koszt jednostkowy [PLN]", "Koszt całkowity [PLN]"
    ]
    
    for col, header in enumerate(headers, 1):
        cell = detail_ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    # Calculate overhead distribution
    if total_parts_qty > 0:
        extra_per_part = (tech_per_order + add_costs_order) / total_parts_qty
        op_cost_per_part = (total_sheets * op_cost_per_sheet) / total_parts_qty
    else:
        extra_per_part = 0.0
        op_cost_per_part = 0.0
    
    # Get TKW rates for calculations
    oxygen_rate_tkw = _parse_float(oxygen_rate_entry_TKW.get()) or 262.50
    nitrogen_rate_tkw = _parse_float(nitrogen_rate_entry_TKW.get()) or 412.50
    al_nitrogen_rate_tkw = _parse_float(al_nitrogen_rate_entry_TKW.get()) or 487.50
    bending_percent_tkw = _parse_float(bending_percent_entry_TKW.get()) or 75.0
    
    # Data for corrected charts
    cost_components = {
        'Materiał': 0.0,
        'Cięcie laserowe': 0.0,
        'Koszty operacyjne': 0.0,
        'Technologia': 0.0,
        'Gięcie': 0.0,
        'Koszty dodatkowe': 0.0
    }
    tkw_total = 0.0
    row_num = 2
    for part in all_parts:
        # Calculate individual cost components
        mat_cost = part['adj_weight'] * part.get('base_price_per_kg', 0.0) 
        
        # Get cutting time per part
        part_cutting_time = part.get('cutting_time', 0.0)
            
        # Calculate TKW cutting cost based on gas type
        gas_type = part.get('gas_type', 'O')
        material_name = part.get('material', '')
        
        if gas_type == 'O':
            cut_cost_tkw = part_cutting_time * oxygen_rate_tkw
        elif gas_type == 'N':
            if 'AL' in material_name.upper():
                cut_cost_tkw = part_cutting_time * al_nitrogen_rate_tkw
            else:
                cut_cost_tkw = part_cutting_time * nitrogen_rate_tkw
        else:
            cut_cost_tkw = part_cutting_time * oxygen_rate_tkw  # Default to O2
        
        bending_cost = part.get('bending_per_unit', 0.0)
        bending_cost_tkw = bending_cost * (bending_percent_tkw / 100.0)
        
        # Accumulate for charts (multiply by quantity for total costs)
        cost_components['Materiał'] += mat_cost * part['qty']
        cost_components['Cięcie laserowe'] += cut_cost_tkw * part['qty']
        cost_components['Koszty operacyjne'] += op_cost_per_part * part['qty']
        cost_components['Technologia'] += extra_per_part * part['qty']
        cost_components['Gięcie'] += bending_cost_tkw * part['qty']
        cost_components['Koszty dodatkowe'] += part.get('additional_per_unit', 0.0) * part['qty']
        
        # Write data with proper column order (Miniatura as column 2)
        detail_ws.cell(row=row_num, column=1, value=part['id'])
        detail_ws.cell(row=row_num, column=2, value='')  # Placeholder for thumbnail
        detail_ws.cell(row=row_num, column=3, value=part['name'])
        detail_ws.cell(row=row_num, column=4, value=part['material'])
        detail_ws.cell(row=row_num, column=5, value=part['thickness'])
        detail_ws.cell(row=row_num, column=6, value=part['qty'])
        
        # Waga jednostkowa [kg]
        cell = detail_ws.cell(row=row_num, column=7, value=float(part.get('raw_weight', 0.0)))
        cell.number_format = '#,##0.00'
        
        # Waga skorygowana [kg]
        cell = detail_ws.cell(row=row_num, column=8, value=float(part.get('adj_weight', 0.0)))
        cell.number_format = '#,##0.00'
        
        # Cena materiału [PLN/kg]
        cell = detail_ws.cell(row=row_num, column=9, value=float(part.get('base_price_per_kg', 0.0)))
        cell.number_format = '#,##0.00'
        
        # Koszt materiału [PLN] - formula
        col_letter = get_column_letter
        cell = detail_ws.cell(row=row_num, column=10, value=f"={col_letter(8)}{row_num}*{col_letter(9)}{row_num}")
        cell.number_format = '#,##0.00'
        
        # Czas cięcia [h]
        cell = detail_ws.cell(row=row_num, column=11, value=float(part_cutting_time))
        cell.number_format = '#,##0.0000'
        
        # Koszt cięcia [PLN] - based on cutting time
        cell = detail_ws.cell(row=row_num, column=12, value=float(cut_cost_tkw))
        cell.number_format = '#,##0.00'
        
        # Koszt operacyjny [PLN]
        cell = detail_ws.cell(row=row_num, column=13, value=float(op_cost_per_part))
        cell.number_format = '#,##0.00'
        
        # Koszt technologii [PLN]
        cell = detail_ws.cell(row=row_num, column=14, value=float(extra_per_part))
        cell.number_format = '#,##0.00'
        
        # Gięcie [PLN]
        cell = detail_ws.cell(row=row_num, column=15, value=float(bending_cost))
        cell.number_format = '#,##0.00'
        
        # Koszty dodatkowe [PLN]
        cell = detail_ws.cell(row=row_num, column=16, value=float(part.get('additional_per_unit', 0.0)))
        cell.number_format = '#,##0.00'
        
        # Koszt jednostkowy [PLN] - sum of all components
        cell = detail_ws.cell(row=row_num, column=17, value=f"=SUM({col_letter(10)}{row_num}:{col_letter(16)}{row_num})")
        cell.number_format = '#,##0.00'
        
        # Koszt całkowity [PLN] - unit cost * quantity
        cell = detail_ws.cell(row=row_num, column=18, value=f"={col_letter(17)}{row_num}*{col_letter(6)}{row_num}")
        cell.number_format = '#,##0.00'
        
        # Add thumbnail in column 2 (B)
        if part.get('thumb_data'):
            try:
                img = OpenpyxlImage(io.BytesIO(part['thumb_data']))
                img.width = 60
                img.height = 40
                detail_ws.add_image(img, f'B{row_num}')
                detail_ws.row_dimensions[row_num].height = 45
            except Exception as e:
                analysis_logger.log(f"Failed to add thumbnail for part {part['id']}: {str(e)}", "WARNING")
        
        # Format cells
        for col in range(1, 19):
            cell = detail_ws.cell(row=row_num, column=col)
            if col >= 4:  # Numeric columns (after material name)
                cell.alignment = Alignment(horizontal="right")
        
        # Calculate TKW total
        tkw_total += (mat_cost + cut_cost_tkw + op_cost_per_part + extra_per_part + bending_cost_tkw) * part['qty']
        row_num += 1
    
    # Add totals row with formulas
    total_row = row_num
    detail_ws.cell(row=total_row, column=3, value="SUMA").font = Font(bold=True)
    
    # Total cost formula
    detail_ws.cell(row=total_row, column=18).font = Font(bold=True)
    cell = detail_ws.cell(row=total_row, column=18, value = f"=SUM({col_letter(18)}2:{col_letter(18)}{row_num-1})")
    cell.number_format = '#,##0.00'
    
    # Autofit columns with specific widths
    column_widths = {
        'A': 8,   # ID
        'B': 12,  # Miniatura (thumbnail)
        'C': 25,  # Nazwa części
        'D': 15,  # Materiał
    }
    
    for col_letter_key, width in column_widths.items():
        detail_ws.column_dimensions[col_letter_key].width = width
    
    # Auto-width for remaining columns
    for column in detail_ws.columns:
        column_letter = column[0].column_letter
        if column_letter not in column_widths:
            max_length = 0
            for cell in column:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            adjusted_width = min(max_length + 2, 30)
            detail_ws.column_dimensions[column_letter].width = adjusted_width
    
    # Sheet 2: Margin Analysis Summary
    margin_ws = cost_wb.create_sheet("Analiza marż")
    
    # Title
    margin_ws['A1'] = "ANALIZA DYNAMICZNYCH MARŻ - PODSUMOWANIE"
    margin_ws['A1'].font = Font(bold=True, size=16)
    margin_ws.merge_cells('A1:F1')
    
    # File-by-file margin breakdown
    margin_ws['A3'] = "Nazwa pliku"
    margin_ws['B3'] = "Marża materiałowa [%]"
    margin_ws['C3'] = "Marża cięcia [%]"
    margin_ws['D3'] = "Powierzchnia całkowita [m²]"
    margin_ws['E3'] = "Długość cięcia całkowita [mm]"
    margin_ws['F3'] = "Liczba przetworzonych wierszy"
    
    for cell in ['A3', 'B3', 'C3', 'D3', 'E3', 'F3']:
        margin_ws[cell].font = Font(bold=True)
        margin_ws[cell].fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    
    row = 4
    for fm in file_margins:
        margin_ws.cell(row=row, column=1, value=fm['filename'])
        cell = margin_ws.cell(row=row, column=2, value=float(fm['material_margin']))
        cell.number_format = '#,##0.00'  
        cell = margin_ws.cell(row=row, column=3, value=float(fm['cutting_margin']))
        cell.number_format = '#,##0.00'  
        cell = margin_ws.cell(row=row, column=4, value=float(fm['total_area']))
        cell.number_format = '#,##0.00'  
        cell = margin_ws.cell(row=row, column=5, value=float(fm['total_cutting']))
        cell.number_format = '#,##0.00'  
        margin_ws.cell(row=row, column=6, value=fm['row_count'])
        row += 1
    
    # Overall averages with formulas
    margin_ws.cell(row=row+1, column=1, value="ŚREDNIE WARTOŚCI").font = Font(bold=True)
    margin_ws.cell(row=row+2, column=1, value="Średnia marża materiałowa:")
    cell = margin_ws.cell(row=row+2, column=2, value=float(avg_material_margin/100))
    cell.font = Font(bold=True, color="008000")
    cell.number_format = '0.00 %'  
    margin_ws.cell(row=row+3, column=1, value="Średnia marża cięcia:")
    cell = margin_ws.cell(row=row+3, column=2, value=float(avg_cutting_margin/100))
    cell.font = Font(bold=True, color="008000")
    cell.number_format = '0.00 %'  
    
    # Autofit columns for margin sheet
    for column_cells in margin_ws.columns:
        max_length = 0
        column_letter = None
        for cell in column_cells:
            if hasattr(cell, 'column_letter'):
                if column_letter is None:
                    column_letter = cell.column_letter
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
        if column_letter:
            adjusted_width = min(max_length + 2, 40)
            margin_ws.column_dimensions[column_letter].width = adjusted_width

    # Sheet 3: Financial Analysis (ANALIZA FINANSOWA ZLECENIA)
    analysis_ws = cost_wb.create_sheet("ANALIZA FINANSOWA ZLECENIA")
    
    # Title
    analysis_ws['A1'] = "ANALIZA FINANSOWA ZLECENIA"
    analysis_ws['A1'].font = Font(bold=True, size=14)
    analysis_ws.merge_cells('A1:C1')
    
    # Cost breakdown table
    analysis_ws['A3'] = "Składnik kosztów"
    analysis_ws['B3'] = "Wartość [PLN]"
    analysis_ws['C3'] = "Udział [%]"
    
    for cell in ['A3', 'B3', 'C3']:
        analysis_ws[cell].font = Font(bold=True)
        analysis_ws[cell].fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    
    # Filter out zero-value components for cleaner display
    total_costs = tkw_total  # Use TKW total instead of sum(cost_components.values())
    active_components = [(k, v) for k, v in cost_components.items() if v > 0]
    
    row = 4
    for name, value in active_components:
        analysis_ws.cell(row=row, column=1, value=name)
        cell = analysis_ws.cell(row=row, column=2, value=float(value))
        cell.number_format = '#,##0.00'  
        # Percentage formula
        cell = analysis_ws.cell(row=row, column=3, value=f"=B{row}/{total_costs}*100")
        cell.number_format = '#,##0.00'  
        row += 1
    
    # Total row with formula
    analysis_ws.cell(row=row, column=1, value="RAZEM").font = Font(bold=True)
    cell = analysis_ws.cell(row=row, column=2, value=f"=SUM(B4:B{row-1})")
    cell.number_format = '#,##0.00'  
    analysis_ws.cell(row=row, column=2).font = Font(bold=True)
    cell = analysis_ws.cell(row=row, column=3, value=100)
    cell.number_format = '#,##0.00'  
    analysis_ws.cell(row=row, column=3).font = Font(bold=True)
    
    # Add separator
    row += 2
    
    # Financial Result
    analysis_ws.cell(row=row, column=1, value="WYNIK FINANSOWY").font = Font(bold=True, size=12)
    row += 2
    
    client_price = total_price_per_order
    
    analysis_ws.cell(row=row, column=1, value="Rzeczywiste koszty TKW:")
    cell = analysis_ws.cell(row=row, column=2, value=float(total_costs))
    cell.number_format = '#,##0.00 PLN'  
    row += 1
    
    analysis_ws.cell(row=row, column=1, value="Cena dla klienta:")
    cell = analysis_ws.cell(row=row, column=2, value=float(client_price))
    cell.number_format = '#,##0.00 PLN'  
    row += 1
    
    analysis_ws.cell(row=row, column=1, value="Marża:")
    # Margin formula
    cell = analysis_ws.cell(row=row, column=2, value=f"=B{row-1}-B{row-2}")
    cell.number_format = '#,##0.00 PLN'  
    analysis_ws[f'B{row}'].font = Font(bold=True, color="008000")
    row += 1
    
    analysis_ws.cell(row=row, column=1, value="Marża [%]:")
    # Margin percentage formula
    cell = analysis_ws.cell(row=row, column=2, value=f"=B{row-1}/B{row-3}")
    cell.number_format = '0.00%'  
    analysis_ws[f'B{row}'].font = Font(bold=True, color="008000")
    
    # Add additional TKW rates information
    row += 2
    analysis_ws.cell(row=row, column=1, value="STAWKI TKW UŻYTE W KALKULACJI:").font = Font(bold=True)
    row += 1
    
    analysis_ws.cell(row=row, column=1, value="O₂ cutting rate TKW:")
    cell = analysis_ws.cell(row=row, column=2, value=float(oxygen_rate_tkw))
    cell.number_format = '#,##0.00 PLN/h'
    row += 1
    
    analysis_ws.cell(row=row, column=1, value="N₂ cutting rate TKW:")
    cell = analysis_ws.cell(row=row, column=2, value=float(nitrogen_rate_tkw))
    cell.number_format = '#,##0.00 PLN/h'
    row += 1
    
    analysis_ws.cell(row=row, column=1, value="AL N₂ cutting rate TKW:")
    cell = analysis_ws.cell(row=row, column=2, value=float(al_nitrogen_rate_tkw))
    cell.number_format = '#,##0.00 PLN/h'
    row += 1
    
    analysis_ws.cell(row=row, column=1, value="Bending percent TKW:")
    cell = analysis_ws.cell(row=row, column=2, value=float(bending_percent_tkw/100))
    cell.number_format = '0.00%'
    
    # Create pie chart for cost structure
    if len(active_components) > 0:
        pie = PieChart()
        pie.title = "Struktura kosztów TKW [%]"
        
        # Prepare data for chart
        labels = Reference(analysis_ws, min_col=1, min_row=4, max_row=3+len(active_components))
        data = Reference(analysis_ws, min_col=3, min_row=3, max_row=3+len(active_components))
        
        pie.add_data(data, titles_from_data=True)
        pie.set_categories(labels)
        pie.width = 20
        pie.height = 21

        # Add data labels
        pie.dataLabels = DataLabelList()
        pie.dataLabels.showPercent = True
        pie.dataLabels.showCatName = True
        
        # Position the chart
        analysis_ws.add_chart(pie, "E3")
    
    # Autofit columns
    for column_cells in analysis_ws.columns:
        max_length = 0
        column_letter = None
        for cell in column_cells:
            if hasattr(cell, 'column_letter'):
                if column_letter is None:
                    column_letter = cell.column_letter
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
        if column_letter and column_letter in ['A', 'B', 'C']:
            adjusted_width = min(max_length + 2, 30)
            analysis_ws.column_dimensions[column_letter].width = adjusted_width

    # Save the enhanced cost report
    try:
        fname = f"Raport_kosztowy_{offer_number.replace('/', '-')}.xlsx"
        cost_wb.save(os.path.join(raporty_path, fname))
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save Enhanced Cost Report: {str(e)}")
    
    # Generate Client report with margins.xlsx
    client_wb = Workbook()
    client_ws = client_wb.active
    client_ws.title = "Offer for client"
    
    # Add header with company info
    client_ws.merge_cells('A1:I1')
    client_ws['A1'] = "LP KONSTAL Sp. z o.o."
    client_ws['A1'].font = Font(bold=True, size=16, color="366092")
    client_ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
    
    # Add offer details
    client_ws.merge_cells('A3:D3')
    client_ws['A3'] = f"OFERTA NR: {offer_number}"
    client_ws['A3'].font = Font(bold=True, size=14)
    
    client_ws.merge_cells('F3:I3')
    client_ws['F3'] = f"Data: {datetime.datetime.now().strftime('%d.%m.%Y')}"
    client_ws['F3'].alignment = Alignment(horizontal="right")
    
    client_ws.merge_cells('A4:D4')
    client_ws['A4'] = f"Dla: {customer_name}"
    client_ws['A4'].font = Font(size=12)
    
    client_ws.merge_cells('F4:I4')
    client_ws['F4'] = f"Ważna do: {validity}"
    client_ws['F4'].alignment = Alignment(horizontal="right")
    
    # Add a separator row
    client_ws.merge_cells('A6:I6')
    client_ws['A6'].fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    client_ws.row_dimensions[6].height = 3
    
     # Column headers for client report
    headers = [
        "ID", "Miniatura", "Part name", "Material", 
        "Thickness [mm]", "Unit weight [kg]", 
        "Quantity [pcs]", "Unit cost [PLN]", "Total cost [PLN]"
    ]
    
    header_row = 8
    for col, header in enumerate(headers, 1):
        cell = client_ws.cell(row=header_row, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        ) 
        
    
    # Add data rows with alternating colors
    data_start_row = header_row + 1
    client_total = 0.0
    
    for idx, part in enumerate(all_parts):
        row_num = data_start_row + idx
        
        # Alternate row colors for better readability
        fill_color = "F2F2F2" if idx % 2 == 0 else "FFFFFF"
        
        unit_total = part['cost_per_unit'] + part['bending_per_unit'] + part['additional_per_unit']
        total_part = unit_total * part['qty']
        client_total += total_part
        
        # ID
        cell = client_ws.cell(row=row_num, column=1, value=part['id'])
        cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
        cell.border = Border(left=Side(style='thin'), right=Side(style='thin'))
        
        
        # ADD IMAGE IN CELL (column 2)
        # Helper functions for image sizing and placement
        EMU_PER_PIXEL = 9525  # Excel drawing units

        def _excel_col_width_to_pixels(col_width: float | None) -> int:
            """
            Approx. convert Excel column width (characters) to pixels.
            8.43 -> ~64 px with Calibri 11. Use standard approximation.
            """
            if not col_width:  # None means default 8.43
                col_width = 8.43
            if col_width <= 0:
                return 0
            # Microsoft’s approximation widely used in openpyxl examples:
            return int(round(col_width * 7 + 5))

        def _points_to_pixels(points: float | None) -> int:
            # Excel row height is in points; 1 pt = 1/72 in; screen ~96 dpi
            if not points:  # None means default ~15 points (~20 px)
                points = 15
            return int(round(points * 96 / 72))

        def add_image_inside_cell(ws, row: int, col: int, img_bytes: bytes, padding_px: int = 2):
            """
            Insert an image anchored to a single cell so it is contained within it
            (top-left at cell corner, bottom-right at cell bottom-right).
            """
            # Compute cell pixel size
            col_letter = get_column_letter(col)
            col_px = _excel_col_width_to_pixels(ws.column_dimensions[col_letter].width)
            row_px = _points_to_pixels(ws.row_dimensions[row].height)

            # Fallback: if the row is too small, set a reasonable height
            if row_px < 60:
                ws.row_dimensions[row].height = 60  # points
                row_px = _points_to_pixels(ws.row_dimensions[row].height)

            # Prepare image
            img = OpenpyxlImage(io.BytesIO(img_bytes))

            # Bound the image to cell interior (with small padding)
            draw_w = max(1, col_px - 2 * padding_px)
            draw_h = max(1, row_px - 2 * padding_px)

            # Build two-cell anchor from cell top-left to cell bottom-right
            # openpyxl uses 0-based row/col in AnchorMarker
            start = AnchorMarker(col=col - 1, row=row - 1,
                                    colOff=padding_px * EMU_PER_PIXEL,
                                    rowOff=padding_px * EMU_PER_PIXEL)
            end = AnchorMarker(col=col - 1, row=row - 1,
                                colOff=(padding_px + draw_w) * EMU_PER_PIXEL,
                                rowOff=(padding_px + draw_h) * EMU_PER_PIXEL)
            img.anchor = TwoCellAnchor(_from=start, to=end, editAs="twoCell")

            ws.add_image(img)

        # Miniatura (column 2)
        cell = client_ws.cell(row=row_num, column=2, value='')
        cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")

        if part.get('thumb_data'):
            try:
                add_image_inside_cell(client_ws, row=row_num, col=2, img_bytes=part['thumb_data'], padding_px=2)
            except Exception as e:
                # optionally log error
                pass

               
        # Other columns
        values = [
            part['name'],
            part['material'],
            part['thickness'],
            f"{part.get('raw_weight', 0.0):.3f}",
            part['qty'],
            f"{unit_total:.2f}",
            f"{total_part:.2f}"
        ]
        
        for col, value in enumerate(values, 3):
            cell = client_ws.cell(row=row_num, column=col, value=value)
            cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
            if col in [5, 6, 8, 9]:  # Numeric columns
                cell.alignment = Alignment(horizontal="right",wrap_text=True)
            if col in [2, 3]:  # text
                cell.alignment = Alignment(horizontal="left",vertical="center", wrap_text=True)
            cell.border = Border(left=Side(style='thin'), right=Side(style='thin'))
    
    # Total row
    total_row = data_start_row + len(all_parts)
    client_ws.merge_cells(f'A{total_row}:F{total_row}')
    cell = client_ws.cell(row=total_row, column=1, value="TOTAL")
    cell.font = Font(bold=True, size=12)
    cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    cell.alignment = Alignment(horizontal="right")
    
    for col in range(7, 9):
        cell = client_ws.cell(row=total_row, column=col, value="")
        cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    
    cell = client_ws.cell(row=total_row, column=9, value=f"{client_total:.2f}")
    cell.font = Font(bold=True, size=12)
    cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    cell.alignment = Alignment(horizontal="right")
    cell.border = Border(top=Side(style='double'), bottom=Side(style='double'))
    
    # Add closing text with disclaimers
    disclaimer_start = total_row + 3
    client_ws.merge_cells(f'A{disclaimer_start}:I{disclaimer_start}')
    client_ws[f'A{disclaimer_start}'] = "IMPLEMENTATION CONDITIONS"
    client_ws[f'A{disclaimer_start}'].font = Font(bold=True, size=11)
    
    disclaimers = [
        "1. Order implementation follows offer acceptance and submission of technical documentation.",
        "2. Delivery time: to be agreed, standard 5-7 working days.",
        "3. Prices do not include transportation costs.",
        "4. Liability exclusions:",
        "   • For errors in the provided technical documentation",
        "   • For defects in the material provided",
        "   • For damages resulting from force majeure",
        "5. Payment: transfer 14 days from the invoice VAT date."
    ]
    
    for idx, text in enumerate(disclaimers):
        row = disclaimer_start + idx + 1
        client_ws.merge_cells(f'A{row}:I{row}')
        client_ws[f'A{row}'] = text
        client_ws[f'A{row}'].font = Font(size=9)
        client_ws[f'A{row}'].alignment = Alignment(wrap_text=True)
    
    # Footer with contact info
    footer_row = disclaimer_start + len(disclaimers) + 3
    client_ws.merge_cells(f'A{footer_row}:I{footer_row}')
    client_ws[f'A{footer_row}'] = "Laser Team | Tel: +48 537 883 393 | Email: laser@konstal.com"
    client_ws[f'A{footer_row}'].font = Font(size=10, italic=True)
    client_ws[f'A{footer_row}'].alignment = Alignment(horizontal="center")
    
    # Set column widths
    column_widths = {
        'A': 8,   # ID
        'B': 20,  # Miniatura
        'C': 35,  # Part name
        'D': 15,  # Material
        'E': 12,  # Thickness
        'F': 18,  # Weight
        'G': 10,  # Quantity
        'H': 18,  # Unit cost
        'I': 18   # Total cost
    }
    
    for col, width in column_widths.items():
        client_ws.column_dimensions[col].width = width
    
    # Add print settings
    client_ws.page_setup.orientation = 'landscape'
    client_ws.page_setup.fitToWidth = 1
    client_ws.page_setup.fitToHeight = 0
    
    # Save the client report
    fname = f"Raport_klienta_{offer_number.replace('/', '-')}.xlsx"
    client_wb.save(os.path.join(raporty_path, fname))
    
    messagebox.showinfo("Sukces", f"Raporty wygenerowane w folderze Raporty!\n\n"
                                  f"Utworzone pliki:\n"
                                  f"• {fname}\n"
                                  f"• Raport kosztów\n"
                                  f"• Raport klienta\n"
                                  f"• cost_calculation_log.txt")

# left buttons (use grid instead of pack)
btn_analyze = ttk.Button(buttons_frame, text="Analyze XLSX", command=analyze_xlsx_folder)
btn_analyze.grid(row=1, column=0, padx=5, pady=5, sticky="we")

btn_report = ttk.Button(buttons_frame, text="Generate report", command=generate_report)
btn_report.grid(row=1, column=1, padx=5, pady=5, sticky="we")

# make columns expand nicely (do once for buttons_frame)
buttons_frame.grid_columnconfigure(0, weight=1)
buttons_frame.grid_columnconfigure(1, weight=1)

# ---- sash setup ----
def set_sash_positions(attempt=1):
    try:
        root.update_idletasks()
        panes = panel_a.panes()
        sash_count = max(len(panes) - 1, 0)
        if sash_count == 0:
            root.after(60, lambda: set_sash_positions(attempt+1)); return

        h = panel_a.winfo_height()
        if h < 400 and attempt < 10:
            root.after(80, lambda: set_sash_positions(attempt+1)); return

        y1 = max(220, int(h * 0.50))
        try: panel_a.sash_place(0, 0, y1)
        except Exception: pass

        if sash_count >= 2:
            y2 = min(h - 200, y1 + panel2_height)
            try: panel_a.sash_place(1, 0, y2)
            except Exception: pass
    except tk.TclError:
        if attempt < 10:
            root.after(80, lambda: set_sash_positions(attempt+1))

root.after_idle(set_sash_positions)

# run
root.geometry("2100x1200")
root.mainloop()